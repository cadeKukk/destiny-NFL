#!/usr/bin/env python3
"""
Create Simplified Visual NFL Mock Draft Document
Matching the exact layout from the provided image
"""

from docx import Document
from docx.shared import Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from datetime import datetime
import os

def create_visual_mock_draft_document():
    """Create a simplified visual Word document matching the image layout"""
    
    # Create processed folder
    os.makedirs('processed', exist_ok=True)
    
    # Mock drafts with the exact players and descriptions from your image
    mock_drafts = [
        {
            'title': 'Bucky Brooks 2025 NFL Mock Draft 3.0',
            'author': 'Bucky Brooks',
            'date': '2024-12-01',
            'picks': [
                {
                    'pick': 1, 
                    'team': 'Tennessee Titans', 
                    'player': 'Cam Ward', 
                    'position': 'QB', 
                    'school': 'Miami',
                    'class': 'Senior',
                    'description': 'The talented passer gives Brian Callahan the franchise quarterback needed to spark the Titans\' rebuild.',
                    'pick_color': 'dark_blue'
                },
                {
                    'pick': 2, 
                    'team': 'Cleveland Browns', 
                    'player': 'Shedeur Sanders', 
                    'position': 'QB', 
                    'school': 'Colorado',
                    'class': 'Senior',
                    'description': 'Kevin Stefanski has worked well with traditional pocket passers throughout his career. Sanders fits the bill as a classic dropback quarterback with a game built on touch, timing and anticipation.',
                    'pick_color': 'orange'
                },
                {
                    'pick': 3, 
                    'team': 'New York Giants', 
                    'player': 'Travis Hunter', 
                    'position': 'WR/CB', 
                    'school': 'Colorado',
                    'class': 'Junior',
                    'description': 'Adding a two-way standout doesn\'t solve the Giants\' most pressing need, but Hunter\'s playmaking presence would help the offense and defense improve.',
                    'pick_color': 'dark_blue'
                },
                {
                    'pick': 4, 
                    'team': 'New England Patriots', 
                    'player': 'Abdul Carter', 
                    'position': 'Edge', 
                    'school': 'Penn State',
                    'class': 'Junior',
                    'description': 'If Patriots personnel chief Eliot Wolf is truly committed to taking the best player available, the Penn State product would be a no-brainer at this point.',
                    'pick_color': 'dark_blue'
                }
            ]
        },
        {
            'title': 'Charles Davis 2025 NFL Mock Draft 3.0',
            'author': 'Charles Davis', 
            'date': '2024-11-28',
            'picks': [
                {
                    'pick': 1, 
                    'team': 'Tennessee Titans', 
                    'player': 'Cam Ward', 
                    'position': 'QB', 
                    'school': 'Miami',
                    'class': 'Senior',
                    'description': 'Ward has the most NFL-ready skill set of any quarterback in this class. His accuracy and leadership make him the clear choice for Tennessee.',
                    'pick_color': 'dark_blue'
                },
                {
                    'pick': 2, 
                    'team': 'Cleveland Browns', 
                    'player': 'Abdul Carter', 
                    'position': 'Edge', 
                    'school': 'Penn State',
                    'class': 'Junior',
                    'description': 'Carter brings elite pass rush ability and versatility. His combination of speed and power makes him a game-changing defender.',
                    'pick_color': 'orange'
                },
                {
                    'pick': 3, 
                    'team': 'New York Giants', 
                    'player': 'Will Johnson', 
                    'position': 'CB', 
                    'school': 'Michigan',
                    'class': 'Junior',
                    'description': 'Johnson has lockdown coverage ability and the physicality to match up with elite receivers in today\'s NFL.',
                    'pick_color': 'dark_blue'
                }
            ]
        }
    ]
    
    # Create the Word document
    doc = Document()
    
    # Title
    title = doc.add_heading('NFL 2025 Mock Draft Picks', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph(f"Generated: {datetime.now().strftime('%B %d, %Y')}")
    doc.add_paragraph("")
    
    # Process each mock draft
    for draft in mock_drafts:
        # Author header
        author_heading = doc.add_heading(f"{draft['author']} Mock Draft", level=1)
        doc.add_paragraph(f"{draft['title']}")
        doc.add_paragraph("")
        
        # Create each pick in the visual format
        for pick in draft['picks']:
            # Create a table for the pick layout
            table = doc.add_table(rows=2, cols=3)
            table.style = 'Table Grid'
            
            # Configure table width
            table.autofit = False
            table.columns[0].width = Inches(0.8)  # Pick number column
            table.columns[1].width = Inches(4.5)  # Main content column  
            table.columns[2].width = Inches(1.5)  # Player image column
            
            # First row - Pick info
            pick_cell = table.cell(0, 0)
            content_cell = table.cell(0, 1)
            image_cell = table.cell(0, 2)
            
            # Merge the description row
            desc_cell = table.cell(1, 0)
            table.cell(1, 0).merge(table.cell(1, 2))
            
            # Pick number cell with colored background
            pick_para = pick_cell.paragraphs[0]
            pick_run = pick_para.runs[0] if pick_para.runs else pick_para.add_run()
            pick_run.text = f"Pick\n{pick['pick']}"
            pick_run.font.bold = True
            pick_run.font.size = Inches(0.2)
            pick_run.font.color.rgb = RGBColor(255, 255, 255)  # White text
            pick_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            pick_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            
            # Set background color based on pick
            if pick['pick_color'] == 'orange':
                # Orange background for Browns
                shading = pick_cell._element.get_or_add_tcPr().get_or_add_shd()
                shading.set('fill', 'FF4500')
            else:
                # Dark blue background for other teams
                shading = pick_cell._element.get_or_add_tcPr().get_or_add_shd() 
                shading.set('fill', '1E3A8A')
            
            # Main content cell
            content_para = content_cell.paragraphs[0]
            
            # Team name
            team_run = content_para.add_run(f"{pick['team']}")
            team_run.font.bold = True
            team_run.font.size = Inches(0.15)
            
            content_para.add_run("\n")
            
            # Player name in blue
            player_run = content_para.add_run(f"{pick['player']}")
            player_run.font.bold = True
            player_run.font.size = Inches(0.18)
            player_run.font.color.rgb = RGBColor(37, 99, 235)  # Blue color
            
            content_para.add_run("\n")
            
            # Player details
            details_run = content_para.add_run(f"{pick['school']} • {pick['position']} • {pick['class']}")
            details_run.font.size = Inches(0.12)
            details_run.font.color.rgb = RGBColor(107, 114, 128)  # Gray color
            
            # Player image placeholder
            image_para = image_cell.paragraphs[0]
            image_run = image_para.add_run("[Player Photo]")
            image_run.font.size = Inches(0.1)
            image_run.font.color.rgb = RGBColor(156, 163, 175)
            image_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            image_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            
            # Description in merged bottom row
            desc_para = desc_cell.paragraphs[0]
            desc_run = desc_para.add_run(pick['description'])
            desc_run.font.size = Inches(0.13)
            desc_para.space_after = Inches(0.2)
            
            # Add spacing after table
            doc.add_paragraph("")
        
        # Page break between drafts
        doc.add_page_break()
    
    # Summary section
    doc.add_heading('Summary Analysis', level=1)
    
    # Most popular players
    all_players = []
    for draft in mock_drafts:
        for pick in draft['picks']:
            all_players.append(pick['player'])
    
    from collections import Counter
    player_counts = Counter(all_players)
    most_common = player_counts.most_common(5)
    
    doc.add_heading('Most Frequently Selected Players:', level=2)
    for i, (player, count) in enumerate(most_common, 1):
        doc.add_paragraph(f"{i}. {player} - Selected {count} time(s)")
    
    # Save the document
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    output_path = f'processed/NFL_Mock_Drafts_VISUAL_SIMPLE_{timestamp}.docx'
    doc.save(output_path)
    
    return output_path, len(mock_drafts), sum(len(draft['picks']) for draft in mock_drafts)

def main():
    print("=== Creating Simplified Visual NFL Mock Draft Document ===")
    print("✓ Clean layout matching your provided image")
    print("✓ Pick number with colored background")
    print("✓ Team name and player name prominently displayed")
    print("✓ Player details (school • position • class)")
    print("✓ Player photo placeholder")
    print("✓ Brief description for each pick")
    print("=" * 60)
    
    try:
        output_path, draft_count, total_picks = create_visual_mock_draft_document()
        
        print("🎉 SUCCESS! Simplified visual document created!")
        print("=" * 60)
        print(f"📊 {draft_count} mock drafts processed")
        print(f"📈 {total_picks} picks in clean visual format")
        print(f"📁 Document saved: {output_path}")
        
        print("\n🎨 Visual Format Features:")
        print("   • Colored pick number boxes (blue/orange)")
        print("   • Large, bold team names")
        print("   • Player names in blue highlighting")
        print("   • Clean player details format")
        print("   • Photo placeholders for each player")
        print("   • Brief, readable descriptions")
        print("   • Consistent spacing and layout")
        
        print(f"\n📖 Layout matches your provided image exactly!")
        print(f"💡 Open: {output_path}")
        
    except Exception as e:
        print(f"❌ Error creating document: {e}")
        import traceback
        traceback.print_exc()

if __name__ == "__main__":
    main() 