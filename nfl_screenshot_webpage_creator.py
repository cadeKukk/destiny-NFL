#!/usr/bin/env python3
"""
NFL Screenshot Webpage Creator
Takes actual screenshots of NFL.com mock draft pages and creates Word documents
with the exact visual layouts from the webpages
"""

import os
import time
from datetime import datetime
from docx import Document
from docx.shared import Inches, RGBColor, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import requests
from selenium import webdriver
from selenium.webdriver.chrome.options import Options
from selenium.webdriver.common.by import By
from selenium.webdriver.support.ui import WebDriverWait
from selenium.webdriver.support import expected_conditions as EC
from PIL import Image

class NFLScreenshotCreator:
    def __init__(self):
        self.setup_selenium()
        
        # All NFL.com mock draft URLs
        self.author_urls = {
            'Bucky Brooks': 'https://www.nfl.com/news/bucky-brooks-2025-nfl-mock-draft-4-0-steelers-land-shedeur-sanders-cowboys-broncos-select-rbs',
            'Daniel Jeremiah': 'https://www.nfl.com/news/daniel-jeremiah-2025-nfl-mock-draft-4-0-broncos-giants-trade-up-steelers-pick-shedeur-sanders',
            'Lance Zierlein': 'https://www.nfl.com/news/lance-zierlein-2025-nfl-mock-draft-4-0-colts-trade-up-for-colston-loveland-saints-go-get-jaxson-dart',
            'Charles Davis': 'https://www.nfl.com/news/charles-davis-2025-nfl-mock-draft-3-0-cam-ward-only-qb-in-round-1-eagles-pick-te-mason-taylor',
            'Eric Edholm': 'https://www.nfl.com/news/eric-edholm-2025-nfl-mock-draft-3-0-four-first-round-quarterbacks-jaguars-take-rb-ashton-jeanty',
            'Dan Parr': 'https://www.nfl.com/news/dan-parr-2025-nfl-mock-draft-2-0-offensive-linemen-dominate-top-10-bears-grab-tight-end-tyler-warren',
            'Chad Reuter': 'https://www.nfl.com/news/seven-round-2025-nfl-mock-draft-patriots-pick-ashton-jeanty-in-round-1-packers-trade-up',
            'Gennaro Filice': 'https://www.nfl.com/news/gennaro-filice-2025-nfl-mock-draft-2-0-rb-ashton-jeanty-goes-top-5-cowboys-jump-for-jalon-walker',
            'Marc Ross': 'https://www.nfl.com/news/marc-ross-2025-nfl-mock-draft-1-0-three-qbs-selected-in-top-10-jets-snag-rb-ashton-jeanty'
        }
        
        os.makedirs('processed/webpage_screenshots', exist_ok=True)

    def setup_selenium(self):
        """Setup Selenium WebDriver for taking screenshots"""
        try:
            chrome_options = Options()
            chrome_options.add_argument('--no-sandbox')
            chrome_options.add_argument('--disable-dev-shm-usage')
            chrome_options.add_argument('--window-size=1600,1200')  # Larger window for better capture
            chrome_options.add_argument('--user-agent=Mozilla/5.0 (Macintosh; Intel Mac OS X 10_15_7) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/120.0.0.0 Safari/537.36')
            chrome_options.add_argument('--disable-blink-features=AutomationControlled')
            
            self.driver = webdriver.Chrome(options=chrome_options)
            print("✓ Selenium WebDriver setup complete")
        except Exception as e:
            print(f"⚠️ Selenium setup failed: {e}")
            print("   Please install ChromeDriver: brew install chromedriver")
            self.driver = None

    def screenshot_webpage_content(self, url, author):
        """Take screenshots of the NFL.com webpage content"""
        if not self.driver:
            print(f"⚠️ No WebDriver available for {author}")
            return []
            
        screenshots = []
        
        try:
            print(f"📸 Capturing screenshots for {author}...")
            
            # Navigate to the page
            self.driver.get(url)
            time.sleep(8)  # Longer wait for page to fully load
            
            # Remove any overlay/cookie banners that might block content
            self.remove_overlays()
            
            # Wait for content to load
            time.sleep(3)
            
            # Get article title screenshot
            title_screenshot = self.screenshot_article_title(author)
            if title_screenshot:
                screenshots.append(title_screenshot)
            
            # Get main content area screenshot
            content_screenshot = self.screenshot_main_content(author)
            if content_screenshot:
                screenshots.append(content_screenshot)
            
            # Find and screenshot individual NFL draft picks using correct selectors
            pick_screenshots = self.screenshot_nfl_draft_picks(author)
            screenshots.extend(pick_screenshots)
            
            # If no individual picks found, take strategic page sections
            if not pick_screenshots:
                print(f"   📄 Taking strategic page sections for {author}...")
                section_screenshots = self.screenshot_strategic_sections(author)
                screenshots.extend(section_screenshots)
                
        except Exception as e:
            print(f"   ⚠️ Error taking screenshots for {author}: {e}")
            
        return screenshots

    def remove_overlays(self):
        """Remove cookie banners and overlays that might block content"""
        try:
            overlay_selectors = [
                '[data-module="CookieBanner"]',
                '.onetrust-banner-sdk',
                '.cookie-banner',
                '[class*="overlay"]',
                '[class*="modal"]',
                '.nfl-banner',
                '[id*="onetrust"]'
            ]
            
            for selector in overlay_selectors:
                try:
                    overlays = self.driver.find_elements(By.CSS_SELECTOR, selector)
                    for overlay in overlays:
                        if overlay.is_displayed():
                            self.driver.execute_script("arguments[0].style.display = 'none';", overlay)
                except:
                    continue
                    
            # Also try to close any visible close buttons
            close_button_selectors = [
                '[aria-label="Close"]',
                '.close-button',
                '[data-dismiss]'
            ]
            
            for selector in close_button_selectors:
                try:
                    buttons = self.driver.find_elements(By.CSS_SELECTOR, selector)
                    for button in buttons:
                        if button.is_displayed():
                            button.click()
                            time.sleep(1)
                except:
                    continue
                    
        except Exception as e:
            print(f"   ⚠️ Error removing overlays: {e}")

    def screenshot_article_title(self, author):
        """Screenshot the article title and header area"""
        try:
            # Look for the NFL article header
            header_selectors = [
                '.nfl-c-article__header',
                '.nfl-c-article__title',
                'h1',
                '.article-header'
            ]
            
            for selector in header_selectors:
                try:
                    header_element = self.driver.find_element(By.CSS_SELECTOR, selector)
                    if header_element.is_displayed():
                        # Scroll to element and ensure it's visible
                        self.driver.execute_script("arguments[0].scrollIntoView({block: 'center'});", header_element)
                        time.sleep(2)
                        
                        # Take screenshot with more padding
                        screenshot_path = f"processed/webpage_screenshots/{author}_article_title.png"
                        
                        # Get element location and size, add padding
                        location = header_element.location
                        size = header_element.size
                        
                        # Take full page screenshot first
                        temp_screenshot = f"processed/webpage_screenshots/temp_full_{author}.png"
                        self.driver.save_screenshot(temp_screenshot)
                        
                        # Crop to show header with padding
                        with Image.open(temp_screenshot) as img:
                            left = max(0, location['x'] - 50)
                            top = max(0, location['y'] - 50)
                            right = min(img.width, location['x'] + size['width'] + 50)
                            bottom = min(img.height, location['y'] + size['height'] + 200)  # More bottom padding
                            
                            cropped = img.crop((left, top, right, bottom))
                            cropped.save(screenshot_path)
                        
                        # Clean up temp file
                        os.remove(temp_screenshot)
                        
                        print(f"   ✓ Article title screenshot: {author}_article_title.png")
                        return screenshot_path
                except:
                    continue
                    
        except Exception as e:
            print(f"   ⚠️ Could not capture article title for {author}: {e}")
            
        return None

    def screenshot_main_content(self, author):
        """Screenshot the main article content area"""
        try:
            # Look for the main article body
            content_selectors = [
                '.nfl-c-article__body',
                '.nfl-c-article__container',
                '.article-body',
                'main'
            ]
            
            for selector in content_selectors:
                try:
                    content_element = self.driver.find_element(By.CSS_SELECTOR, selector)
                    if content_element.is_displayed():
                        # Scroll to element
                        self.driver.execute_script("arguments[0].scrollIntoView({block: 'start'});", content_element)
                        time.sleep(2)
                        
                        # Take screenshot
                        screenshot_path = f"processed/webpage_screenshots/{author}_main_content.png"
                        content_element.screenshot(screenshot_path)
                        
                        print(f"   ✓ Main content screenshot: {author}_main_content.png")
                        return screenshot_path
                except:
                    continue
                    
        except Exception as e:
            print(f"   ⚠️ Could not capture main content for {author}: {e}")
            
        return None

    def screenshot_nfl_draft_picks(self, author):
        """Screenshot individual NFL draft picks using correct selectors"""
        screenshots = []
        
        try:
            # Wait for content to load
            WebDriverWait(self.driver, 10).until(
                EC.presence_of_element_located((By.CSS_SELECTOR, 'article'))
            )
            
            # Look for NFL.com specific draft pick elements
            pick_selectors = [
                '.nfl-o-ranked-item.nfl-o-ranked-item--side-by-side',  # Primary NFL.com selector
                '.nfl-o-ranked-item',  # Fallback
                '[class*="ranked-item"]',  # Any ranked item
                '.mock-draft-pick',  # Generic mock draft pick
                '.draft-pick'  # Generic draft pick
            ]
            
            pick_elements = []
            for selector in pick_selectors:
                try:
                    elements = self.driver.find_elements(By.CSS_SELECTOR, selector)
                    if elements:
                        pick_elements = elements[:15]  # Take first 15 picks
                        print(f"   📋 Found {len(pick_elements)} pick elements using selector: {selector}")
                        break
                except:
                    continue
            
            if pick_elements:
                for i, pick_element in enumerate(pick_elements, 1):
                    try:
                        # Scroll to the pick with more context
                        self.driver.execute_script("arguments[0].scrollIntoView({block: 'center'});", pick_element)
                        time.sleep(2)
                        
                        # Take screenshot with additional context around the pick
                        screenshot_path = f"processed/webpage_screenshots/{author}_draft_pick_{i}.png"
                        
                        # Get element location and size for expanded capture
                        location = pick_element.location
                        size = pick_element.size
                        
                        # Take full page screenshot
                        temp_screenshot = f"processed/webpage_screenshots/temp_pick_{author}_{i}.png"
                        self.driver.save_screenshot(temp_screenshot)
                        
                        # Crop with generous padding to show complete pick context
                        with Image.open(temp_screenshot) as img:
                            left = max(0, location['x'] - 100)
                            top = max(0, location['y'] - 50)
                            right = min(img.width, location['x'] + size['width'] + 100)
                            bottom = min(img.height, location['y'] + size['height'] + 100)
                            
                            # Ensure minimum screenshot size for readability
                            if right - left < 800:
                                center_x = (left + right) // 2
                                left = max(0, center_x - 400)
                                right = min(img.width, center_x + 400)
                            
                            if bottom - top < 400:
                                center_y = (top + bottom) // 2
                                top = max(0, center_y - 200)
                                bottom = min(img.height, center_y + 200)
                            
                            cropped = img.crop((left, top, right, bottom))
                            cropped.save(screenshot_path)
                        
                        # Clean up temp file
                        os.remove(temp_screenshot)
                        
                        screenshots.append(screenshot_path)
                        print(f"   ✓ Draft pick {i} screenshot captured")
                        
                    except Exception as e:
                        print(f"   ⚠️ Error capturing pick {i}: {e}")
                        continue
            else:
                print(f"   ⚠️ No draft pick elements found for {author}")
                
        except Exception as e:
            print(f"   ⚠️ Error finding draft picks for {author}: {e}")
            
        return screenshots

    def screenshot_strategic_sections(self, author):
        """Take strategic screenshots of page sections when individual picks aren't found"""
        screenshots = []
        
        try:
            # Scroll to top first
            self.driver.execute_script("window.scrollTo(0, 0);")
            time.sleep(2)
            
            # Get total page height
            total_height = self.driver.execute_script("return document.body.scrollHeight")
            viewport_height = self.driver.execute_script("return window.innerHeight")
            
            # Take screenshots in strategic sections (smaller chunks for better quality)
            section_height = viewport_height // 2  # Half viewport for better detail
            sections = max(1, total_height // section_height)
            
            for i in range(min(sections, 8)):  # Limit to 8 sections max
                scroll_position = i * section_height
                
                # Scroll to position
                self.driver.execute_script(f"window.scrollTo(0, {scroll_position});")
                time.sleep(3)  # Longer wait for content to load
                
                # Take screenshot
                screenshot_path = f"processed/webpage_screenshots/{author}_content_section_{i+1}.png"
                self.driver.save_screenshot(screenshot_path)
                
                screenshots.append(screenshot_path)
                print(f"   ✓ Content section {i+1} screenshot captured")
                
        except Exception as e:
            print(f"   ⚠️ Error taking strategic screenshots for {author}: {e}")
            
        return screenshots

    def create_screenshot_document(self, all_screenshots):
        """Create Word document with all webpage screenshots"""
        print("📄 Creating Word document with webpage screenshots...")
        
        doc = Document()
        
        # Set margins
        sections = doc.sections
        for section in sections:
            section.top_margin = Inches(0.5)
            section.bottom_margin = Inches(0.5)
            section.left_margin = Inches(0.5)
            section.right_margin = Inches(0.5)
        
        # Title
        title = doc.add_heading('NFL 2025 Mock Draft Collection', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_run = title.runs[0]
        title_run.font.size = Pt(24)
        title_run.font.color.rgb = RGBColor(0, 53, 148)
        
        # Subtitle
        subtitle = doc.add_paragraph(f'Exact NFL.com Webpage Screenshots - Generated: {datetime.now().strftime("%B %d, %Y")}')
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        subtitle_run = subtitle.runs[0]
        subtitle_run.font.size = Pt(12)
        subtitle_run.font.color.rgb = RGBColor(107, 114, 128)
        
        # Group screenshots by author
        authors_screenshots = {}
        for screenshot_data in all_screenshots:
            author = screenshot_data['author']
            if author not in authors_screenshots:
                authors_screenshots[author] = []
            authors_screenshots[author].extend(screenshot_data['screenshots'])
        
        # Add each author's screenshots
        for author, screenshots in authors_screenshots.items():
            
            # Author header
            doc.add_page_break()
            
            author_para = doc.add_paragraph()
            author_run = author_para.add_run(f"{author} - NFL.com Mock Draft")
            author_run.font.size = Pt(20)
            author_run.font.bold = True
            author_run.font.color.rgb = RGBColor(0, 53, 148)
            
            # Add each screenshot
            for screenshot_path in screenshots:
                try:
                    if os.path.exists(screenshot_path):
                        # Add with appropriate width (maintain aspect ratio)
                        doc.add_picture(screenshot_path, width=Inches(7.5))
                        
                        # Center the image
                        last_paragraph = doc.paragraphs[-1]
                        last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        last_paragraph.space_before = Pt(8)
                        last_paragraph.space_after = Pt(8)
                        
                        # Add caption with filename
                        caption = doc.add_paragraph(f"📸 {os.path.basename(screenshot_path)}")
                        caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        caption_run = caption.runs[0]
                        caption_run.font.size = Pt(8)
                        caption_run.font.color.rgb = RGBColor(107, 114, 128)
                        
                except Exception as e:
                    print(f"⚠️ Could not add screenshot {screenshot_path}: {e}")
                    
                    # Add text fallback
                    fallback_para = doc.add_paragraph(f"Screenshot: {os.path.basename(screenshot_path)}")
                    fallback_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Save document
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        output_path = f'processed/NFL_WEBPAGE_SCREENSHOTS_ENHANCED_{timestamp}.docx'
        doc.save(output_path)
        
        print(f"✓ Enhanced screenshot document saved: {output_path}")
        return output_path

    def cleanup(self):
        """Clean up resources"""
        if self.driver:
            self.driver.quit()

def main():
    print("=== NFL Webpage Screenshot Creator (Enhanced) ===")
    print("📸 Taking actual screenshots of NFL.com pages")
    print("🎯 Targeting draft pick content and player information")
    print("📄 Creating Word document with exact layouts")
    print("✓ All target authors included")
    print("=" * 55)
    
    creator = NFLScreenshotCreator()
    
    if not creator.driver:
        print("❌ Cannot proceed without WebDriver")
        return
    
    try:
        all_screenshots = []
        
        # Screenshot each author's webpage
        for author, url in creator.author_urls.items():
            screenshots = creator.screenshot_webpage_content(url, author)
            
            all_screenshots.append({
                'author': author,
                'screenshots': screenshots
            })
        
        # Create Word document with screenshots
        if all_screenshots:
            output_path = creator.create_screenshot_document(all_screenshots)
            
            print(f"\n🎉 SUCCESS! Enhanced NFL.com webpage screenshots captured!")
            print("=" * 55)
            print(f"📁 Document: {output_path}")
            print(f"📸 Screenshots: processed/webpage_screenshots/")
            
            total_screenshots = sum(len(author_data['screenshots']) for author_data in all_screenshots)
            print(f"\n📊 Summary:")
            print(f"   • {len(creator.author_urls)} authors processed")
            print(f"   • {total_screenshots} total screenshots captured")
            print(f"   • Enhanced capture with player information")
            print(f"   • Exact NFL.com visual layouts preserved")
            print(f"   • Ready for viewing in Word")
        else:
            print("❌ No screenshots could be captured!")
    
    finally:
        creator.cleanup()

if __name__ == "__main__":
    main() 