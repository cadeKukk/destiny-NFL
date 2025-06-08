#!/usr/bin/env python3
"""
Enhanced NFL Mock Draft Scraper Runner
Generates detailed Word documents with individual picks, player images, and analysis
"""

import sys
import os

# Add current directory to path
sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))

def main():
    try:
        print("=== Enhanced NFL Mock Draft Scraper ===")
        print("This version includes:")
        print("✓ Individual player picks for each draft")
        print("✓ Player images (when available)")
        print("✓ Author analysis and reasoning for each pick")
        print("✓ Enhanced Word document formatting")
        print("=" * 50)
        
        from enhanced_nfl_scraper import EnhancedNFLMockDraftScraper
        
        url = "https://www.nfl.com/news/bucky-brooks-2025-nfl-mock-draft-3-0-browns-take-shedeur-sanders-two-running-backs-in-top-10-picks"
        
        scraper = EnhancedNFLMockDraftScraper()
        mock_drafts = scraper.run(url)
        
        print("\n" + "=" * 50)
        print("🎉 ENHANCED SCRAPING COMPLETED! 🎉")
        print("=" * 50)
        
        if mock_drafts:
            print(f"📊 Successfully processed {len(mock_drafts)} mock drafts")
            print("\nDetailed breakdown:")
            
            total_picks_with_analysis = 0
            for i, draft in enumerate(mock_drafts, 1):
                picks_count = len(draft['picks'])
                total_picks_with_analysis += picks_count
                print(f"{i}. {draft['author']}")
                print(f"   📝 {picks_count} picks with detailed analysis")
                print(f"   📅 Published: {draft['date']}")
                
                # Show sample pick
                if draft['picks']:
                    sample = draft['picks'][0]
                    player = sample.get('player', 'N/A')
                    team = sample.get('team', 'N/A')
                    print(f"   🏈 Sample: Pick #{sample.get('pick', 'N/A')} - {team} selects {player}")
                print()
            
            print(f"📈 Total individual picks analyzed: {total_picks_with_analysis}")
            print("\n📁 Output files created:")
            print("   • Enhanced Word document in 'processed' folder")
            print("   • Player images in 'processed/images' folder")
            
            print("\n🔍 What's included in the Word document:")
            print("   • Complete table of contents")
            print("   • Individual sections for each author")
            print("   • Detailed breakdown of every pick")
            print("   • Player information (name, position, school)")
            print("   • Author's reasoning for each selection")
            print("   • Player images (when available)")
            print("   • Summary analysis of most popular players")
            
        else:
            print("❌ No mock drafts were found or processed.")
            
        print("\n💡 Tip: Open the Word document to see the detailed analysis!")
        
    except Exception as e:
        print(f"❌ ERROR: {e}")
        import traceback
        print("\nFull error details:")
        traceback.print_exc()
        
        print("\n🔧 Creating fallback document with sample data...")
        create_enhanced_fallback()

def create_enhanced_fallback():
    """Create enhanced fallback document with sample detailed picks"""
    try:
        from docx import Document
        from docx.shared import Inches
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from datetime import datetime
        import os
        
        os.makedirs('processed', exist_ok=True)
        
        # Enhanced sample data with detailed analysis
        sample_drafts = [
            {
                'title': 'Bucky Brooks 2025 NFL Mock Draft 3.0 - Browns Take Shedeur Sanders',
                'author': 'Bucky Brooks',
                'date': '2024-12-01',
                'url': 'https://www.nfl.com/news/bucky-brooks-2025-nfl-mock-draft-3-0-browns-take-shedeur-sanders-two-running-backs-in-top-10-picks',
                'picks': [
                    {
                        'pick': 1, 'team': 'Tennessee Titans', 'player': 'Cam Ward', 'position': 'QB', 'school': 'Miami',
                        'reasoning': 'Ward has shown exceptional arm talent and poise in the pocket throughout his college career. His ability to make throws from multiple arm angles and his mobility make him the perfect fit for a Titans offense that needs a dynamic playmaker at quarterback. Ward\'s leadership qualities and clutch gene separate him from other prospects in this class.'
                    },
                    {
                        'pick': 2, 'team': 'Cleveland Browns', 'player': 'Shedeur Sanders', 'position': 'QB', 'school': 'Colorado',
                        'reasoning': 'Sanders brings elite accuracy and football IQ to Cleveland. His ability to process defenses quickly and make precise throws in tight windows is exactly what the Browns need. Sanders has shown he can perform under pressure and has the mental toughness to succeed in the NFL immediately.'
                    },
                    {
                        'pick': 3, 'team': 'New York Giants', 'player': 'Travis Hunter', 'position': 'CB/WR', 'school': 'Colorado',
                        'reasoning': 'Hunter is a generational talent who can impact the game on both sides of the ball. His rare combination of size, speed, and ball skills makes him invaluable. The Giants can use him as a shutdown corner while also utilizing his receiving abilities in offensive packages, creating matchup nightmares for opponents.'
                    },
                    {
                        'pick': 4, 'team': 'New England Patriots', 'player': 'Ashton Jeanty', 'position': 'RB', 'school': 'Boise State',
                        'reasoning': 'Jeanty has been the most dominant offensive player in college football this season. His combination of power, speed, and vision reminds scouts of elite NFL backs. New England needs an offensive spark, and Jeanty can provide that immediately while taking pressure off their quarterback.'
                    },
                    {
                        'pick': 5, 'team': 'Jacksonville Jaguars', 'player': 'Tetairoa McMillan', 'position': 'WR', 'school': 'Arizona',
                        'reasoning': 'McMillan has the size and athleticism to be a true #1 receiver in the NFL. His ability to win contested catches and create separation makes him an ideal target for Trevor Lawrence. The Jaguars need a reliable receiving threat, and McMillan fits perfectly into their offensive system.'
                    }
                ]
            },
            {
                'title': 'Charles Davis 2025 Mock Draft Analysis',
                'author': 'Charles Davis',
                'date': '2024-11-28',
                'url': 'https://www.nfl.com/news/charles-davis-mock-draft',
                'picks': [
                    {
                        'pick': 1, 'team': 'Tennessee Titans', 'player': 'Shedeur Sanders', 'position': 'QB', 'school': 'Colorado',
                        'reasoning': 'Davis believes Sanders has the most NFL-ready skill set of any quarterback in this class. His accuracy on intermediate routes and ability to read defenses pre-snap make him the safest pick for a Titans team that needs immediate help at quarterback.'
                    },
                    {
                        'pick': 2, 'team': 'Cleveland Browns', 'player': 'Cam Ward', 'position': 'QB', 'school': 'Miami',
                        'reasoning': 'Ward\'s arm strength and mobility give Cleveland the dynamic playmaker they\'ve been missing. Davis sees Ward as having the highest ceiling of any quarterback in the draft, with the ability to make throws that few others can attempt.'
                    },
                    {
                        'pick': 3, 'team': 'New York Giants', 'player': 'Abdul Carter', 'position': 'EDGE', 'school': 'Penn State',
                        'reasoning': 'Carter has shown elite pass rush ability and the versatility to play multiple positions along the defensive line. Davis believes the Giants need to prioritize pass rush, and Carter has the potential to be a perennial Pro Bowl player.'
                    }
                ]
            }
        ]
        
        # Create enhanced document
        doc = Document()
        title = doc.add_heading('NFL 2025 Mock Draft Analysis - Enhanced Fallback Edition', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        doc.add_paragraph(f"Generated on: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
        doc.add_paragraph("Note: This is a fallback document with sample enhanced data showing the format.")
        
        total_picks = sum(len(draft['picks']) for draft in sample_drafts)
        doc.add_paragraph(f"Total Individual Picks Detailed: {total_picks}")
        
        # Table of contents
        doc.add_heading('Mock Drafts Included:', level=1)
        for i, draft in enumerate(sample_drafts, 1):
            picks_count = len(draft['picks'])
            doc.add_paragraph(f"{i}. {draft['author']} - {picks_count} picks analyzed")
            doc.add_paragraph(f"   {draft['title']}")
        
        doc.add_page_break()
        
        # Detailed mock drafts
        for draft in sample_drafts:
            doc.add_heading(f"{draft['author']} - Complete Mock Draft Analysis", level=1)
            doc.add_paragraph(f"Title: {draft['title']}")
            doc.add_paragraph(f"Published: {draft['date']}")
            
            doc.add_heading(f'Individual Player Selections ({len(draft["picks"])} picks):', level=2)
            
            for pick in draft['picks']:
                # Pick header
                pick_title = f"Pick #{pick['pick']}: {pick['team']}"
                doc.add_heading(pick_title, level=3)
                
                # Player details
                player_info = f"SELECTION: {pick['player']} | {pick['position']} | {pick['school']}"
                player_para = doc.add_paragraph(player_info)
                player_para.bold = True
                
                # Analysis
                doc.add_heading('Draft Analysis:', level=4)
                doc.add_paragraph(pick['reasoning'])
                
                doc.add_paragraph("")  # Spacing
            
            doc.add_page_break()
        
        # Save
        output_path = os.path.join('processed', f'NFL_Mock_Drafts_ENHANCED_FALLBACK_{datetime.now().strftime("%Y%m%d_%H%M%S")}.docx')
        doc.save(output_path)
        print(f"✅ Enhanced fallback document created: {output_path}")
        
    except Exception as e:
        print(f"❌ Failed to create enhanced fallback: {e}")

if __name__ == "__main__":
    main() 