#!/usr/bin/env python3
"""
NFL Player Ranking Analyzer Enhanced
Analyzes NFL.com mock draft pages to create a ranked list of players by frequency
Enhanced version with better player name extraction
"""

import os
import time
from datetime import datetime
from collections import Counter
from docx import Document
from docx.shared import Inches, RGBColor, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from selenium import webdriver
from selenium.webdriver.chrome.options import Options
from selenium.webdriver.common.by import By
from selenium.webdriver.support.ui import WebDriverWait
from selenium.webdriver.support import expected_conditions as EC
import re

class NFLPlayerRankingAnalyzerEnhanced:
    def __init__(self):
        self.setup_selenium()
        
        # UPDATED URLs - REMOVED Lance Zierlein and Chad Reuter as requested
        self.author_urls = {
            'Bucky Brooks': 'https://www.nfl.com/news/bucky-brooks-2025-nfl-mock-draft-4-0-steelers-land-shedeur-sanders-cowboys-broncos-select-rbs',
            'Daniel Jeremiah': 'https://www.nfl.com/news/daniel-jeremiah-2025-nfl-mock-draft-4-0',
            'Charles Davis': 'https://www.nfl.com/news/charles-davis-2025-nfl-mock-draft-3-0-cam-ward-only-qb-in-round-1-eagles-pick-te-mason-taylor',
            'Eric Edholm': 'https://www.nfl.com/news/eric-edholm-2025-nfl-mock-draft-3-0-four-first-round-quarterbacks-jaguars-take-rb-ashton-jeanty',
            'Dan Parr': 'https://www.nfl.com/news/dan-parr-2025-nfl-mock-draft-2-0-offensive-linemen-dominate-top-10-bears-grab-tight-end-tyler-warren',
            'Gennaro Filice': 'https://www.nfl.com/news/gennaro-filice-2025-nfl-mock-draft-2-0-rb-ashton-jeanty-goes-top-5-cowboys-jump-for-jalon-walker',
            'Marc Ross': 'https://www.nfl.com/news/marc-ross-2025-nfl-mock-draft-1-0-three-qbs-selected-in-top-10-jets-snag-rb-ashton-jeanty'
        }
        
        self.all_players = []
        self.player_selections = {}  # {player_name: {author: pick_number}}

    def setup_selenium(self):
        """Setup Selenium WebDriver"""
        try:
            chrome_options = Options()
            chrome_options.add_argument('--no-sandbox')
            chrome_options.add_argument('--disable-dev-shm-usage')
            chrome_options.add_argument('--window-size=1800,1400')
            chrome_options.add_argument('--user-agent=Mozilla/5.0 (Macintosh; Intel Mac OS X 10_15_7) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/120.0.0.0 Safari/537.36')
            chrome_options.add_argument('--disable-blink-features=AutomationControlled')
            chrome_options.add_argument('--headless')  # Run headless for faster processing
            
            self.driver = webdriver.Chrome(options=chrome_options)
            print("✓ Selenium WebDriver setup complete")
        except Exception as e:
            print(f"⚠️ Selenium setup failed: {e}")
            self.driver = None

    def extract_players_from_author(self, url, author):
        """Extract player names from an author's mock draft"""
        if not self.driver:
            return []
            
        players = []
        
        try:
            print(f"🔍 Analyzing {author}'s mock draft...")
            
            self.driver.get(url)
            time.sleep(10)  # Increased wait time
            
            # Remove overlays
            self.remove_overlays()
            time.sleep(2)
            
            # Enhanced element selection - look for different patterns
            pick_elements = self.find_draft_pick_elements()
            
            if pick_elements:
                print(f"   📋 Found {len(pick_elements)} draft picks for {author}")
                
                for i, pick_element in enumerate(pick_elements[:32], 1):  # Top 32 picks
                    try:
                        # Extract player name from the pick element
                        player_name = self.extract_player_name_comprehensive(pick_element, i)
                        
                        if player_name:
                            players.append({
                                'name': player_name,
                                'pick': i,
                                'author': author
                            })
                            print(f"   ✓ Pick {i}: {player_name}")
                            
                            # Track in master list
                            if player_name not in self.player_selections:
                                self.player_selections[player_name] = {}
                            self.player_selections[player_name][author] = i
                        else:
                            print(f"   ⚠️ No player name found for pick {i}")
                        
                    except Exception as e:
                        print(f"   ⚠️ Error extracting pick {i}: {e}")
                        continue
            else:
                print(f"   ⚠️ No draft pick elements found for {author}")
                
        except Exception as e:
            print(f"   ⚠️ Error analyzing {author}: {e}")
            
        return players

    def find_draft_pick_elements(self):
        """Find draft pick elements using multiple strategies"""
        pick_elements = []
        
        # Try different selectors
        selectors = [
            '.nfl-o-ranked-item.nfl-o-ranked-item--side-by-side',
            '.nfl-o-ranked-item',
            '[class*="ranked-item"]',
            '.mock-draft-pick',
            '.draft-pick'
        ]
        
        for selector in selectors:
            try:
                elements = self.driver.find_elements(By.CSS_SELECTOR, selector)
                if elements:
                    pick_elements = elements[:32]  # Top 32 picks
                    print(f"   ✓ Using selector: {selector}")
                    break
            except:
                continue
        
        return pick_elements

    def extract_player_name_comprehensive(self, pick_element, pick_number):
        """Enhanced player name extraction using multiple methods"""
        try:
            # Method 1: Look for specific NFL.com player name patterns
            player_name = self.extract_from_title_elements(pick_element)
            if player_name:
                return player_name
            
            # Method 2: Look for names in the full text using patterns
            player_name = self.extract_from_text_patterns(pick_element)
            if player_name:
                return player_name
            
            # Method 3: Manual extraction for known patterns
            player_name = self.extract_from_manual_patterns(pick_element, pick_number)
            if player_name:
                return player_name
                
        except Exception as e:
            print(f"   ⚠️ Error in comprehensive extraction: {e}")
            
        return None

    def extract_from_title_elements(self, pick_element):
        """Extract from title-like elements"""
        try:
            # Look for various title elements
            title_selectors = [
                '.nfl-o-ranked-item__title',
                '[class*="title"]',
                'h3',
                'h4',
                'h5',
                '.player-name',
                '.name'
            ]
            
            for selector in title_selectors:
                try:
                    elements = pick_element.find_elements(By.CSS_SELECTOR, selector)
                    for element in elements:
                        text = element.get_attribute('textContent').strip()
                        if text:
                            cleaned = self.clean_and_validate_name(text)
                            if cleaned:
                                return cleaned
                except:
                    continue
                    
        except:
            pass
            
        return None

    def extract_from_text_patterns(self, pick_element):
        """Extract using text pattern matching"""
        try:
            full_text = pick_element.get_attribute('textContent')
            if not full_text:
                return None
            
            # Pattern 1: Look for "Player Name, Position, School" patterns
            patterns = [
                r'([A-Z][a-z]+\s+[A-Z][a-z]+),\s*(?:QB|RB|WR|TE|OL|DL|LB|CB|S)',
                r'([A-Z][a-z]+\s+[A-Z][a-z]+)\s*\|\s*(?:QB|RB|WR|TE|OL|DL|LB|CB|S)',
                r'([A-Z][a-z]+\s+[A-Z][a-z]+)\s*-\s*(?:QB|RB|WR|TE|OL|DL|LB|CB|S)',
                r'([A-Z][a-z]+\s+[A-Z][a-z]+)\s*(?:QB|RB|WR|TE|OL|DL|LB|CB|S)',
                r'(?:select|draft)s?\s+([A-Z][a-z]+\s+[A-Z][a-z]+)',
                r'([A-Z][a-z]+\s+[A-Z][a-z]+)\s*(?:from|at)\s+[A-Z]'
            ]
            
            for pattern in patterns:
                matches = re.findall(pattern, full_text, re.IGNORECASE)
                for match in matches:
                    cleaned = self.clean_and_validate_name(match)
                    if cleaned:
                        return cleaned
                        
        except:
            pass
            
        return None

    def extract_from_manual_patterns(self, pick_element, pick_number):
        """Manual extraction for specific known patterns"""
        try:
            # Get all text and split into lines
            full_text = pick_element.get_attribute('textContent')
            if not full_text:
                return None
            
            lines = [line.strip() for line in full_text.split('\n') if line.strip()]
            
            # Look for likely player names in the lines
            for line in lines:
                # Skip obvious non-player lines
                if any(skip in line.lower() for skip in ['pick', 'round', 'draft', 'team', 'position']):
                    continue
                
                # Look for names that follow standard patterns
                if len(line.split()) == 2:  # First Last
                    cleaned = self.clean_and_validate_name(line)
                    if cleaned:
                        return cleaned
                elif len(line.split()) == 3:  # First Middle Last
                    cleaned = self.clean_and_validate_name(line)
                    if cleaned:
                        return cleaned
                        
        except:
            pass
            
        return None

    def clean_and_validate_name(self, text):
        """Clean text and validate if it's a player name"""
        try:
            if not text:
                return None
            
            # Basic cleaning
            text = text.strip()
            text = re.sub(r'^(Pick|#\d+|Round \d+)', '', text, flags=re.IGNORECASE).strip()
            text = re.sub(r'(QB|RB|WR|TE|OL|DL|LB|CB|S|K|P)$', '', text, flags=re.IGNORECASE).strip()
            text = re.sub(r'[,\-\|].*$', '', text).strip()  # Remove everything after comma, dash, or pipe
            
            # Team name filtering
            team_keywords = [
                'titans', 'browns', 'giants', 'patriots', 'jaguars', 'raiders', 'jets', 'panthers',
                'saints', 'bears', 'francisco', 'cowboys', 'dolphins', 'colts', 'falcons',
                'cardinals', 'bengals', 'seahawks', 'buccaneers', 'broncos', 'packers', 'chargers',
                'chiefs', 'bills', 'steelers', 'ravens', 'lions', 'vikings', 'rams', 'eagles',
                'commanders', 'texans', 'titan', 'brown', 'giant', 'patriot', 'jaguar', 'raider',
                'jet', 'panther', 'saint', 'bear', 'cowboy', 'dolphin', 'colt', 'falcon',
                'cardinal', 'bengal', 'seahawk', 'buccaneer', 'bronco', 'packer', 'charger'
            ]
            
            if any(team in text.lower() for team in team_keywords):
                return None
            
            # Skip common non-name words
            skip_words = ['pick', 'round', 'draft', 'team', 'position', 'college', 'university', 'select']
            if any(skip in text.lower() for skip in skip_words):
                return None
            
            # Name validation
            words = text.split()
            if len(words) < 2 or len(words) > 3:
                return None
            
            # Check if all words look like names
            for word in words:
                clean_word = word.replace('.', '').replace("'", '')
                if not clean_word.isalpha():
                    return None
                if not (word[0].isupper() and word[1:].islower()):
                    return None
            
            return text
            
        except:
            return None

    def remove_overlays(self):
        """Remove cookie banners and overlays"""
        try:
            overlay_selectors = [
                '[data-module="CookieBanner"]',
                '.onetrust-banner-sdk',
                '.cookie-banner',
                '.overlay',
                '.modal'
            ]
            
            for selector in overlay_selectors:
                try:
                    overlays = self.driver.find_elements(By.CSS_SELECTOR, selector)
                    for overlay in overlays:
                        if overlay.is_displayed():
                            self.driver.execute_script("arguments[0].style.display = 'none';", overlay)
                except:
                    continue
        except:
            pass

    def analyze_all_authors(self):
        """Analyze all authors and extract player data"""
        print("🔍 Analyzing all authors for player selections...")
        
        for author, url in self.author_urls.items():
            players = self.extract_players_from_author(url, author)
            self.all_players.extend(players)
        
        print(f"\n✓ Analysis complete! Found {len(self.all_players)} total player selections")

    def create_player_ranking_document(self):
        """Create a Word document with player rankings"""
        print("📊 Creating player ranking document...")
        
        # Count player frequencies
        player_counts = Counter(selection['name'] for selection in self.all_players)
        
        # Sort by frequency (most picked first)
        ranked_players = player_counts.most_common()
        
        # Create Word document
        doc = Document()
        
        # Set margins
        sections = doc.sections
        for section in sections:
            section.top_margin = Inches(0.5)
            section.bottom_margin = Inches(0.5)
            section.left_margin = Inches(0.5)
            section.right_margin = Inches(0.5)
        
        # Title
        title = doc.add_heading('NFL 2025 Mock Draft Player Rankings', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_run = title.runs[0]
        title_run.font.size = Pt(24)
        title_run.font.color.rgb = RGBColor(0, 53, 148)
        
        # Subtitle
        subtitle = doc.add_paragraph(f'Players Ranked by Selection Frequency Across 7 NFL.com Experts - {datetime.now().strftime("%B %d, %Y")}')
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        subtitle_run = subtitle.runs[0]
        subtitle_run.font.size = Pt(12)
        subtitle_run.font.color.rgb = RGBColor(107, 114, 128)
        
        # Summary stats
        summary = doc.add_paragraph()
        summary_run = summary.add_run(f"📊 Analysis Summary: {len(ranked_players)} unique players • {len(self.all_players)} total selections • 7 expert analysts")
        summary_run.font.size = Pt(11)
        summary_run.font.color.rgb = RGBColor(107, 114, 128)
        summary.alignment = WD_ALIGN_PARAGRAPH.CENTER
        summary.space_after = Pt(12)
        
        # Player rankings
        ranking_header = doc.add_heading('🏆 Player Rankings by Selection Frequency', level=1)
        ranking_header_run = ranking_header.runs[0]
        ranking_header_run.font.color.rgb = RGBColor(0, 53, 148)
        
        # Add each player with details
        for rank, (player_name, count) in enumerate(ranked_players, 1):
            # Player entry
            player_para = doc.add_paragraph()
            
            # Rank and count
            rank_run = player_para.add_run(f"{rank:2d}. ")
            rank_run.font.size = Pt(14)
            rank_run.font.bold = True
            rank_run.font.color.rgb = RGBColor(0, 53, 148)
            
            # Player name
            name_run = player_para.add_run(f"{player_name}")
            name_run.font.size = Pt(14)
            name_run.font.bold = True
            
            # Selection count
            count_run = player_para.add_run(f" ({count} selections)")
            count_run.font.size = Pt(12)
            count_run.font.color.rgb = RGBColor(220, 38, 127)  # Pink accent
            
            # Show which authors picked this player
            if player_name in self.player_selections:
                authors_info = []
                for author, pick_num in self.player_selections[player_name].items():
                    authors_info.append(f"{author} (#{pick_num})")
                
                if authors_info:
                    detail_para = doc.add_paragraph()
                    detail_run = detail_para.add_run(f"    Selected by: {', '.join(authors_info)}")
                    detail_run.font.size = Pt(10)
                    detail_run.font.color.rgb = RGBColor(107, 114, 128)
                    detail_para.space_after = Pt(4)
            
            player_para.space_after = Pt(8)
        
        # Save document
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        output_path = f'processed/NFL_PLAYER_RANKINGS_ENHANCED_{timestamp}.docx'
        doc.save(output_path)
        
        print(f"✓ Enhanced player ranking document saved: {output_path}")
        return output_path, ranked_players

    def cleanup(self):
        """Clean up resources"""
        if self.driver:
            self.driver.quit()

def main():
    print("=== NFL Player Ranking Analyzer Enhanced ===")
    print("🔍 Analyzing player selections across all 7 NFL.com experts")
    print("📊 Enhanced player name extraction and ranking")
    print("=" * 60)
    
    analyzer = NFLPlayerRankingAnalyzerEnhanced()
    
    if not analyzer.driver:
        print("❌ Cannot proceed without WebDriver")
        return
    
    try:
        # Analyze all authors
        analyzer.analyze_all_authors()
        
        # Create ranking document
        output_path, ranked_players = analyzer.create_player_ranking_document()
        
        print(f"\n🎉 SUCCESS! Enhanced player ranking analysis complete!")
        print("=" * 60)
        print(f"📁 Document: {output_path}")
        print(f"📊 Top 15 Most Selected Players:")
        
        for rank, (player_name, count) in enumerate(ranked_players[:15], 1):
            print(f"   {rank:2d}. {player_name} ({count} selections)")
        
        print(f"\n📈 Analysis Summary:")
        print(f"   • {len(ranked_players)} unique players identified")
        print(f"   • {len(analyzer.all_players)} total selections analyzed")
        print(f"   • 7 NFL.com expert mock drafts processed")
        print(f"   • Enhanced name extraction with pattern matching")
        
    finally:
        analyzer.cleanup()

if __name__ == "__main__":
    main() 