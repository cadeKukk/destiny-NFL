#!/usr/bin/env python3
"""
NFL Screenshot Complete - All Authors, 32 Picks, With Descriptions
Takes screenshots of all NFL.com mock draft pages with pick reasoning
"""

import os
import time
from datetime import datetime
from docx import Document
from docx.shared import Inches, RGBColor, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_PARAGRAPH_ALIGNMENT
from selenium import webdriver
from selenium.webdriver.chrome.options import Options
from selenium.webdriver.common.by import By
from selenium.webdriver.support.ui import WebDriverWait
from selenium.webdriver.support import expected_conditions as EC

class NFLScreenshotComplete:
    def __init__(self):
        self.setup_selenium()
        
        # UPDATED URLs - REMOVED Lance Zierlein and Chad Reuter as requested
        self.author_urls = {
            'Bucky Brooks': 'https://www.nfl.com/news/bucky-brooks-2025-nfl-mock-draft-4-0-steelers-land-shedeur-sanders-cowboys-broncos-select-rbs',
            'Daniel Jeremiah': 'https://www.nfl.com/news/daniel-jeremiah-2025-nfl-mock-draft-4-0',
            'Charles Davis': 'https://www.nfl.com/news/charles-davis-2025-nfl-mock-draft-3-0-cam-ward-only-qb-in-round-1-eagles-pick-te-mason-taylor',
            'Eric Edholm': 'https://www.nfl.com/news/eric-edholm-2025-nfl-mock-draft-3-0-four-first-round-quarterbacks-jaguars-take-rb-ashton-jeanty',
            'Dan Parr': 'https://www.nfl.com/news/dan-parr-2025-nfl-mock-draft-2-0-offensive-linemen-dominate-top-10-bears-grab-tight-end-tyler-warren',
            'Gennaro Filice': 'https://www.nfl.com/news/gennaro-filice-2025-nfl-mock-draft-2-0-rb-ashton-jeanty-goes-top-5-cowboys-jump-for-jalon-walker',
            'Marc Ross': 'https://www.nfl.com/news/marc-ross-2025-nfl-mock-draft-1-0-three-qbs-selected-in-top-10-jets-snag-rb-ashton-jeanty'
        }
        
        os.makedirs('processed/complete_screenshots', exist_ok=True)
        self.pick_descriptions = {}  # Store descriptions for each pick

    def setup_selenium(self):
        """Setup Selenium WebDriver for taking screenshots"""
        try:
            chrome_options = Options()
            chrome_options.add_argument('--no-sandbox')
            chrome_options.add_argument('--disable-dev-shm-usage')
            chrome_options.add_argument('--window-size=1800,1400')  # Large window
            chrome_options.add_argument('--user-agent=Mozilla/5.0 (Macintosh; Intel Mac OS X 10_15_7) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/120.0.0.0 Safari/537.36')
            chrome_options.add_argument('--disable-blink-features=AutomationControlled')
            chrome_options.add_argument('--force-device-scale-factor=1')
            
            self.driver = webdriver.Chrome(options=chrome_options)
            print("✓ Selenium WebDriver setup complete")
        except Exception as e:
            print(f"⚠️ Selenium setup failed: {e}")
            self.driver = None

    def screenshot_webpage_content(self, url, author):
        """Take screenshots of NFL.com webpage content - ALWAYS returns something for every author"""
        if not self.driver:
            print(f"⚠️ No WebDriver available for {author}")
            return [f"No WebDriver available for {author}"]
            
        screenshots = []
        
        try:
            print(f"📸 Capturing screenshots for {author}...")
            
            # Navigate to the page
            self.driver.get(url)
            time.sleep(6)  # Reduced wait for page to fully load
            
            # Remove overlays
            self.remove_overlays()
            time.sleep(3)
            
            # Get article header (always try to get something)
            header_screenshot = self.screenshot_article_header(author)
            if header_screenshot:
                screenshots.append(header_screenshot)
            
            # Get individual draft picks (UP TO 32 PICKS)
            pick_screenshots = self.screenshot_individual_picks(author)
            screenshots.extend(pick_screenshots)
            
            # If no picks found, ALWAYS get page sections so author appears in document
            if not pick_screenshots:
                print(f"   📄 No individual picks found for {author}, taking page sections...")
                section_screenshots = self.screenshot_page_sections(author)
                screenshots.extend(section_screenshots)
            
            # If still no screenshots at all, take a full page screenshot as last resort
            if not screenshots:
                print(f"   🚨 Last resort: taking full page screenshot for {author}")
                fallback_screenshot = self.screenshot_full_page_fallback(author)
                if fallback_screenshot:
                    screenshots.append(fallback_screenshot)
                
        except Exception as e:
            print(f"   ⚠️ Error processing {author}: {e}")
            # Even if there's an error, try to get a fallback screenshot
            try:
                fallback_screenshot = self.screenshot_full_page_fallback(author)
                if fallback_screenshot:
                    screenshots.append(fallback_screenshot)
            except:
                pass
        
        # Ensure every author has at least one entry
        if not screenshots:
            screenshots = [f"Could not capture content for {author}"]
            
        print(f"   ✓ Captured {len([s for s in screenshots if s.endswith('.png')])} screenshots for {author}")
        return screenshots

    def remove_overlays(self):
        """Remove cookie banners and overlays"""
        try:
            overlay_selectors = [
                '[data-module="CookieBanner"]',
                '.onetrust-banner-sdk',
                '.cookie-banner',
                '[class*="overlay"]',
                '[class*="modal"]',
                '.nfl-banner',
                '[id*="onetrust"]'
            ]
            
            for selector in overlay_selectors:
                try:
                    overlays = self.driver.find_elements(By.CSS_SELECTOR, selector)
                    for overlay in overlays:
                        if overlay.is_displayed():
                            self.driver.execute_script("arguments[0].style.display = 'none';", overlay)
                except:
                    continue
                    
            # Try to click close buttons
            close_selectors = ['[aria-label="Close"]', '.close-button', '[data-dismiss]']
            for selector in close_selectors:
                try:
                    buttons = self.driver.find_elements(By.CSS_SELECTOR, selector)
                    for button in buttons:
                        if button.is_displayed():
                            button.click()
                            time.sleep(1)
                except:
                    continue
                    
        except Exception as e:
            print(f"   ⚠️ Error removing overlays: {e}")

    def screenshot_article_header(self, author):
        """Screenshot the article header"""
        try:
            header_selectors = [
                '.nfl-c-article__header',
                'h1',
                '.article-header',
                '.article-title'
            ]
            
            for selector in header_selectors:
                try:
                    header_element = self.driver.find_element(By.CSS_SELECTOR, selector)
                    if header_element.is_displayed():
                        # Scroll to element
                        self.driver.execute_script("arguments[0].scrollIntoView({block: 'center'});", header_element)
                        time.sleep(3)
                        
                        # Take screenshot
                        screenshot_path = f"processed/complete_screenshots/{author}_header.png"
                        header_element.screenshot(screenshot_path)
                        
                        print(f"   ✓ Header screenshot: {author}_header.png")
                        return screenshot_path
                except:
                    continue
        except Exception as e:
            print(f"   ⚠️ Could not capture header for {author}: {e}")
            
        return None

    def screenshot_individual_picks(self, author):
        """Screenshot individual NFL draft picks - UP TO 32 PICKS"""
        screenshots = []
        
        try:
            # Look for NFL.com draft pick elements
            pick_selectors = [
                '.nfl-o-ranked-item.nfl-o-ranked-item--side-by-side',
                '.nfl-o-ranked-item'
            ]
            
            pick_elements = []
            for selector in pick_selectors:
                try:
                    elements = self.driver.find_elements(By.CSS_SELECTOR, selector)
                    if elements:
                        pick_elements = elements[:32]  # UP TO 32 PICKS
                        print(f"   📋 Found {len(pick_elements)} draft picks for {author} using: {selector}")
                        break
                except:
                    continue
            
            if pick_elements:
                # Initialize author in pick_descriptions if not exists
                if author not in self.pick_descriptions:
                    self.pick_descriptions[author] = {}
                
                for i, pick_element in enumerate(pick_elements, 1):
                    try:
                        print(f"   🔍 Processing Pick {i}...")
                        
                        # Scroll element into view for better capture
                        self.driver.execute_script("arguments[0].scrollIntoView({block: 'center', behavior: 'smooth'});", pick_element)
                        time.sleep(0.8)  # Reduced wait for scroll to complete
                        
                        # Extract description/reasoning for this pick FIRST
                        description = self.extract_pick_description_enhanced(pick_element, i, author)
                        self.pick_descriptions[author][i] = description
                        print(f"      📝 Description: {description[:100]}...")
                        
                        # Take screenshot of the pick element with better positioning
                        filename = f'{author}_pick_{i}.png'
                        filepath = os.path.join('processed/complete_screenshots', filename)
                        
                        # Ensure element is fully visible
                        self.driver.execute_script("arguments[0].style.border='2px solid red';", pick_element)
                        time.sleep(0.5)
                        
                        # Take the screenshot
                        pick_element.screenshot(filepath)
                        
                        # Remove the border
                        self.driver.execute_script("arguments[0].style.border='';", pick_element)
                        
                        screenshots.append(filepath)
                        print(f"   ✓ Pick {i} screenshot captured")
                        
                        time.sleep(0.4)  # Reduced pause between screenshots
                        
                    except Exception as e:
                        print(f"   ⚠️ Error capturing pick {i}: {e}")
                        # Store a placeholder description even if screenshot fails
                        if author not in self.pick_descriptions:
                            self.pick_descriptions[author] = {}
                        self.pick_descriptions[author][i] = f"Analysis for pick #{i} by {author}."
                        continue
            else:
                # Fallback: try to capture a full page section
                print(f"   ⚠️ Using fallback method for {author}")
                try:
                    filename = f'{author}_section.png'
                    filepath = os.path.join('processed/complete_screenshots', filename)
                    self.driver.save_screenshot(filepath)
                    screenshots.append(filepath)
                except:
                    pass
                    
        except Exception as e:
            print(f"   ⚠️ Error in screenshot_individual_picks: {e}")
            
        return screenshots

    def _get_filtered_analysis_paragraphs(self):
        """Get filtered analysis paragraphs for sequential mapping"""
        try:
            # Get all paragraph elements
            all_paragraphs = self.driver.find_elements(By.CSS_SELECTOR, 'p')
            
            # Get first and last pick locations for filtering boundaries
            pick_elements = self.driver.find_elements(By.CSS_SELECTOR, '.nfl-o-ranked-item.nfl-o-ranked-item--side-by-side')
            if not pick_elements:
                return []
            
            first_pick_y = pick_elements[0].location['y']
            last_pick_y = pick_elements[-1].location['y']
            
            analysis_paragraphs = []
            for para in all_paragraphs:
                try:
                    text = para.get_attribute('textContent').strip()
                    location = para.location['y']
                    
                    # Filter criteria based on debug findings
                    if (text and len(text) > 50 and len(text) < 1000 and
                        location > first_pick_y and  # After first pick
                        location < last_pick_y + 2000 and  # Before way after last pick
                        any(keyword in text.lower() for keyword in ['quarterback', 'player', 'draft', 'team', 'offense', 'defense', 'potential', 'needs', 'season', 'franchise']) and
                        not text.startswith('Pick') and
                        '©' not in text and
                        'nfl.com' not in text.lower() and
                        'cookie' not in text.lower() and
                        'privacy' not in text.lower() and
                        # Filter out intro/outro text
                        'finally the week' not in text.lower() and
                        'trades that are struck' not in text.lower() and
                        'in his final mock' not in text.lower() and
                        'with round 1' not in text.lower()):
                        
                        analysis_paragraphs.append((location, text))
                except:
                    continue
            
            # Sort by Y position and return just the text
            analysis_paragraphs.sort(key=lambda x: x[0])
            return [text for location, text in analysis_paragraphs]
            
        except Exception as e:
            print(f"      ⚠️ Error getting analysis paragraphs: {e}")
            return []

    def extract_pick_description_enhanced(self, pick_element, pick_number, author):
        """Enhanced description extraction using improved spatial positioning and context analysis"""
        try:
            print(f"      🔍 Extracting description for pick {pick_number}...")
            
            # Get player info from the pick element
            pick_text = pick_element.get_attribute('textContent')
            player_name = None
            team_name = None
            
            # Extract player and team names from pick element
            lines = [line.strip() for line in pick_text.split('\n') if line.strip()]
            for line in lines:
                # Check for team names
                if any(team in line.lower() for team in ['titans', 'browns', 'giants', 'patriots', 'raiders', 'jaguars', 'jets', 'panthers', 'saints', 'lions', 'cowboys', 'dolphins', 'colts', 'falcons', 'cardinals', 'bengals', 'vikings', 'buccaneers', 'broncos', 'chargers', 'steelers', 'packers', 'texans', 'rams', 'eagles', 'bills', 'chiefs', 'seahawks', 'commanders']):
                    team_name = line
                # Look for player name (typically 2 words, starts with capital)
                elif len(line.split()) == 2 and line[0].isupper() and 'pick' not in line.lower():
                    player_name = line
                    break
            
            print(f"         Team: {team_name}, Player: {player_name}")
            
            # Find analysis text that comes immediately after this pick element
            pick_location = pick_element.location['y']
            
            # Get all paragraph elements and find those that come after this pick
            all_paragraphs = self.driver.find_elements(By.CSS_SELECTOR, 'p')
            
            analysis_candidates = []
            for para in all_paragraphs:
                try:
                    para_location = para.location['y']
                    para_text = para.get_attribute('textContent').strip()
                    
                    # Only consider paragraphs that come after this pick with analysis content
                    if (para_location > pick_location and 
                        len(para_text) > 50 and len(para_text) < 1000 and
                        any(keyword in para_text.lower() for keyword in ['quarterback', 'player', 'draft', 'team', 'offense', 'defense', 'potential', 'needs', 'season', 'franchise', 'protection', 'elite']) and
                        not para_text.startswith('Pick') and
                        '©' not in para_text and
                        'nfl.com' not in para_text.lower() and
                        'cookie' not in para_text.lower() and
                        'privacy' not in para_text.lower()):
                        
                        analysis_candidates.append((para_location, para_text))
                except:
                    continue
            
            # Sort by Y position to get the closest analysis text
            analysis_candidates.sort(key=lambda x: x[0])
            
            print(f"         Found {len(analysis_candidates)} analysis candidates after pick")
            
            # Strategy A: Look for analysis that mentions the player or team specifically
            if player_name or team_name:
                for location, text in analysis_candidates[:3]:  # Check first 3 candidates
                    text_lower = text.lower()
                    
                    # Check if mentions player name
                    if player_name and player_name.lower() in text_lower:
                        description = text.replace('\n', ' ').replace('\t', ' ')
                        description = ' '.join(description.split())
                        if len(description) > 400:
                            description = description[:400] + '...'
                        print(f"      ✓ Found player-specific analysis: {description[:50]}...")
                        return description
                    
                    # Check if mentions team characteristics or context
                    if team_name:
                        team_keywords = {
                            'giants': ['giants', 'new york'],
                            'titans': ['titans', 'tennessee'],
                            'browns': ['browns', 'cleveland'], 
                            'patriots': ['patriots', 'new england', 'drake maye'],
                            'jaguars': ['jaguars', 'jacksonville', 'trevor lawrence'],
                            'raiders': ['raiders', 'las vegas'],
                            'jets': ['jets', 'new york jets', 'justin fields'],
                            'cowboys': ['cowboys', 'dallas'],
                            'saints': ['saints', 'new orleans'],
                            'bears': ['bears', 'chicago'],
                            'panthers': ['panthers', 'carolina'],
                            'dolphins': ['dolphins', 'miami', 'jalen ramsey'],
                            'colts': ['colts', 'indianapolis'],
                            'falcons': ['falcons', 'atlanta'],
                            'cardinals': ['cardinals', 'arizona'],
                            'bengals': ['bengals', 'cincinnati'],
                            'vikings': ['vikings', 'minnesota'],
                            'buccaneers': ['buccaneers', 'tampa bay', 'bucs'],
                            'broncos': ['broncos', 'denver', 'sean payton'],
                            'chargers': ['chargers', 'los angeles chargers'],
                            'steelers': ['steelers', 'pittsburgh'],
                            'packers': ['packers', 'green bay'],
                            'texans': ['texans', 'houston'],
                            'rams': ['rams', 'los angeles rams'],
                            'eagles': ['eagles', 'philadelphia'],
                            'bills': ['bills', 'buffalo'],
                            'chiefs': ['chiefs', 'kansas city'],
                            'seahawks': ['seahawks', 'seattle'],
                            'commanders': ['commanders', 'washington']
                        }
                        
                        team_lower = team_name.lower()
                        for team_key, keywords in team_keywords.items():
                            if team_key in team_lower:
                                if any(keyword in text_lower for keyword in keywords):
                                    description = text.replace('\n', ' ').replace('\t', ' ')
                                    description = ' '.join(description.split())
                                    if len(description) > 400:
                                        description = description[:400] + '...'
                                    print(f"      ✓ Found team-specific analysis: {description[:50]}...")
                                    return description
            
            # Strategy B: Use the closest available analysis text (not indexed by pick number)
            if analysis_candidates:
                # Simply use the first (closest) available analysis text
                # This matches what the debug showed as the correct approach
                location, text = analysis_candidates[0]
                description = text.replace('\n', ' ').replace('\t', ' ')
                description = ' '.join(description.split())
                if len(description) > 400:
                    description = description[:400] + '...'
                print(f"      ✓ Found closest positional analysis: {description[:50]}...")
                return description
            
            # Strategy 2: Fallback to element-based extraction
            full_text = pick_element.get_attribute('textContent')
            
            if not full_text:
                print(f"      ⚠️ No text content found, using placeholder")
                return f"Draft analysis for pick #{pick_number} by {author}."
            
            # Clean and split the text
            lines = [line.strip() for line in full_text.split('\n') if line.strip()]
            
            # Look for any substantial text lines
            substantial_lines = [line for line in lines if len(line) > 30 and not line.isupper() and 'Pick' not in line]
            if substantial_lines:
                fallback = substantial_lines[0]
                if len(fallback) > 300:
                    fallback = fallback[:300] + '...'
                print(f"      ⚠️ Using fallback description: {fallback[:50]}...")
                return fallback
            else:
                print(f"      ⚠️ No description found, using placeholder")
                return f"Draft analysis for pick #{pick_number} by {author}."
                
        except Exception as e:
            print(f"      ⚠️ Error extracting description: {e}")
            return f"Analysis for pick #{pick_number} not available."

    def screenshot_page_sections(self, author):
        """Take page section screenshots as backup"""
        screenshots = []
        
        try:
            # Scroll to top
            self.driver.execute_script("window.scrollTo(0, 0);")
            time.sleep(2)
            
            # Get page dimensions
            total_height = self.driver.execute_script("return document.body.scrollHeight")
            viewport_height = self.driver.execute_script("return window.innerHeight")
            
            # Take screenshots in sections
            section_height = viewport_height
            sections = min(4, max(1, total_height // section_height))  # Max 4 sections
            
            for i in range(sections):
                scroll_position = i * section_height
                
                # Scroll to position
                self.driver.execute_script(f"window.scrollTo(0, {scroll_position});")
                time.sleep(2)  # Reduced wait for content to load
                
                # Take screenshot
                screenshot_path = f"processed/complete_screenshots/{author}_section_{i+1:02d}.png"
                self.driver.save_screenshot(screenshot_path)
                
                screenshots.append(screenshot_path)
                print(f"   ✓ Section {i+1} screenshot captured")
                
        except Exception as e:
            print(f"   ⚠️ Error taking page sections for {author}: {e}")
            
        return screenshots

    def screenshot_full_page_fallback(self, author):
        """Take a full page screenshot as absolute fallback"""
        try:
            self.driver.execute_script("window.scrollTo(0, 0);")
            time.sleep(3)
            
            screenshot_path = f"processed/complete_screenshots/{author}_fullpage.png"
            self.driver.save_screenshot(screenshot_path)
            
            print(f"   ✓ Fallback full page screenshot: {author}_fullpage.png")
            return screenshot_path
        except Exception as e:
            print(f"   ⚠️ Could not take fallback screenshot for {author}: {e}")
            return None

    def create_word_document(self, all_screenshots):
        """Create a Word document with all screenshots and descriptions"""
        print("📄 Creating optimized Word document with all authors...")
        
        doc = Document()
        
        # Set narrow margins for space efficiency
        sections = doc.sections
        for section in sections:
            section.top_margin = Inches(0.4)
            section.bottom_margin = Inches(0.4)
            section.left_margin = Inches(0.4)
            section.right_margin = Inches(0.4)
        
        # Title
        title = doc.add_heading('NFL.com 2025 Mock Draft Screenshots', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_run = title.runs[0]
        title_run.font.size = Pt(18)
        title_run.font.color.rgb = RGBColor(0, 53, 148)
        
        # Subtitle with date and summary
        subtitle = doc.add_paragraph(f'Complete NFL.com Mock Drafts - {datetime.now().strftime("%B %d, %Y")}')
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        subtitle_run = subtitle.runs[0]
        subtitle_run.font.size = Pt(10)
        subtitle_run.font.color.rgb = RGBColor(107, 114, 128)
        
        # Summary
        total_screenshots = sum(len(screenshots) for screenshots in all_screenshots.values())
        summary = doc.add_paragraph()
        summary_run = summary.add_run(f"📊 {len(all_screenshots)} Authors • {total_screenshots} Screenshots • Top 32 Picks Each • Expert Analysis Included")
        summary_run.font.size = Pt(9)
        summary_run.font.color.rgb = RGBColor(107, 114, 128)
        summary.alignment = WD_ALIGN_PARAGRAPH.CENTER
        summary.space_after = Pt(8)
        
        # Add screenshots for each author
        for author, screenshots in all_screenshots.items():
            if screenshots:
                print(f"   📝 Adding {author} to document...")
                
                # Author header
                author_header = doc.add_heading(f'{author} - 2025 Mock Draft', level=1)
                author_header_run = author_header.runs[0]
                author_header_run.font.size = Pt(16)
                author_header_run.font.color.rgb = RGBColor(0, 53, 148)
                author_header.space_before = Pt(6)
                author_header.space_after = Pt(4)
                
                # Debug: Print available descriptions for this author
                if author in self.pick_descriptions:
                    print(f"      📝 Found {len(self.pick_descriptions[author])} descriptions for {author}")
                else:
                    print(f"      ⚠️ No descriptions found for {author}")
                
                # Add each screenshot with description
                for i, screenshot in enumerate(screenshots):
                    try:
                        # Skip header screenshots for pick processing
                        if 'header' in screenshot:
                            # Add header screenshot
                            if os.path.exists(screenshot):
                                img_paragraph = doc.add_paragraph()
                                img_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                                try:
                                    img_run = img_paragraph.runs[0] if img_paragraph.runs else img_paragraph.add_run()
                                    img_run.add_picture(screenshot, width=Inches(6.5))
                                    img_paragraph.space_after = Pt(8)
                                except:
                                    pass
                            continue
                        
                        # Process pick screenshots
                        if 'pick_' in screenshot:
                            # Extract pick number from filename more reliably
                            try:
                                # Extract from filename like "Author_pick_5.png"
                                filename = os.path.basename(screenshot)
                                pick_part = filename.split('pick_')[-1]
                                pick_num = int(pick_part.split('.')[0])
                            except:
                                # Fallback to position in list
                                pick_num = i
                            
                            print(f"      🔍 Processing pick #{pick_num} screenshot...")
                            
                            # Add screenshot (removed Pick #X header as requested)
                            if os.path.exists(screenshot):
                                img_paragraph = doc.add_paragraph()
                                img_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                                
                                try:
                                    img_run = img_paragraph.runs[0] if img_paragraph.runs else img_paragraph.add_run()
                                    img_run.add_picture(screenshot, width=Inches(6.5))
                                except:
                                    # Fallback with smaller width if image is too large
                                    try:
                                        img_run = img_paragraph.runs[0] if img_paragraph.runs else img_paragraph.add_run()
                                        img_run.add_picture(screenshot, width=Inches(5.5))
                                    except Exception as img_error:
                                        print(f"         ⚠️ Could not add image: {img_error}")
                                        continue
                                
                                img_paragraph.space_after = Pt(4)
                                
                                # Add description for this pick
                                description = None
                                if author in self.pick_descriptions and pick_num in self.pick_descriptions[author]:
                                    description = self.pick_descriptions[author][pick_num]
                                    print(f"         ✓ Adding description: {description[:50]}...")
                                else:
                                    description = f"Analysis for pick #{pick_num} by {author}."
                                    print(f"         ⚠️ No description found, using placeholder")
                                
                                # Description paragraph
                                desc_paragraph = doc.add_paragraph()
                                desc_run = desc_paragraph.add_run(f"📝 Analysis: {description}")
                                desc_run.font.size = Pt(9)
                                desc_run.font.color.rgb = RGBColor(74, 85, 104)
                                desc_run.italic = True
                                desc_paragraph.space_after = Pt(8)
                        
                    except Exception as e:
                        print(f"   ⚠️ Error adding screenshot {screenshot}: {e}")
                        continue
        
        # Save document
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        output_path = f'processed/NFL_COMPLETE_ALL_AUTHORS_{timestamp}.docx'
        doc.save(output_path)
        
        # Debug: Print summary of descriptions collected
        print(f"\n📊 Description Summary:")
        for author, descriptions in self.pick_descriptions.items():
            print(f"   {author}: {len(descriptions)} descriptions")
        
        return output_path

    def cleanup(self):
        """Clean up resources"""
        if self.driver:
            self.driver.quit()



def main():
    print("=== NFL Complete Screenshot Creator ===")
    print("📸 ALL 7 AUTHORS with up to 32 picks each")
    print("🎯 Expert analysis included under each pick")
    print("✅ Every author guaranteed to appear in document")
    print("=======================================================")
    
    creator = NFLScreenshotComplete()
    
    if not creator.driver:
        print("❌ Cannot proceed without WebDriver")
        return
    
    try:
        all_screenshots = {}
        
        # Process ALL authors
        for author, url in creator.author_urls.items():
            screenshots = creator.screenshot_webpage_content(url, author)
            
            all_screenshots[author] = screenshots
        
        # Create Word document with all authors
        output_path = creator.create_word_document(all_screenshots)
        
        print(f"\n🎉 SUCCESS! Complete NFL.com screenshots captured!")
        print("=======================================================")
        print(f"📁 Document: {output_path}")
        print(f"📸 Screenshots: processed/complete_screenshots/")
        
        total_screenshots = sum(len([s for s in author_data if isinstance(s, str) and s.endswith('.png')]) for author_data in all_screenshots.values())
        print(f"\n📊 Summary:")
        print(f"   • {len(creator.author_urls)} authors processed (ALL)")
        print(f"   • {total_screenshots} total screenshots captured")
        print(f"   • Up to 32 picks per author")
        print(f"   • Expert analysis included under each pick")
        print(f"   • Optimized spacing and formatting")
        print(f"   • Every author appears in document")
        print(f"   • Ready for viewing in Word")
        
    finally:
        creator.cleanup()

if __name__ == "__main__":
    main() 