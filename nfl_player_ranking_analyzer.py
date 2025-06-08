#!/usr/bin/env python3
"""
NFL Player Ranking Analyzer
Analyzes NFL.com mock draft pages to create a ranked list of players by frequency
"""

import os
import time
from datetime import datetime
from collections import Counter
from docx import Document
from docx.shared import Inches, RGBColor, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from selenium import webdriver
from selenium.webdriver.chrome.options import Options
from selenium.webdriver.common.by import By
from selenium.webdriver.support.ui import WebDriverWait
from selenium.webdriver.support import expected_conditions as EC
import re

class NFLPlayerRankingAnalyzer:
    def __init__(self):
        self.setup_selenium()
        
        # UPDATED URLs as provided by user
        self.author_urls = {
            'Bucky Brooks': 'https://www.nfl.com/news/bucky-brooks-2025-nfl-mock-draft-4-0-steelers-land-shedeur-sanders-cowboys-broncos-select-rbs',
            'Daniel Jeremiah': 'https://www.nfl.com/news/daniel-jeremiah-2025-nfl-mock-draft-4-0',
            'Lance Zierlein': 'https://www.nfl.com/news/lance-zierlein-2025-nfl-mock-draft-4-0-colts-trade-up-for-colston-loveland-saints-go-get-jaxson-dart',
            'Charles Davis': 'https://www.nfl.com/news/charles-davis-2025-nfl-mock-draft-3-0-cam-ward-only-qb-in-round-1-eagles-pick-te-mason-taylor',
            'Eric Edholm': 'https://www.nfl.com/news/eric-edholm-2025-nfl-mock-draft-3-0-four-first-round-quarterbacks-jaguars-take-rb-ashton-jeanty',
            'Dan Parr': 'https://www.nfl.com/news/dan-parr-2025-nfl-mock-draft-2-0-offensive-linemen-dominate-top-10-bears-grab-tight-end-tyler-warren',
            'Chad Reuter': 'https://www.nfl.com/news/seven-round-2025-nfl-mock-draft-round-one',
            'Gennaro Filice': 'https://www.nfl.com/news/gennaro-filice-2025-nfl-mock-draft-2-0-rb-ashton-jeanty-goes-top-5-cowboys-jump-for-jalon-walker',
            'Marc Ross': 'https://www.nfl.com/news/marc-ross-2025-nfl-mock-draft-1-0-three-qbs-selected-in-top-10-jets-snag-rb-ashton-jeanty'
        }
        
        self.all_players = []
        self.player_selections = {}  # {player_name: {author: pick_number}}

    def setup_selenium(self):
        """Setup Selenium WebDriver"""
        try:
            chrome_options = Options()
            chrome_options.add_argument('--no-sandbox')
            chrome_options.add_argument('--disable-dev-shm-usage')
            chrome_options.add_argument('--window-size=1800,1400')
            chrome_options.add_argument('--user-agent=Mozilla/5.0 (Macintosh; Intel Mac OS X 10_15_7) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/120.0.0.0 Safari/537.36')
            chrome_options.add_argument('--disable-blink-features=AutomationControlled')
            chrome_options.add_argument('--headless')  # Run headless for faster processing
            
            self.driver = webdriver.Chrome(options=chrome_options)
            print("✓ Selenium WebDriver setup complete")
        except Exception as e:
            print(f"⚠️ Selenium setup failed: {e}")
            self.driver = None

    def extract_players_from_author(self, url, author):
        """Extract player names from an author's mock draft"""
        if not self.driver:
            return []
            
        players = []
        
        try:
            print(f"🔍 Analyzing {author}'s mock draft...")
            
            self.driver.get(url)
            time.sleep(8)
            
            # Remove overlays
            self.remove_overlays()
            time.sleep(2)
            
            # Look for NFL.com draft pick elements
            pick_selectors = [
                '.nfl-o-ranked-item.nfl-o-ranked-item--side-by-side',
                '.nfl-o-ranked-item'
            ]
            
            pick_elements = []
            for selector in pick_selectors:
                try:
                    elements = self.driver.find_elements(By.CSS_SELECTOR, selector)
                    if elements:
                        pick_elements = elements[:20]  # Top 20 picks
                        print(f"   📋 Found {len(pick_elements)} draft picks for {author}")
                        break
                except:
                    continue
            
            if pick_elements:
                for i, pick_element in enumerate(pick_elements, 1):
                    try:
                        # Extract player name from the pick element
                        player_name = self.extract_player_name_from_element(pick_element)
                        
                        if player_name:
                            players.append({
                                'name': player_name,
                                'pick': i,
                                'author': author
                            })
                            print(f"   ✓ Pick {i}: {player_name}")
                            
                            # Track in master list
                            if player_name not in self.player_selections:
                                self.player_selections[player_name] = {}
                            self.player_selections[player_name][author] = i
                        
                    except Exception as e:
                        print(f"   ⚠️ Error extracting pick {i}: {e}")
                        continue
            else:
                print(f"   ⚠️ No draft pick elements found for {author}")
                
        except Exception as e:
            print(f"   ⚠️ Error analyzing {author}: {e}")
            
        return players

    def extract_player_name_from_element(self, pick_element):
        """Extract player name from a draft pick element"""
        try:
            # Look for player name in various selectors
            name_selectors = [
                '.nfl-o-ranked-item__title',
                '[class*="title"]',
                'h3',
                'h4',
                '.player-name'
            ]
            
            for selector in name_selectors:
                try:
                    name_elements = pick_element.find_elements(By.CSS_SELECTOR, selector)
                    for name_element in name_elements:
                        text = name_element.get_attribute('textContent').strip()
                        
                        # Look for player names (typically in format "FirstName LastName")
                        # Skip team names, positions, etc.
                        if text and len(text.split()) >= 2 and len(text) < 50:
                            # Clean up the name
                            cleaned_name = self.clean_player_name(text)
                            if cleaned_name and self.is_likely_player_name(cleaned_name):
                                return cleaned_name
                except:
                    continue
            
            # If no name found in structured elements, try to extract from all text
            full_text = pick_element.get_attribute('textContent')
            if full_text:
                potential_name = self.extract_name_from_text(full_text)
                if potential_name:
                    return potential_name
                    
        except Exception as e:
            pass
            
        return None

    def clean_player_name(self, text):
        """Clean and extract player name from text"""
        try:
            # Remove common prefixes/suffixes
            text = re.sub(r'^(Pick|#\d+|Round \d+)', '', text, flags=re.IGNORECASE).strip()
            text = re.sub(r'(QB|RB|WR|TE|OL|DL|LB|CB|S|K|P)$', '', text, flags=re.IGNORECASE).strip()
            text = re.sub(r'(Quarterback|Running Back|Wide Receiver|Tight End)', '', text, flags=re.IGNORECASE).strip()
            
            # Skip team names
            team_names = [
                'Titan', 'Brown', 'Giant', 'Patriot', 'Jaguar', 'Raider', 'Jet', 'Panther',
                'Saint', 'Bear', 'Francisco', 'Cowboy', 'Dolphin', 'Colt', 'Falcon',
                'Cardinal', 'Bengal', 'Seahawk', 'Buccaneer', 'Bronco', 'Packer', 'Charger',
                'Chiefs', 'Bills', 'Steelers', 'Ravens', 'Lions', 'Vikings', 'Packers',
                'Bears', 'Saints', 'Falcons', 'Panthers', 'Buccaneers', 'Rams', 'Cardinals',
                'Seahawks', 'Niners', 'Cowboys', 'Giants', 'Eagles', 'Commanders', 'Titans',
                'Colts', 'Texans', 'Jaguars', 'Broncos', 'Chiefs', 'Raiders', 'Chargers'
            ]
            
            if any(team in text for team in team_names):
                return None
            
            # Extract name patterns (First Last or First Middle Last)
            name_pattern = r'^([A-Z][a-z]+(?:\s+[A-Z][a-z]*)*\s+[A-Z][a-z]+)'
            match = re.search(name_pattern, text)
            if match:
                return match.group(1).strip()
            
            # Fallback: just return cleaned text if it looks like a name
            if len(text.split()) in [2, 3] and all(word.replace('.', '').isalpha() for word in text.split()):
                return text
                
        except:
            pass
            
        return None

    def extract_name_from_text(self, text):
        """Extract player name from full text content"""
        try:
            # Look for common name patterns in the full text
            lines = [line.strip() for line in text.split('\n') if line.strip()]
            
            # Common player name patterns in NFL.com structure
            for line in lines:
                # Skip lines that are clearly not names
                if any(skip in line.lower() for skip in ['pick', 'round', 'draft', 'team', 'position']):
                    continue
                    
                # Look for player names - typically appear after team names
                if len(line.split()) in [2, 3] and len(line) < 40:
                    # Check if this could be a player name
                    potential_name = self.clean_player_name(line)
                    if potential_name and self.is_likely_player_name(potential_name):
                        return potential_name
            
            # Try to find names in a structured way - look for text that follows certain patterns
            # NFL.com often has structure like: "Pick X. Team Name selects Player Name, Position, School"
            pick_pattern = r'(?:pick\s+\d+[.\s]+)?(?:[a-z\s]+(?:select|draft)s?\s+)?([A-Z][a-z]+\s+[A-Z][a-z]+)'
            matches = re.findall(pick_pattern, text, re.IGNORECASE)
            
            for match in matches:
                cleaned = self.clean_player_name(match)
                if cleaned and self.is_likely_player_name(cleaned):
                    return cleaned
                    
        except:
            pass
            
        return None

    def is_likely_player_name(self, name):
        """Check if text is likely a player name"""
        if not name or len(name) < 4:
            return False
        
        # Skip obvious non-names
        skip_words = [
            'pick', 'round', 'draft', 'team', 'position', 'college', 'university',
            'quarterback', 'running', 'wide', 'tight', 'offensive', 'defensive',
            'linebacker', 'cornerback', 'safety', 'kicker', 'punter'
        ]
        
        if any(skip in name.lower() for skip in skip_words):
            return False
        
        # Skip team names (more comprehensive list)
        team_names = [
            'Titan', 'Brown', 'Giant', 'Patriot', 'Jaguar', 'Raider', 'Jet', 'Panther',
            'Saint', 'Bear', 'Francisco', 'Cowboy', 'Dolphin', 'Colt', 'Falcon',
            'Cardinal', 'Bengal', 'Seahawk', 'Buccaneer', 'Bronco', 'Packer', 'Charger',
            'Chiefs', 'Bills', 'Steelers', 'Ravens', 'Lions', 'Vikings', 'Packers',
            'Bears', 'Saints', 'Falcons', 'Panthers', 'Buccaneers', 'Rams', 'Cardinals',
            'Seahawks', 'Niners', 'Cowboys', 'Giants', 'Eagles', 'Commanders', 'Titans',
            'Colts', 'Texans', 'Jaguars', 'Broncos', 'Chiefs', 'Raiders', 'Chargers'
        ]
        
        for team in team_names:
            if team.lower() in name.lower():
                return False
        
        # Should have 2-3 words
        words = name.split()
        if len(words) < 2 or len(words) > 3:
            return False
        
        # All words should be properly capitalized and alpha
        for word in words:
            clean_word = word.replace('.', '').replace("'", '')
            if not clean_word.isalpha() or not (word[0].isupper() and word[1:].islower()):
                return False
        
        return True

    def remove_overlays(self):
        """Remove cookie banners and overlays"""
        try:
            overlay_selectors = [
                '[data-module="CookieBanner"]',
                '.onetrust-banner-sdk',
                '.cookie-banner'
            ]
            
            for selector in overlay_selectors:
                try:
                    overlays = self.driver.find_elements(By.CSS_SELECTOR, selector)
                    for overlay in overlays:
                        if overlay.is_displayed():
                            self.driver.execute_script("arguments[0].style.display = 'none';", overlay)
                except:
                    continue
        except:
            pass

    def analyze_all_authors(self):
        """Analyze all authors and extract player data"""
        print("🔍 Analyzing all authors for player selections...")
        
        for author, url in self.author_urls.items():
            players = self.extract_players_from_author(url, author)
            self.all_players.extend(players)
        
        print(f"\n✓ Analysis complete! Found {len(self.all_players)} total player selections")

    def create_player_ranking_document(self):
        """Create a Word document with player rankings"""
        print("📊 Creating player ranking document...")
        
        # Count player frequencies
        player_counts = Counter(selection['name'] for selection in self.all_players)
        
        # Sort by frequency (most picked first)
        ranked_players = player_counts.most_common()
        
        # Create Word document
        doc = Document()
        
        # Set margins
        sections = doc.sections
        for section in sections:
            section.top_margin = Inches(0.5)
            section.bottom_margin = Inches(0.5)
            section.left_margin = Inches(0.5)
            section.right_margin = Inches(0.5)
        
        # Title
        title = doc.add_heading('NFL 2025 Mock Draft Player Rankings', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_run = title.runs[0]
        title_run.font.size = Pt(24)
        title_run.font.color.rgb = RGBColor(0, 53, 148)
        
        # Subtitle
        subtitle = doc.add_paragraph(f'Players Ranked by Selection Frequency Across 9 NFL.com Experts - {datetime.now().strftime("%B %d, %Y")}')
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        subtitle_run = subtitle.runs[0]
        subtitle_run.font.size = Pt(12)
        subtitle_run.font.color.rgb = RGBColor(107, 114, 128)
        
        # Summary stats
        summary = doc.add_paragraph()
        summary_run = summary.add_run(f"📊 Analysis Summary: {len(ranked_players)} unique players • {len(self.all_players)} total selections • 9 expert analysts")
        summary_run.font.size = Pt(11)
        summary_run.font.color.rgb = RGBColor(107, 114, 128)
        summary.alignment = WD_ALIGN_PARAGRAPH.CENTER
        summary.space_after = Pt(12)
        
        # Player rankings
        ranking_header = doc.add_heading('🏆 Player Rankings by Selection Frequency', level=1)
        ranking_header_run = ranking_header.runs[0]
        ranking_header_run.font.color.rgb = RGBColor(0, 53, 148)
        
        # Add each player
        for rank, (player_name, count) in enumerate(ranked_players, 1):
            # Player entry
            player_para = doc.add_paragraph()
            
            # Rank and count
            rank_run = player_para.add_run(f"{rank:2d}. ")
            rank_run.font.size = Pt(12)
            rank_run.font.bold = True
            rank_run.font.color.rgb = RGBColor(0, 53, 148)
            
            # Player name
            name_run = player_para.add_run(f"{player_name}")
            name_run.font.size = Pt(12)
            name_run.font.bold = True
            
            # Selection count
            count_run = player_para.add_run(f" ({count} selections)")
            count_run.font.size = Pt(11)
            count_run.font.color.rgb = RGBColor(220, 38, 127)  # Pink accent
            
            # Show which authors picked this player
            if player_name in self.player_selections:
                authors_info = []
                for author, pick_num in self.player_selections[player_name].items():
                    authors_info.append(f"{author} (#{pick_num})")
                
                if authors_info:
                    detail_para = doc.add_paragraph()
                    detail_run = detail_para.add_run(f"    Selected by: {', '.join(authors_info)}")
                    detail_run.font.size = Pt(9)
                    detail_run.font.color.rgb = RGBColor(107, 114, 128)
                    detail_para.space_after = Pt(4)
            
            player_para.space_after = Pt(6)
        
        # Add methodology section
        doc.add_page_break()
        method_header = doc.add_heading('📋 Methodology', level=1)
        method_header_run = method_header.runs[0]
        method_header_run.font.color.rgb = RGBColor(0, 53, 148)
        
        methodology_text = [
            "• Analyzed mock drafts from 9 NFL.com expert analysts",
            "• Extracted player names from first 20 picks of each mock draft", 
            "• Ranked players by total number of selections across all analysts",
            "• Included pick position for each selection to show expert consensus",
            f"• Data collected on {datetime.now().strftime('%B %d, %Y')}"
        ]
        
        for item in methodology_text:
            method_para = doc.add_paragraph(item)
            method_run = method_para.runs[0]
            method_run.font.size = Pt(11)
            
        # Save document
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        output_path = f'processed/NFL_PLAYER_RANKINGS_{timestamp}.docx'
        doc.save(output_path)
        
        print(f"✓ Player ranking document saved: {output_path}")
        return output_path, ranked_players

    def cleanup(self):
        """Clean up resources"""
        if self.driver:
            self.driver.quit()

def main():
    print("=== NFL Player Ranking Analyzer ===")
    print("🔍 Analyzing player selections across all 9 NFL.com experts")
    print("📊 Creating ranked list by selection frequency")
    print("=" * 55)
    
    analyzer = NFLPlayerRankingAnalyzer()
    
    if not analyzer.driver:
        print("❌ Cannot proceed without WebDriver")
        return
    
    try:
        # Analyze all authors
        analyzer.analyze_all_authors()
        
        # Create ranking document
        output_path, ranked_players = analyzer.create_player_ranking_document()
        
        print(f"\n🎉 SUCCESS! Player ranking analysis complete!")
        print("=" * 55)
        print(f"📁 Document: {output_path}")
        print(f"📊 Top 10 Most Selected Players:")
        
        for rank, (player_name, count) in enumerate(ranked_players[:10], 1):
            print(f"   {rank:2d}. {player_name} ({count} selections)")
        
        print(f"\n📈 Analysis Summary:")
        print(f"   • {len(ranked_players)} unique players identified")
        print(f"   • {len(analyzer.all_players)} total selections analyzed")
        print(f"   • 9 NFL.com expert mock drafts processed")
        print(f"   • Rankings by selection frequency across all experts")
        
    finally:
        analyzer.cleanup()

if __name__ == "__main__":
    main() 