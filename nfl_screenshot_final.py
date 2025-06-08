#!/usr/bin/env python3
"""
NFL Screenshot Final - Simplified Version
Takes reliable screenshots of NFL.com mock draft pages using direct element capture
"""

import os
import time
from datetime import datetime
from docx import Document
from docx.shared import Inches, RGBColor, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from selenium import webdriver
from selenium.webdriver.chrome.options import Options
from selenium.webdriver.common.by import By
from selenium.webdriver.support.ui import WebDriverWait
from selenium.webdriver.support import expected_conditions as EC

class NFLScreenshotFinal:
    def __init__(self):
        self.setup_selenium()
        
        # All NFL.com mock draft URLs
        self.author_urls = {
            'Bucky Brooks': 'https://www.nfl.com/news/bucky-brooks-2025-nfl-mock-draft-4-0-steelers-land-shedeur-sanders-cowboys-broncos-select-rbs',
            'Daniel Jeremiah': 'https://www.nfl.com/news/daniel-jeremiah-2025-nfl-mock-draft-4-0-broncos-giants-trade-up-steelers-pick-shedeur-sanders',
            'Lance Zierlein': 'https://www.nfl.com/news/lance-zierlein-2025-nfl-mock-draft-4-0-colts-trade-up-for-colston-loveland-saints-go-get-jaxson-dart',
            'Charles Davis': 'https://www.nfl.com/news/charles-davis-2025-nfl-mock-draft-3-0-cam-ward-only-qb-in-round-1-eagles-pick-te-mason-taylor',
            'Eric Edholm': 'https://www.nfl.com/news/eric-edholm-2025-nfl-mock-draft-3-0-four-first-round-quarterbacks-jaguars-take-rb-ashton-jeanty'
        }
        
        os.makedirs('processed/final_screenshots', exist_ok=True)

    def setup_selenium(self):
        """Setup Selenium WebDriver for taking screenshots"""
        try:
            chrome_options = Options()
            chrome_options.add_argument('--no-sandbox')
            chrome_options.add_argument('--disable-dev-shm-usage')
            chrome_options.add_argument('--window-size=1800,1400')  # Large window
            chrome_options.add_argument('--user-agent=Mozilla/5.0 (Macintosh; Intel Mac OS X 10_15_7) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/120.0.0.0 Safari/537.36')
            chrome_options.add_argument('--disable-blink-features=AutomationControlled')
            chrome_options.add_argument('--force-device-scale-factor=1')  # Ensure consistent scaling
            
            self.driver = webdriver.Chrome(options=chrome_options)
            print("✓ Selenium WebDriver setup complete")
        except Exception as e:
            print(f"⚠️ Selenium setup failed: {e}")
            self.driver = None

    def screenshot_webpage_content(self, url, author):
        """Take screenshots of NFL.com webpage content"""
        if not self.driver:
            print(f"⚠️ No WebDriver available for {author}")
            return []
            
        screenshots = []
        
        try:
            print(f"📸 Capturing screenshots for {author}...")
            
            # Navigate to the page
            self.driver.get(url)
            time.sleep(10)  # Long wait for page to fully load
            
            # Remove overlays
            self.remove_overlays()
            time.sleep(3)
            
            # Get article header
            header_screenshot = self.screenshot_article_header(author)
            if header_screenshot:
                screenshots.append(header_screenshot)
            
            # Get all draft picks in one content area
            content_screenshot = self.screenshot_draft_content_area(author)
            if content_screenshot:
                screenshots.append(content_screenshot)
            
            # Get individual draft picks
            pick_screenshots = self.screenshot_individual_picks(author)
            screenshots.extend(pick_screenshots)
            
            # If no picks found, take page sections
            if not pick_screenshots:
                print(f"   📄 Taking page sections for {author}...")
                section_screenshots = self.screenshot_page_sections(author)
                screenshots.extend(section_screenshots)
                
        except Exception as e:
            print(f"   ⚠️ Error taking screenshots for {author}: {e}")
            
        return screenshots

    def remove_overlays(self):
        """Remove cookie banners and overlays"""
        try:
            overlay_selectors = [
                '[data-module="CookieBanner"]',
                '.onetrust-banner-sdk',
                '.cookie-banner',
                '[class*="overlay"]',
                '[class*="modal"]',
                '.nfl-banner'
            ]
            
            for selector in overlay_selectors:
                try:
                    overlays = self.driver.find_elements(By.CSS_SELECTOR, selector)
                    for overlay in overlays:
                        if overlay.is_displayed():
                            self.driver.execute_script("arguments[0].style.display = 'none';", overlay)
                except:
                    continue
        except Exception as e:
            print(f"   ⚠️ Error removing overlays: {e}")

    def screenshot_article_header(self, author):
        """Screenshot the article header"""
        try:
            header_selectors = [
                '.nfl-c-article__header',
                'h1',
                '.article-header'
            ]
            
            for selector in header_selectors:
                try:
                    header_element = self.driver.find_element(By.CSS_SELECTOR, selector)
                    if header_element.is_displayed():
                        # Scroll to element
                        self.driver.execute_script("arguments[0].scrollIntoView({block: 'center'});", header_element)
                        time.sleep(3)
                        
                        # Take screenshot
                        screenshot_path = f"processed/final_screenshots/{author}_header.png"
                        header_element.screenshot(screenshot_path)
                        
                        print(f"   ✓ Header screenshot: {author}_header.png")
                        return screenshot_path
                except:
                    continue
        except Exception as e:
            print(f"   ⚠️ Could not capture header for {author}: {e}")
            
        return None

    def screenshot_draft_content_area(self, author):
        """Screenshot the main draft content area"""
        try:
            content_selectors = [
                '.nfl-c-article__body',
                '.nfl-c-article__container',
                'main',
                'article'
            ]
            
            for selector in content_selectors:
                try:
                    content_element = self.driver.find_element(By.CSS_SELECTOR, selector)
                    if content_element.is_displayed():
                        # Scroll to element
                        self.driver.execute_script("arguments[0].scrollIntoView({block: 'start'});", content_element)
                        time.sleep(3)
                        
                        # Take screenshot
                        screenshot_path = f"processed/final_screenshots/{author}_content.png"
                        content_element.screenshot(screenshot_path)
                        
                        print(f"   ✓ Content area screenshot: {author}_content.png")
                        return screenshot_path
                except:
                    continue
        except Exception as e:
            print(f"   ⚠️ Could not capture content area for {author}: {e}")
            
        return None

    def screenshot_individual_picks(self, author):
        """Screenshot individual NFL draft picks"""
        screenshots = []
        
        try:
            # Wait for content to load
            WebDriverWait(self.driver, 15).until(
                EC.presence_of_element_located((By.CSS_SELECTOR, 'article'))
            )
            
            # Look for NFL.com draft pick elements
            pick_selectors = [
                '.nfl-o-ranked-item.nfl-o-ranked-item--side-by-side',
                '.nfl-o-ranked-item'
            ]
            
            pick_elements = []
            for selector in pick_selectors:
                try:
                    elements = self.driver.find_elements(By.CSS_SELECTOR, selector)
                    if elements:
                        pick_elements = elements[:10]  # Take first 10 picks
                        print(f"   📋 Found {len(pick_elements)} draft picks using: {selector}")
                        break
                except:
                    continue
            
            if pick_elements:
                for i, pick_element in enumerate(pick_elements, 1):
                    try:
                        # Scroll to the pick
                        self.driver.execute_script("arguments[0].scrollIntoView({block: 'center'});", pick_element)
                        time.sleep(2)
                        
                        # Take direct element screenshot
                        screenshot_path = f"processed/final_screenshots/{author}_pick_{i:02d}.png"
                        pick_element.screenshot(screenshot_path)
                        
                        screenshots.append(screenshot_path)
                        print(f"   ✓ Pick {i} screenshot captured")
                        
                    except Exception as e:
                        print(f"   ⚠️ Error capturing pick {i}: {e}")
                        continue
            else:
                print(f"   ⚠️ No draft pick elements found for {author}")
                
        except Exception as e:
            print(f"   ⚠️ Error finding draft picks for {author}: {e}")
            
        return screenshots

    def screenshot_page_sections(self, author):
        """Take full page section screenshots"""
        screenshots = []
        
        try:
            # Scroll to top
            self.driver.execute_script("window.scrollTo(0, 0);")
            time.sleep(2)
            
            # Get page dimensions
            total_height = self.driver.execute_script("return document.body.scrollHeight")
            viewport_height = self.driver.execute_script("return window.innerHeight")
            
            # Take screenshots in sections
            section_height = viewport_height
            sections = min(6, max(1, total_height // section_height))  # Max 6 sections
            
            for i in range(sections):
                scroll_position = i * section_height
                
                # Scroll to position
                self.driver.execute_script(f"window.scrollTo(0, {scroll_position});")
                time.sleep(4)  # Wait for content to load
                
                # Take screenshot
                screenshot_path = f"processed/final_screenshots/{author}_section_{i+1:02d}.png"
                self.driver.save_screenshot(screenshot_path)
                
                screenshots.append(screenshot_path)
                print(f"   ✓ Section {i+1} screenshot captured")
                
        except Exception as e:
            print(f"   ⚠️ Error taking page sections for {author}: {e}")
            
        return screenshots

    def create_screenshot_document(self, all_screenshots):
        """Create Word document with all screenshots"""
        print("📄 Creating final Word document with screenshots...")
        
        doc = Document()
        
        # Set margins
        sections = doc.sections
        for section in sections:
            section.top_margin = Inches(0.5)
            section.bottom_margin = Inches(0.5)
            section.left_margin = Inches(0.5)
            section.right_margin = Inches(0.5)
        
        # Title
        title = doc.add_heading('NFL 2025 Mock Draft Collection', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_run = title.runs[0]
        title_run.font.size = Pt(24)
        title_run.font.color.rgb = RGBColor(0, 53, 148)
        
        # Subtitle
        subtitle = doc.add_paragraph(f'NFL.com Webpage Screenshots - Generated: {datetime.now().strftime("%B %d, %Y at %I:%M %p")}')
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        subtitle_run = subtitle.runs[0]
        subtitle_run.font.size = Pt(12)
        subtitle_run.font.color.rgb = RGBColor(107, 114, 128)
        
        # Group screenshots by author
        authors_screenshots = {}
        for screenshot_data in all_screenshots:
            author = screenshot_data['author']
            if author not in authors_screenshots:
                authors_screenshots[author] = []
            authors_screenshots[author].extend(screenshot_data['screenshots'])
        
        # Add each author's screenshots
        for author, screenshots in authors_screenshots.items():
            
            # Author section
            doc.add_page_break()
            
            author_para = doc.add_paragraph()
            author_run = author_para.add_run(f"🏈 {author} - NFL.com Mock Draft")
            author_run.font.size = Pt(20)
            author_run.font.bold = True
            author_run.font.color.rgb = RGBColor(0, 53, 148)
            
            # Add description
            desc_para = doc.add_paragraph(f"Screenshots captured from NFL.com showing draft picks and player information")
            desc_run = desc_para.runs[0]
            desc_run.font.size = Pt(10)
            desc_run.font.color.rgb = RGBColor(107, 114, 128)
            
            # Add each screenshot
            for screenshot_path in screenshots:
                try:
                    if os.path.exists(screenshot_path):
                        # Add screenshot with full width
                        doc.add_picture(screenshot_path, width=Inches(7.5))
                        
                        # Center the image
                        last_paragraph = doc.paragraphs[-1]
                        last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        last_paragraph.space_before = Pt(12)
                        last_paragraph.space_after = Pt(12)
                        
                        # Add small caption
                        caption = doc.add_paragraph(f"📸 {os.path.basename(screenshot_path)}")
                        caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        caption_run = caption.runs[0]
                        caption_run.font.size = Pt(8)
                        caption_run.font.color.rgb = RGBColor(128, 128, 128)
                        
                except Exception as e:
                    print(f"⚠️ Could not add screenshot {screenshot_path}: {e}")
        
        # Save document
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        output_path = f'processed/NFL_FINAL_SCREENSHOTS_{timestamp}.docx'
        doc.save(output_path)
        
        print(f"✓ Final screenshot document saved: {output_path}")
        return output_path

    def cleanup(self):
        """Clean up resources"""
        if self.driver:
            self.driver.quit()

def main():
    print("=== NFL Final Screenshot Creator ===")
    print("📸 Taking reliable screenshots of NFL.com pages")
    print("🎯 Capturing draft picks and player information")
    print("📄 Creating comprehensive Word document")
    print("=" * 50)
    
    creator = NFLScreenshotFinal()
    
    if not creator.driver:
        print("❌ Cannot proceed without WebDriver")
        return
    
    try:
        all_screenshots = []
        
        # Screenshot each author's webpage
        for author, url in creator.author_urls.items():
            screenshots = creator.screenshot_webpage_content(url, author)
            
            all_screenshots.append({
                'author': author,
                'screenshots': screenshots
            })
        
        # Create Word document with screenshots
        if all_screenshots:
            output_path = creator.create_screenshot_document(all_screenshots)
            
            print(f"\n🎉 SUCCESS! NFL.com screenshots captured!")
            print("=" * 50)
            print(f"📁 Document: {output_path}")
            print(f"📸 Screenshots: processed/final_screenshots/")
            
            total_screenshots = sum(len(author_data['screenshots']) for author_data in all_screenshots)
            print(f"\n📊 Summary:")
            print(f"   • {len(creator.author_urls)} authors processed")
            print(f"   • {total_screenshots} total screenshots captured")
            print(f"   • Player information and draft content preserved")
            print(f"   • Exact NFL.com visual layouts")
            print(f"   • Ready for viewing in Word")
        else:
            print("❌ No screenshots could be captured!")
    
    finally:
        creator.cleanup()

if __name__ == "__main__":
    main() 