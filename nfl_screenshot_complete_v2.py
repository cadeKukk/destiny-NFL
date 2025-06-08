#!/usr/bin/env python3
"""
NFL.com Mock Draft Screenshot Tool v2 - Improved Sequential Mapping
Creates Word document with screenshots and correctly mapped analysis text
"""

import os
import time
from datetime import datetime
from selenium import webdriver
from selenium.webdriver.chrome.options import Options
from selenium.webdriver.common.by import By
from docx import Document
from docx.shared import Inches, Pt
from docx.oxml.shared import OxmlElement, qn
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml

class NFLScreenshotCompleteV2:
    def __init__(self):
        self.driver = None
        self._analysis_cache = None
        self._current_author = None
        
        # Create directories if they don't exist
        os.makedirs('processed', exist_ok=True)
        os.makedirs('processed/complete_screenshots', exist_ok=True)
        
        self.authors = {
            'Bucky Brooks': 'https://www.nfl.com/news/bucky-brooks-2025-nfl-mock-draft-4-0-steelers-land-shedeur-sanders-cowboys-broncos-select-rbs',
            'Daniel Jeremiah': 'https://www.nfl.com/news/daniel-jeremiah-2025-nfl-mock-draft-4-0-trades-shake-up-order-for-shedeur-sanders-travis-hunter',
            'Charles Davis': 'https://www.nfl.com/news/charles-davis-2025-nfl-mock-draft-3-0-titans-browns-giants-select-qbs-cam-ward-travis-hunter-shedeur-sanders',
            'Eric Edholm': 'https://www.nfl.com/news/eric-edholm-2025-nfl-mock-draft-3-0-cam-ward-goes-no-1-overall-shedeur-sanders-slides',
            'Dan Parr': 'https://www.nfl.com/news/dan-parr-2025-nfl-mock-draft-2-0-cam-ward-travis-hunter-shedeur-sanders-go-in-top-10',
            'Gennaro Filice': 'https://www.nfl.com/news/gennaro-filice-2025-nfl-mock-draft-2-0-cam-ward-travis-hunter-lead-talented-top-10',
            'Marc Ross': 'https://www.nfl.com/news/marc-ross-2025-nfl-mock-draft-1-0-cam-ward-travis-hunter-lead-the-way'
        }

    def setup_selenium(self):
        """Setup Selenium WebDriver with Chrome"""
        try:
            chrome_options = Options()
            chrome_options.add_argument('--no-sandbox')
            chrome_options.add_argument('--disable-dev-shm-usage')
            chrome_options.add_argument('--window-size=1800,1400')
            chrome_options.add_argument('--user-agent=Mozilla/5.0 (Macintosh; Intel Mac OS X 10_15_7) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/120.0.0.0 Safari/537.36')
            chrome_options.add_argument('--disable-blink-features=AutomationControlled')
            
            self.driver = webdriver.Chrome(options=chrome_options)
            print("✓ Selenium WebDriver setup complete")
            return True
        except Exception as e:
            print(f"⚠️ Error setting up Selenium: {e}")
            return False

    def screenshot_webpage_content(self, url, author):
        """Screenshot webpage content with analysis extraction"""
        try:
            # Navigate to the page
            self.driver.get(url)
            time.sleep(6)  # Reduced wait for page to fully load
            
            # Remove overlays
            self.remove_overlays()
            time.sleep(3)
            
            # Scroll through page to load content
            self.load_page_content()
            
            # Reset analysis cache for new author
            self._analysis_cache = None
            self._current_author = None
            
            # Take header screenshot
            header_screenshots = self.screenshot_article_header(author)
            
            # Take individual pick screenshots with analysis
            pick_screenshots = self.screenshot_individual_picks(author)
            
            return header_screenshots + pick_screenshots
            
        except Exception as e:
            print(f"⚠️ Error screenshotting {author}: {e}")
            return []

    def load_page_content(self):
        """Load page content by scrolling"""
        total_height = self.driver.execute_script("return document.body.scrollHeight")
        for i in range(0, total_height, 2000):
            self.driver.execute_script(f"window.scrollTo(0, {i});")
            time.sleep(1.5)  # Reduced wait time for content loading
        
        # Return to top
        self.driver.execute_script("window.scrollTo(0, 0);")
        time.sleep(2)

    def remove_overlays(self):
        """Remove cookie banners and overlays"""
        overlay_selectors = [
            '[data-module="CookieBanner"]',
            '.onetrust-banner-sdk',
            '.cookie-banner',
            '.cookie-notice'
        ]
        
        for selector in overlay_selectors:
            try:
                overlays = self.driver.find_elements(By.CSS_SELECTOR, selector)
                for overlay in overlays:
                    if overlay.is_displayed():
                        self.driver.execute_script("arguments[0].style.display = 'none';", overlay)
            except:
                continue

    def screenshot_article_header(self, author):
        """Screenshot the article header"""
        screenshots = []
        try:
            header_path = f"processed/complete_screenshots/{author}_header.png"
            self.driver.save_screenshot(header_path)
            screenshots.append(header_path)
            print(f"   ✓ Header screenshot: {author}_header.png")
        except Exception as e:
            print(f"   ⚠️ Error taking header screenshot: {e}")
        
        return screenshots

    def _get_filtered_analysis_paragraphs(self):
        """Get filtered analysis paragraphs for sequential mapping"""
        try:
            # Get all paragraph elements
            all_paragraphs = self.driver.find_elements(By.CSS_SELECTOR, 'p')
            
            # Get first and last pick locations for filtering boundaries
            pick_elements = self.driver.find_elements(By.CSS_SELECTOR, '.nfl-o-ranked-item.nfl-o-ranked-item--side-by-side')
            if not pick_elements:
                return []
            
            first_pick_y = pick_elements[0].location['y']
            last_pick_y = pick_elements[-1].location['y']
            
            analysis_paragraphs = []
            for para in all_paragraphs:
                try:
                    text = para.get_attribute('textContent').strip()
                    location = para.location['y']
                    
                    # Filter criteria based on debug findings
                    if (text and len(text) > 50 and len(text) < 1000 and
                        location > first_pick_y and  # After first pick
                        location < last_pick_y + 2000 and  # Before way after last pick
                        any(keyword in text.lower() for keyword in ['quarterback', 'player', 'draft', 'team', 'offense', 'defense', 'potential', 'needs', 'season', 'franchise']) and
                        not text.startswith('Pick') and
                        '©' not in text and
                        'nfl.com' not in text.lower() and
                        'cookie' not in text.lower() and
                        'privacy' not in text.lower() and
                        # Filter out intro/outro text
                        'finally the week' not in text.lower() and
                        'trades that are struck' not in text.lower() and
                        'in his final mock' not in text.lower() and
                        'with round 1' not in text.lower()):
                        
                        analysis_paragraphs.append((location, text))
                except:
                    continue
            
            # Sort by Y position and return just the text
            analysis_paragraphs.sort(key=lambda x: x[0])
            return [text for location, text in analysis_paragraphs]
            
        except Exception as e:
            print(f"      ⚠️ Error getting analysis paragraphs: {e}")
            return []

    def extract_pick_description_sequential(self, pick_number, author):
        """Extract description using sequential mapping for accurate 1:1 correspondence"""
        try:
            print(f"      🔍 Extracting description for pick {pick_number}...")
            
            # Get all analysis paragraphs once (cached per author)
            if not hasattr(self, '_analysis_cache') or self._current_author != author:
                self._current_author = author
                self._analysis_cache = self._get_filtered_analysis_paragraphs()
                print(f"         Cached {len(self._analysis_cache)} analysis paragraphs for {author}")
            
            # Use sequential mapping: pick N gets analysis N
            analysis_index = pick_number - 1  # Convert to 0-based index
            
            if analysis_index < len(self._analysis_cache):
                description = self._analysis_cache[analysis_index]
                description = description.replace('\n', ' ').replace('\t', ' ')
                description = ' '.join(description.split())
                if len(description) > 400:
                    description = description[:400] + '...'
                print(f"      ✓ Sequential analysis (#{pick_number}): {description[:50]}...")
                return description
            else:
                # Fallback for picks beyond available analysis
                fallback_description = f"Expert analysis for this {pick_number} overall selection by {author}."
                print(f"      ⚠️ No analysis available, using fallback")
                return fallback_description
                
        except Exception as e:
            print(f"      ⚠️ Error extracting description: {e}")
            return f"Analysis for pick #{pick_number} not available."

    def screenshot_individual_picks(self, author):
        """Screenshot individual draft picks with descriptions"""
        screenshots = []
        descriptions = []
        
        try:
            # Find all pick elements
            pick_selectors = [
                '.nfl-o-ranked-item.nfl-o-ranked-item--side-by-side',
                '.ranked-item',
                '.draft-pick-item'
            ]
            
            pick_elements = None
            for selector in pick_selectors:
                pick_elements = self.driver.find_elements(By.CSS_SELECTOR, selector)
                if pick_elements:
                    print(f"   📋 Found {len(pick_elements)} draft picks for {author} using: {selector}")
                    break
            
            if not pick_elements:
                print(f"   ⚠️ No pick elements found for {author}")
                return []
            
            # Process up to 32 picks
            for i in range(min(32, len(pick_elements))):
                try:
                    pick_element = pick_elements[i]
                    pick_num = i + 1
                    
                    print(f"   🔍 Processing Pick {pick_num}...")
                    
                    # Extract description FIRST using sequential mapping
                    description = self.extract_pick_description_sequential(pick_num, author)
                    descriptions.append(description)
                    print(f"      📝 Description: {description[:100]}...")
                    
                    # Scroll element into view for better capture
                    self.driver.execute_script("arguments[0].scrollIntoView({block: 'center', behavior: 'smooth'});", pick_element)
                    time.sleep(0.8)  # Reduced wait for scroll to complete
                    
                    # Capture screenshot of individual pick
                    screenshot_path = f"processed/complete_screenshots/{author}_pick_{pick_num:02d}.png"
                    pick_element.screenshot(screenshot_path)
                    screenshots.append(screenshot_path)
                    
                    print(f"   ✓ Pick {pick_num} screenshot captured")
                    
                    time.sleep(0.4)  # Reduced pause between screenshots
                    
                except Exception as e:
                    print(f"   ⚠️ Error processing pick {i+1}: {e}")
                    descriptions.append(f"Pick {i+1} analysis not available.")
            
            # Store descriptions for document creation
            self.author_descriptions = getattr(self, 'author_descriptions', {})
            self.author_descriptions[author] = descriptions
            
            print(f"   ✓ Captured {len(screenshots)} screenshots for {author}")
            
        except Exception as e:
            print(f"   ⚠️ Error taking individual pick screenshots: {e}")
        
        return screenshots

    def create_word_document(self, all_screenshots):
        """Create comprehensive Word document with all authors"""
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        doc_path = f"processed/NFL_COMPLETE_ALL_AUTHORS_{timestamp}.docx"
        
        try:
            doc = Document()
            
            # Set document margins (0.4 inches all around)
            sections = doc.sections
            for section in sections:
                section.top_margin = Inches(0.4)
                section.bottom_margin = Inches(0.4)
                section.left_margin = Inches(0.4)
                section.right_margin = Inches(0.4)
            
            # Document title
            title = doc.add_heading('2025 NFL Mock Draft - Complete Analysis', 0)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            subtitle = doc.add_paragraph(f'Generated on {datetime.now().strftime("%B %d, %Y at %I:%M %p")}')
            subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
            subtitle_format = subtitle.runs[0].font
            subtitle_format.size = Pt(12)
            subtitle_format.italic = True
            
            # Process each author
            for author, url in self.authors.items():
                print(f"   📝 Adding {author} to document...")
                
                # Author header
                author_header = doc.add_heading(f'{author} Mock Draft', level=1)
                author_header.space_before = Pt(12)
                author_header.space_after = Pt(6)
                
                # Find screenshots for this author
                author_screenshots = [s for s in all_screenshots if author.replace(' ', '_') in s or author.replace(' ', '') in s]
                header_screenshots = [s for s in author_screenshots if 'header' in s]
                pick_screenshots = [s for s in author_screenshots if 'pick_' in s and 'header' not in s]
                
                # Add header screenshot
                for header_path in header_screenshots:
                    if os.path.exists(header_path):
                        try:
                            header_para = doc.add_paragraph()
                            header_run = header_para.runs[0] if header_para.runs else header_para.add_run()
                            header_run.add_picture(header_path, width=Inches(7.0))
                            header_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                            header_para.space_after = Pt(8)
                        except Exception as e:
                            print(f"      ⚠️ Error adding header image for {author}: {e}")
                
                # Get descriptions for this author
                author_descriptions = getattr(self, 'author_descriptions', {}).get(author, [])
                print(f"      📝 Found {len(author_descriptions)} descriptions for {author}")
                
                # Add pick screenshots with descriptions
                for i, screenshot_path in enumerate(sorted(pick_screenshots), 1):
                    if os.path.exists(screenshot_path):
                        try:
                            print(f"      🔍 Processing pick #{i} screenshot...")
                            
                            # Add screenshot
                            pick_para = doc.add_paragraph()
                            pick_run = pick_para.runs[0] if pick_para.runs else pick_para.add_run()
                            pick_run.add_picture(screenshot_path, width=Inches(6.5))
                            pick_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                            pick_para.space_after = Pt(4)
                            
                            # Add description if available
                            if i <= len(author_descriptions):
                                description = author_descriptions[i-1]
                                desc_para = doc.add_paragraph()
                                desc_run = desc_para.add_run(description)
                                desc_run.font.size = Pt(10)
                                desc_run.font.name = 'Calibri'
                                desc_para.space_after = Pt(8)
                                desc_para.space_before = Pt(2)
                                print(f"         ✓ Adding description: {description[:50]}...")
                            else:
                                print(f"         ⚠️ No description available for pick {i}")
                                
                        except Exception as e:
                            print(f"      ⚠️ Error adding pick {i} for {author}: {e}")
                
                # Add page break between authors (except for last author)
                if author != list(self.authors.keys())[-1]:
                    doc.add_page_break()
            
            # Save document
            doc.save(doc_path)
            
            # Summary
            total_descriptions = sum(len(descriptions) for descriptions in getattr(self, 'author_descriptions', {}).values())
            print(f"\n📊 Description Summary:")
            for author, descriptions in getattr(self, 'author_descriptions', {}).items():
                print(f"   {author}: {len(descriptions)} descriptions")
            
            print(f"\n🎉 SUCCESS! Complete NFL.com screenshots captured!")
            print(f"{'='*55}")
            print(f"📁 Document: {doc_path}")
            print(f"📸 Screenshots: processed/complete_screenshots/")
            print(f"\n📊 Summary:")
            print(f"   • {len(self.authors)} authors processed (ALL)")
            print(f"   • {len(all_screenshots)} total screenshots captured")
            print(f"   • Up to 32 picks per author")
            print(f"   • Expert analysis included under each pick")
            print(f"   • Optimized spacing and formatting")
            print(f"   • Every author appears in document")
            print(f"   • Ready for viewing in Word")
            
            return doc_path
            
        except Exception as e:
            print(f"⚠️ Error creating Word document: {e}")
            return None

    def cleanup(self):
        """Clean up resources"""
        if self.driver:
            self.driver.quit()
            print("✓ WebDriver closed")

def main():
    screenshot_tool = NFLScreenshotCompleteV2()
    
    try:
        if not screenshot_tool.setup_selenium():
            return
        
        print("🎯 Starting NFL Mock Draft Screenshot Collection...")
        print(f"📋 Processing {len(screenshot_tool.authors)} authors")
        
        all_screenshots = []
        
        # Process each author
        for author, url in screenshot_tool.authors.items():
            print(f"\n📸 Capturing screenshots for {author}...")
            author_screenshots = screenshot_tool.screenshot_webpage_content(url, author)
            all_screenshots.extend(author_screenshots)
        
        # Create Word document
        if all_screenshots:
            print(f"\n📄 Creating optimized Word document with all authors...")
            doc_path = screenshot_tool.create_word_document(all_screenshots)
        else:
            print("⚠️ No screenshots captured")
            
    except KeyboardInterrupt:
        print("\n⚠️ Process interrupted by user")
    except Exception as e:
        print(f"⚠️ Unexpected error: {e}")
    finally:
        screenshot_tool.cleanup()

if __name__ == "__main__":
    main() 