#!/usr/bin/env python3
"""
Create Clean Layout NFL Mock Draft Document
Simple format matching the image layout
"""

from docx import Document
from docx.shared import Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from datetime import datetime
import os

def create_clean_layout_document():
    """Create a clean, simple Word document matching the image layout"""
    
    # Create processed folder
    os.makedirs('processed', exist_ok=True)
    
    # Mock drafts with exact format from your image
    mock_drafts = [
        {
            'title': 'Bucky Brooks 2025 NFL Mock Draft 3.0',
            'author': 'Bucky Brooks',
            'picks': [
                {
                    'pick': 1, 
                    'team': 'Tennessee Titans', 
                    'player': 'Cam Ward', 
                    'position': 'QB', 
                    'school': 'Miami',
                    'class': 'Senior',
                    'description': 'The talented passer gives Brian Callahan the franchise quarterback needed to spark the Titans\' rebuild.'
                },
                {
                    'pick': 2, 
                    'team': 'Cleveland Browns', 
                    'player': 'Shedeur Sanders', 
                    'position': 'QB', 
                    'school': 'Colorado',
                    'class': 'Senior',
                    'description': 'Kevin Stefanski has worked well with traditional pocket passers throughout his career. Sanders fits the bill as a classic dropback quarterback with a game built on touch, timing and anticipation.'
                },
                {
                    'pick': 3, 
                    'team': 'New York Giants', 
                    'player': 'Travis Hunter', 
                    'position': 'WR/CB', 
                    'school': 'Colorado',
                    'class': 'Junior',
                    'description': 'Adding a two-way standout doesn\'t solve the Giants\' most pressing need, but Hunter\'s playmaking presence would help the offense and defense improve.'
                },
                {
                    'pick': 4, 
                    'team': 'New England Patriots', 
                    'player': 'Abdul Carter', 
                    'position': 'Edge', 
                    'school': 'Penn State',
                    'class': 'Junior',
                    'description': 'If Patriots personnel chief Eliot Wolf is truly committed to taking the best player available, the Penn State product would be a no-brainer at this point.'
                }
            ]
        },
        {
            'title': 'Charles Davis 2025 NFL Mock Draft 3.0',
            'author': 'Charles Davis',
            'picks': [
                {
                    'pick': 1, 
                    'team': 'Tennessee Titans', 
                    'player': 'Cam Ward', 
                    'position': 'QB', 
                    'school': 'Miami',
                    'class': 'Senior',
                    'description': 'Ward has the most NFL-ready skill set of any quarterback in this class. His accuracy and leadership make him the clear choice for Tennessee.'
                },
                {
                    'pick': 2, 
                    'team': 'Cleveland Browns', 
                    'player': 'Abdul Carter', 
                    'position': 'Edge', 
                    'school': 'Penn State',
                    'class': 'Junior',
                    'description': 'Carter brings elite pass rush ability and versatility. His combination of speed and power makes him a game-changing defender.'
                },
                {
                    'pick': 3, 
                    'team': 'New York Giants', 
                    'player': 'Will Johnson', 
                    'position': 'CB', 
                    'school': 'Michigan',
                    'class': 'Junior',
                    'description': 'Johnson has lockdown coverage ability and the physicality to match up with elite receivers in today\'s NFL.'
                }
            ]
        }
    ]
    
    # Create the Word document
    doc = Document()
    
    # Title
    title = doc.add_heading('NFL 2025 Mock Draft Picks', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph(f"Generated: {datetime.now().strftime('%B %d, %Y')}")
    doc.add_paragraph("")
    
    # Process each mock draft
    for draft in mock_drafts:
        # Author header
        doc.add_heading(f"{draft['author']} Mock Draft", level=1)
        doc.add_paragraph(f"{draft['title']}")
        doc.add_paragraph("")
        
        # Create each pick in clean format
        for pick in draft['picks']:
            
            # Pick header with number and team
            pick_header = doc.add_paragraph()
            pick_run = pick_header.add_run(f"Pick {pick['pick']}   ")
            pick_run.font.bold = True
            pick_run.font.size = Inches(0.16)
            pick_run.font.color.rgb = RGBColor(255, 255, 255)  # White text
            
            team_run = pick_header.add_run(f"{pick['team']}")
            team_run.font.bold = True
            team_run.font.size = Inches(0.16)
            team_run.font.color.rgb = RGBColor(0, 0, 0)  # Black text
            
            # Player name in large blue text
            player_para = doc.add_paragraph()
            player_run = player_para.add_run(f"{pick['player']}")
            player_run.font.bold = True
            player_run.font.size = Inches(0.2)
            player_run.font.color.rgb = RGBColor(37, 99, 235)  # Blue color
            
            # Player details
            details_para = doc.add_paragraph()
            details_run = details_para.add_run(f"{pick['school']} • {pick['position']} • {pick['class']}")
            details_run.font.size = Inches(0.13)
            details_run.font.color.rgb = RGBColor(107, 114, 128)  # Gray color
            
            # Player image placeholder
            image_para = doc.add_paragraph()
            image_run = image_para.add_run("[Player Photo]")
            image_run.font.size = Inches(0.11)
            image_run.font.color.rgb = RGBColor(156, 163, 175)
            image_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Description
            desc_para = doc.add_paragraph()
            desc_run = desc_para.add_run(pick['description'])
            desc_run.font.size = Inches(0.14)
            
            # Add spacing
            doc.add_paragraph("")
            doc.add_paragraph("─" * 60)
            doc.add_paragraph("")
        
        # Page break between drafts  
        doc.add_page_break()
    
    # Summary
    doc.add_heading('Summary', level=1)
    doc.add_paragraph("This document contains individual player picks from multiple NFL mock drafts by top analysts.")
    
    all_players = []
    for draft in mock_drafts:
        for pick in draft['picks']:
            all_players.append(pick['player'])
    
    from collections import Counter
    player_counts = Counter(all_players)
    most_common = player_counts.most_common(5)
    
    doc.add_heading('Most Popular Players:', level=2)
    for i, (player, count) in enumerate(most_common, 1):
        doc.add_paragraph(f"{i}. {player} - Selected {count} time(s)")
    
    # Save the document
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    output_path = f'processed/NFL_Mock_Drafts_CLEAN_LAYOUT_{timestamp}.docx'
    doc.save(output_path)
    
    return output_path

def main():
    print("=== Creating Clean Layout NFL Mock Draft Document ===")
    print("✓ Simple format matching your image")
    print("✓ Pick number and team name")
    print("✓ Player name in blue highlighting")
    print("✓ Player details (school • position • class)")
    print("✓ Player photo placeholder")
    print("✓ Brief description")
    print("=" * 55)
    
    try:
        output_path = create_clean_layout_document()
        
        print("🎉 SUCCESS! Clean layout document created!")
        print("=" * 55)
        print(f"📁 Document saved: {output_path}")
        
        print("\n🎨 Clean Format Features:")
        print("   • Pick numbers and team names")
        print("   • Player names in blue")
        print("   • School • Position • Class format")
        print("   • Photo placeholders")
        print("   • Brief descriptions")
        print("   • Clean spacing and separators")
        
        print(f"\n📖 Simplified layout based on your image!")
        
    except Exception as e:
        print(f"❌ Error: {e}")

if __name__ == "__main__":
    main() 