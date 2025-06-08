#!/usr/bin/env python3
"""
NFL Mock Draft Scraper
Scrapes mock draft data from NFL.com and processes it according to specified authors
"""

import requests
from bs4 import BeautifulSoup
import pandas as pd
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os
import re
from datetime import datetime
import time
from urllib.parse import urljoin, urlparse
import json

class NFLMockDraftScraper:
    def __init__(self):
        self.base_url = "https://www.nfl.com"
        self.session = requests.Session()
        self.session.headers.update({
            'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/91.0.4472.124 Safari/537.36'
        })
        
        # Author names from the spreadsheets
        self.target_authors = [
            "Charles Davis",
            "Chad Reuter", 
            "Bucky Brooks",
            "SI Contributors",
            "Mike Band",
            "Lance Zierlien",
            "Gennaro Filice",
            "Eric Edholm",
            "Daniel Jeremiah",
            "Dan Parr",
            "Cynthia Frelund"
        ]
        
        self.mock_drafts = []
        
    def get_page_content(self, url):
        """Fetch page content with error handling"""
        try:
            response = self.session.get(url, timeout=30)
            response.raise_for_status()
            return response.text
        except requests.RequestException as e:
            print(f"Error fetching {url}: {e}")
            return None
            
    def extract_mock_draft_data(self, url):
        """Extract mock draft data from the specific URL"""
        print(f"Scraping: {url}")
        
        content = self.get_page_content(url)
        if not content:
            return []
            
        soup = BeautifulSoup(content, 'html.parser')
        mock_drafts = []
        
        # Look for article content and author information
        article = soup.find('article') or soup.find('div', class_='nfl-c-article')
        if not article:
            article = soup
            
        # Extract title
        title_elem = soup.find('h1') or soup.find('title')
        title = title_elem.get_text(strip=True) if title_elem else "NFL Mock Draft"
        
        # Extract author
        author_elem = soup.find('span', class_='nfl-c-author__name') or \
                     soup.find('div', class_='author') or \
                     soup.find('a', href=re.compile(r'/author/'))
        
        author = "Unknown"
        if author_elem:
            author = author_elem.get_text(strip=True)
        
        # Extract date
        date_elem = soup.find('time') or soup.find('span', class_='date')
        date = datetime.now().strftime("%Y-%m-%d")
        if date_elem:
            date_text = date_elem.get('datetime') or date_elem.get_text(strip=True)
            try:
                # Try to parse various date formats
                if 'datetime' in str(date_elem.attrs):
                    date = date_elem['datetime'][:10]  # Extract YYYY-MM-DD
                else:
                    date = date_text
            except:
                pass
        
        # Extract draft picks
        picks = self.extract_draft_picks(soup)
        
        mock_draft = {
            'title': title,
            'author': author,
            'date': date,
            'url': url,
            'picks': picks
        }
        
        return [mock_draft]
    
    def extract_draft_picks(self, soup):
        """Extract individual draft picks from the page"""
        picks = []
        
        # Look for various patterns of draft pick information
        pick_patterns = [
            # Pattern 1: Numbered lists or ordered lists
            soup.find_all(['ol', 'ul']),
            
            # Pattern 2: Divs with pick information
            soup.find_all('div', class_=re.compile(r'pick|draft', re.I)),
            
            # Pattern 3: Paragraphs with pick numbers
            soup.find_all('p', string=re.compile(r'\d+\.\s+'))
        ]
        
        # Extract from paragraphs containing numbered picks
        paragraphs = soup.find_all('p')
        for p in paragraphs:
            text = p.get_text(strip=True)
            
            # Look for patterns like "1. Team - Player, Position, School"
            pick_match = re.match(r'(\d+)\.\s*([^-]+)\s*-\s*([^,]+),?\s*([^,]*),?\s*(.*)', text)
            if pick_match:
                pick_num, team, player, position, school = pick_match.groups()
                picks.append({
                    'pick': int(pick_num),
                    'team': team.strip(),
                    'player': player.strip(),
                    'position': position.strip(),
                    'school': school.strip()
                })
        
        # If no picks found in paragraphs, try other methods
        if not picks:
            # Look for table data
            tables = soup.find_all('table')
            for table in tables:
                rows = table.find_all('tr')[1:]  # Skip header
                for i, row in enumerate(rows, 1):
                    cells = row.find_all(['td', 'th'])
                    if len(cells) >= 3:
                        picks.append({
                            'pick': i,
                            'team': cells[0].get_text(strip=True),
                            'player': cells[1].get_text(strip=True),
                            'position': cells[2].get_text(strip=True) if len(cells) > 2 else '',
                            'school': cells[3].get_text(strip=True) if len(cells) > 3 else ''
                        })
        
        return picks[:32]  # Limit to first round if many picks found
    
    def find_related_mock_drafts(self, initial_url):
        """Find other 2025 mock draft articles"""
        content = self.get_page_content(initial_url)
        if not content:
            return [initial_url]
            
        soup = BeautifulSoup(content, 'html.parser')
        mock_draft_urls = {initial_url}
        
        # Look for links to other mock drafts
        links = soup.find_all('a', href=True)
        for link in links:
            href = link['href']
            text = link.get_text(strip=True).lower()
            
            # Check if link is related to 2025 mock drafts
            if ('2025' in text and 'mock' in text and 'draft' in text) or \
               ('2025' in href and 'mock' in href and 'draft' in href):
                
                full_url = urljoin(self.base_url, href)
                mock_draft_urls.add(full_url)
        
        return list(mock_draft_urls)
    
    def scrape_all_mock_drafts(self, initial_url):
        """Scrape all related mock draft articles"""
        urls = self.find_related_mock_drafts(initial_url)
        all_mock_drafts = []
        
        for url in urls:
            mock_drafts = self.extract_mock_draft_data(url)
            all_mock_drafts.extend(mock_drafts)
            time.sleep(1)  # Be respectful to the server
            
        return all_mock_drafts
    
    def filter_by_authors(self, mock_drafts):
        """Filter mock drafts by target authors"""
        filtered_drafts = []
        
        for draft in mock_drafts:
            author = draft['author']
            
            # Check if author matches any of our target authors (case insensitive)
            for target_author in self.target_authors:
                if target_author.lower() in author.lower() or \
                   any(word in author.lower() for word in target_author.lower().split()):
                    filtered_drafts.append(draft)
                    break
                    
        return filtered_drafts
    
    def create_word_document(self, mock_drafts, output_path):
        """Create a Word document with the mock draft data"""
        doc = Document()
        
        # Add title
        title = doc.add_heading('NFL 2025 Mock Draft Data', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        doc.add_paragraph(f"Generated on: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
        doc.add_paragraph(f"Total Mock Drafts: {len(mock_drafts)}")
        
        # Add table of contents
        doc.add_heading('Mock Drafts Included:', level=1)
        for i, draft in enumerate(mock_drafts, 1):
            doc.add_paragraph(f"{i}. {draft['author']} - {draft['title']} ({draft['date']})")
        
        doc.add_page_break()
        
        # Add each mock draft
        for draft in mock_drafts:
            # Draft header
            doc.add_heading(f"{draft['author']} - Mock Draft", level=1)
            doc.add_paragraph(f"Title: {draft['title']}")
            doc.add_paragraph(f"Date: {draft['date']}")
            doc.add_paragraph(f"URL: {draft['url']}")
            
            # Draft picks
            if draft['picks']:
                doc.add_heading('Draft Picks:', level=2)
                
                # Create table for picks
                table = doc.add_table(rows=1, cols=5)
                table.style = 'Light Grid Accent 1'
                
                # Header row
                header_cells = table.rows[0].cells
                header_cells[0].text = 'Pick'
                header_cells[1].text = 'Team'
                header_cells[2].text = 'Player'
                header_cells[3].text = 'Position'
                header_cells[4].text = 'School'
                
                # Add picks
                for pick in draft['picks']:
                    row_cells = table.add_row().cells
                    row_cells[0].text = str(pick.get('pick', ''))
                    row_cells[1].text = pick.get('team', '')
                    row_cells[2].text = pick.get('player', '')
                    row_cells[3].text = pick.get('position', '')
                    row_cells[4].text = pick.get('school', '')
            else:
                doc.add_paragraph("No draft picks found or extracted.")
            
            doc.add_page_break()
        
        # Save document
        doc.save(output_path)
        print(f"Word document saved to: {output_path}")
    
    def run(self, url):
        """Main execution method"""
        print("Starting NFL Mock Draft Scraper...")
        print(f"Target URL: {url}")
        print(f"Target Authors: {', '.join(self.target_authors)}")
        
        # Create processed folder
        os.makedirs('processed', exist_ok=True)
        
        # Scrape mock drafts
        all_mock_drafts = self.scrape_all_mock_drafts(url)
        print(f"Found {len(all_mock_drafts)} total mock drafts")
        
        # Filter by authors
        filtered_drafts = self.filter_by_authors(all_mock_drafts)
        print(f"Found {len(filtered_drafts)} mock drafts from target authors")
        
        if not filtered_drafts:
            print("No mock drafts found from target authors. Including all found drafts.")
            filtered_drafts = all_mock_drafts
        
        # Create Word document
        output_path = os.path.join('processed', f'NFL_Mock_Drafts_2025_{datetime.now().strftime("%Y%m%d_%H%M%S")}.docx')
        self.create_word_document(filtered_drafts, output_path)
        
        print("Scraping completed successfully!")
        return filtered_drafts

def main():
    url = "https://www.nfl.com/news/bucky-brooks-2025-nfl-mock-draft-3-0-browns-take-shedeur-sanders-two-running-backs-in-top-10-picks"
    
    scraper = NFLMockDraftScraper()
    mock_drafts = scraper.run(url)
    
    # Print summary
    print(f"\nSummary:")
    print(f"Total mock drafts processed: {len(mock_drafts)}")
    for draft in mock_drafts:
        print(f"- {draft['author']}: {len(draft['picks'])} picks")

if __name__ == "__main__":
    main() 