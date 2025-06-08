#!/usr/bin/env python3
"""
NFL.com Mock Draft Screenshot Tool - Ultra Simple
Takes element screenshots with extra scrolling to ensure description text is visible
"""

import os
import time
from datetime import datetime
from selenium import webdriver
from selenium.webdriver.chrome.options import Options
from selenium.webdriver.common.by import By
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

class NFLScreenshotUltraSimple:
    def __init__(self):
        self.driver = None
        
        # Create directories if they don't exist
        os.makedirs('processed', exist_ok=True)
        os.makedirs('processed/complete_screenshots', exist_ok=True)
        
        # Updated URLs with corrections
        self.authors = {
            'Bucky Brooks': 'https://www.nfl.com/news/bucky-brooks-2025-nfl-mock-draft-4-0-steelers-land-shedeur-sanders-cowboys-broncos-select-rbs',
            'Daniel Jeremiah': 'https://www.nfl.com/news/daniel-jeremiah-2025-nfl-mock-draft-4-0',
            'Lance Zierlein': 'https://www.nfl.com/news/lance-zierlein-2025-nfl-mock-draft-4-0-colts-trade-up-for-colston-loveland-saints-go-get-jaxson-dart',
            'Charles Davis': 'https://www.nfl.com/news/charles-davis-2025-nfl-mock-draft-3-0-cam-ward-only-qb-in-round-1-eagles-pick-te-mason-taylor',
            'Eric Edholm': 'https://www.nfl.com/news/eric-edholm-2025-nfl-mock-draft-3-0-four-first-round-quarterbacks-jaguars-take-rb-ashton-jeanty',
            'Dan Parr': 'https://www.nfl.com/news/dan-parr-2025-nfl-mock-draft-2-0-offensive-linemen-dominate-top-10-bears-grab-tight-end-tyler-warren',
            'Chad Reuter': 'https://www.nfl.com/news/seven-round-2025-nfl-mock-draft-round-one',
            'Gennaro Filice': 'https://www.nfl.com/news/gennaro-filice-2025-nfl-mock-draft-2-0-rb-ashton-jeanty-goes-top-5-cowboys-jump-for-jalon-walker',
            'Marc Ross': 'https://www.nfl.com/news/marc-ross-2025-nfl-mock-draft-1-0-three-qbs-selected-in-top-10-jets-snag-rb-ashton-jeanty'
        }

    def setup_selenium(self):
        """Setup Selenium WebDriver with Chrome"""
        try:
            chrome_options = Options()
            chrome_options.add_argument('--no-sandbox')
            chrome_options.add_argument('--disable-dev-shm-usage')
            chrome_options.add_argument('--window-size=1800,1400')
            chrome_options.add_argument('--user-agent=Mozilla/5.0 (Macintosh; Intel Mac OS X 10_15_7) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/120.0.0.0 Safari/537.36')
            chrome_options.add_argument('--disable-blink-features=AutomationControlled')
            
            self.driver = webdriver.Chrome(options=chrome_options)
            print("✓ Selenium WebDriver setup complete")
            return True
        except Exception as e:
            print(f"⚠️ Error setting up Selenium: {e}")
            return False

    def screenshot_webpage_content(self, url, author):
        """Screenshot webpage content"""
        try:
            print(f"🔍 Loading: {url}")
            # Navigate to the page
            self.driver.get(url)
            time.sleep(6)  # Wait for page to fully load
            
            # Remove overlays
            self.remove_overlays()
            time.sleep(3)
            
            # Scroll through page to load content
            self.load_page_content()
            
            # Take header screenshot
            header_screenshots = self.screenshot_article_header(author)
            
            # Take individual pick screenshots
            pick_screenshots = self.screenshot_individual_picks_simple(author)
            
            return header_screenshots + pick_screenshots
            
        except Exception as e:
            print(f"⚠️ Error screenshotting {author}: {e}")
            return []

    def load_page_content(self):
        """Load page content by scrolling"""
        total_height = self.driver.execute_script("return document.body.scrollHeight")
        for i in range(0, total_height, 2000):
            self.driver.execute_script(f"window.scrollTo(0, {i});")
            time.sleep(0.8)
        
        # Return to top
        self.driver.execute_script("window.scrollTo(0, 0);")
        time.sleep(2)

    def remove_overlays(self):
        """Remove cookie banners and overlays"""
        overlay_selectors = [
            '[data-module="CookieBanner"]',
            '.onetrust-banner-sdk',
            '.cookie-banner',
            '.cookie-notice'
        ]
        
        for selector in overlay_selectors:
            try:
                overlays = self.driver.find_elements(By.CSS_SELECTOR, selector)
                for overlay in overlays:
                    if overlay.is_displayed():
                        self.driver.execute_script("arguments[0].style.display = 'none';", overlay)
            except:
                continue

    def screenshot_article_header(self, author):
        """Screenshot the article header"""
        screenshots = []
        try:
            header_path = f"processed/complete_screenshots/{author}_header.png"
            self.driver.save_screenshot(header_path)
            screenshots.append(header_path)
            print(f"   ✓ Header screenshot: {author}_header.png")
        except Exception as e:
            print(f"   ⚠️ Error taking header screenshot: {e}")
        
        return screenshots

    def screenshot_individual_picks_simple(self, author):
        """Screenshot individual draft picks with simple approach to capture description"""
        screenshots = []
        
        try:
            # Find pick elements
            pick_elements = self.driver.find_elements(By.CSS_SELECTOR, '.nfl-o-ranked-item.nfl-o-ranked-item--side-by-side')
            
            if not pick_elements:
                print(f"   ⚠️ No pick elements found for {author}")
                return []
            
            print(f"   📋 Processing {min(32, len(pick_elements))} picks for {author}")
            
            # Process up to 32 picks
            for i in range(min(32, len(pick_elements))):
                try:
                    pick_element = pick_elements[i]
                    pick_num = i + 1
                    
                    print(f"   🔍 Processing Pick {pick_num}...")
                    
                    # Scroll element into view with extra space below for description
                    # Position element higher in viewport so description is visible
                    self.driver.execute_script("""
                        var element = arguments[0];
                        var rect = element.getBoundingClientRect();
                        var scrollTop = window.pageYOffset || document.documentElement.scrollTop;
                        var targetY = rect.top + scrollTop - 100;  // Position element 100px from top
                        window.scrollTo(0, targetY);
                    """, pick_element)
                    
                    time.sleep(0.8)  # Wait for scroll and any lazy loading
                    
                    # Try to find and capture a larger container that includes description
                    # Look for parent containers that might include the analysis text
                    larger_containers = [
                        pick_element.find_element(By.XPATH, './ancestor::div[contains(@class, "nfl-o-ranked-item")]'),
                        pick_element.find_element(By.XPATH, './parent::*'),
                        pick_element
                    ]
                    
                    # Use the largest container we can find
                    container_to_screenshot = pick_element
                    for container in larger_containers:
                        try:
                            if container.size['height'] > container_to_screenshot.size['height']:
                                container_to_screenshot = container
                        except:
                            continue
                    
                    # Take screenshot of the container (which should include description)
                    screenshot_path = f"processed/complete_screenshots/{author}_pick_{pick_num:02d}.png"
                    container_to_screenshot.screenshot(screenshot_path)
                    screenshots.append(screenshot_path)
                    
                    print(f"   ✓ Pick {pick_num} screenshot captured with description")
                    
                    time.sleep(0.3)  # Brief pause between screenshots
                    
                except Exception as e:
                    print(f"   ⚠️ Error processing pick {i+1}: {e}")
                    # Fallback to basic element screenshot
                    try:
                        screenshot_path = f"processed/complete_screenshots/{author}_pick_{pick_num:02d}.png"
                        pick_element.screenshot(screenshot_path)
                        screenshots.append(screenshot_path)
                        print(f"   ✓ Pick {pick_num} fallback screenshot captured")
                    except:
                        print(f"   ❌ Failed to capture pick {pick_num}")
            
            print(f"   ✓ Captured {len(screenshots)} screenshots for {author}")
            
        except Exception as e:
            print(f"   ⚠️ Error taking individual pick screenshots: {e}")
        
        return screenshots

    def create_word_document(self, all_screenshots):
        """Create comprehensive Word document with all authors"""
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        doc_path = f"processed/NFL_ULTRA_SIMPLE_{timestamp}.docx"
        
        try:
            doc = Document()
            
            # Set document margins (0.4 inches all around)
            sections = doc.sections
            for section in sections:
                section.top_margin = Inches(0.4)
                section.bottom_margin = Inches(0.4)
                section.left_margin = Inches(0.4)
                section.right_margin = Inches(0.4)
            
            # Document title
            title = doc.add_heading('2025 NFL Mock Draft - Ultra Simple Screenshots', 0)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            subtitle = doc.add_paragraph(f'Generated on {datetime.now().strftime("%B %d, %Y at %I:%M %p")} - Pick screenshots with descriptions')
            subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
            subtitle_format = subtitle.runs[0].font
            subtitle_format.size = Pt(12)
            subtitle_format.italic = True
            
            # Process each author
            for author, url in self.authors.items():
                print(f"   📝 Adding {author} to document...")
                
                # Author header
                author_header = doc.add_heading(f'{author} Mock Draft', level=1)
                author_header.space_before = Pt(12)
                author_header.space_after = Pt(6)
                
                # Find screenshots for this author
                author_screenshots = [s for s in all_screenshots if author in s]
                header_screenshots = [s for s in author_screenshots if 'header' in s]
                pick_screenshots = [s for s in author_screenshots if 'pick_' in s and 'header' not in s]
                
                # Add header screenshot
                for header_path in header_screenshots:
                    if os.path.exists(header_path):
                        try:
                            header_para = doc.add_paragraph()
                            header_run = header_para.runs[0] if header_para.runs else header_para.add_run()
                            header_run.add_picture(header_path, width=Inches(7.0))
                            header_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                            header_para.space_after = Pt(8)
                        except Exception as e:
                            print(f"      ⚠️ Error adding header image for {author}: {e}")
                
                # Add pick screenshots
                for i, screenshot_path in enumerate(sorted(pick_screenshots), 1):
                    if os.path.exists(screenshot_path):
                        try:
                            # Add screenshot
                            pick_para = doc.add_paragraph()
                            pick_run = pick_para.runs[0] if pick_para.runs else pick_para.add_run()
                            pick_run.add_picture(screenshot_path, width=Inches(6.5))
                            pick_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                            pick_para.space_after = Pt(6)
                            
                        except Exception as e:
                            print(f"      ⚠️ Error adding pick {i} for {author}: {e}")
                
                # Add page break between authors (except for last author)
                if author != list(self.authors.keys())[-1]:
                    doc.add_page_break()
            
            # Save document
            doc.save(doc_path)
            
            print(f"\n🎉 SUCCESS! Ultra simple NFL.com screenshots captured!")
            print(f"{'='*55}")
            print(f"📁 Document: {doc_path}")
            print(f"📸 Screenshots: processed/complete_screenshots/")
            print(f"\n📊 Summary:")
            print(f"   • {len(self.authors)} authors processed")
            print(f"   • {len(all_screenshots)} total screenshots captured")
            print(f"   • Up to 32 picks per author")
            print(f"   • Simple element screenshots with positioning for descriptions")
            
            return doc_path
            
        except Exception as e:
            print(f"⚠️ Error creating Word document: {e}")
            return None

    def cleanup(self):
        """Clean up resources"""
        if self.driver:
            self.driver.quit()
            print("✓ WebDriver closed")

def main():
    screenshot_tool = NFLScreenshotUltraSimple()
    
    try:
        if not screenshot_tool.setup_selenium():
            return
        
        print("🎯 Starting NFL Mock Draft Screenshot Collection - ULTRA SIMPLE VERSION...")
        print(f"📋 Processing {len(screenshot_tool.authors)} authors")
        
        all_screenshots = []
        
        # Process each author
        for author, url in screenshot_tool.authors.items():
            print(f"\n📸 Capturing screenshots for {author}...")
            author_screenshots = screenshot_tool.screenshot_webpage_content(url, author)
            all_screenshots.extend(author_screenshots)
        
        # Create Word document
        if all_screenshots:
            print(f"\n📄 Creating Word document...")
            doc_path = screenshot_tool.create_word_document(all_screenshots)
        else:
            print("⚠️ No screenshots captured")
            
    except KeyboardInterrupt:
        print("\n⚠️ Process interrupted by user")
    except Exception as e:
        print(f"⚠️ Unexpected error: {e}")
    finally:
        screenshot_tool.cleanup()

if __name__ == "__main__":
    main() 