#!/usr/bin/env python3
"""
NFL Webpage Replica Creator
Creates Word documents that exactly mirror NFL.com mock draft webpages
"""

import os
from datetime import datetime
from docx import Document
from docx.shared import Inches, RGBColor, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_PARAGRAPH_ALIGNMENT
from docx.oxml.shared import OxmlElement, qn

def create_nfl_webpage_replica():
    """Create exact NFL.com webpage replica in Word format"""
    
    print("📄 Creating NFL.com webpage replica...")
    
    doc = Document()
    
    # Set webpage-like margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)
    
    # Add webpage header (like NFL.com)
    header_para = doc.add_paragraph()
    header_run = header_para.add_run("Mock Draft")
    header_run.font.size = Pt(14)
    header_run.font.color.rgb = RGBColor(0, 53, 148)  # NFL Blue
    header_run.font.bold = True
    
    # Main title (exactly like the webpage)
    title_para = doc.add_heading('', level=1)
    title_run = title_para.add_run("Bucky Brooks 2025 NFL mock draft 4.0: Steelers land Shedeur Sanders; Cowboys, Broncos select RBs")
    title_run.font.size = Pt(32)
    title_run.font.color.rgb = RGBColor(0, 0, 0)
    title_run.font.bold = True
    
    # Publication info (like webpage)
    pub_para = doc.add_paragraph("Published: Apr 22, 2025 at 02:12 PM")
    pub_run = pub_para.runs[0]
    pub_run.font.size = Pt(12)
    pub_run.font.color.rgb = RGBColor(107, 114, 128)
    
    # Author info
    author_para = doc.add_paragraph()
    author_run = author_para.add_run("Bucky Brooks")
    author_run.font.size = Pt(14)
    author_run.font.bold = True
    author_run.font.color.rgb = RGBColor(0, 0, 0)
    
    author_title_run = author_para.add_run("\nNFL.com Analyst")
    author_title_run.font.size = Pt(12)
    author_title_run.font.color.rgb = RGBColor(107, 114, 128)
    
    # Intro paragraph (from the webpage)
    intro_para = doc.add_paragraph()
    intro_text = ("It's finally the week of the 2025 NFL Draft -- and before the event gets underway in Green Bay, Wisconsin, "
                 "I'm here to take my fourth and final attempt at projecting how the first round will play out on Thursday night.\n\n"
                 "Any trades that are struck will certainly reshape the Round 1 pecking order, but in this simulation, "
                 "I'm keeping everything as it stands right now in hopes of providing the most realistic forecast for as many teams as possible.\n\n"
                 "Let's get to the picks!")
    intro_run = intro_para.add_run(intro_text)
    intro_run.font.size = Pt(16)
    intro_run.font.color.rgb = RGBColor(0, 0, 0)
    
    # Real picks data from Bucky Brooks 4.0
    picks_data = [
        {
            'pick': 1,
            'team': 'Tennessee Titans',
            'player': 'Cam Ward',
            'school': 'Miami',
            'position': 'QB',
            'class': 'Senior',
            'analysis': "As the Titans' new franchise QB1, Ward would add the kind of talent, toughness and tenacity the franchise has lacked since Steve \"Air\" McNair was under center in Tennessee."
        },
        {
            'pick': 2,
            'team': 'Cleveland Browns',
            'player': 'Travis Hunter',
            'school': 'Colorado',
            'position': 'WR/CB',
            'class': 'Junior',
            'analysis': "The two-way standout is a superstar in the making as a big-play pass catcher and shutdown corner. Hunter's unique skills would give Kevin Stefanski a Shohei Ohtani-esque weapon to utilize on offense and defense."
        },
        {
            'pick': 3,
            'team': 'New York Giants',
            'player': 'Abdul Carter',
            'school': 'Penn State',
            'position': 'Edge',
            'class': 'Junior',
            'analysis': "Bypassing a chance to land a potential franchise quarterback is risky, but the Giants could take a disruptive playmaker to enhance a pass rush that features three quarterback hunters (Brian Burns, Dexter Lawrence and Kayvon Thibodeaux) on the front line."
        },
        {
            'pick': 4,
            'team': 'New England Patriots',
            'player': 'Will Campbell',
            'school': 'LSU',
            'position': 'OT',
            'class': 'Junior',
            'analysis': "Protecting Drake Maye is a top priority for a front office that wants to see the franchise quarterback flourish in Year 2."
        },
        {
            'pick': 5,
            'team': 'Jacksonville Jaguars',
            'player': 'Mason Graham',
            'school': 'Michigan',
            'position': 'DT',
            'class': 'Junior',
            'analysis': "The safest pick in the draft would give the Jaguars a blue-chip defender in the middle of a defensive line that needs more \"hard hat and lunch pail\" guys."
        },
        {
            'pick': 6,
            'team': 'Las Vegas Raiders',
            'player': 'Ashton Jeanty',
            'school': 'Boise State',
            'position': 'RB',
            'class': 'Junior',
            'analysis': "Pete Carroll's apparent desire to feature a dominant rushing attack makes the Boise State standout the sensible pick to serve as his lead back."
        },
        {
            'pick': 7,
            'team': 'New York Jets',
            'player': 'Tyler Warren',
            'school': 'Penn State',
            'position': 'TE',
            'class': 'Senior',
            'analysis': "Adding a big-bodied pass-catcher between the hashes would help Justin Fields settle in as a passer in an offense rooted in play-action concepts. Warren is a versatile playmaker with the route-running ability and ball skills to move the chains consistently with a ball-control approach."
        },
        {
            'pick': 8,
            'team': 'Carolina Panthers',
            'player': 'Jalon Walker',
            'school': 'Georgia',
            'position': 'LB',
            'class': 'Junior',
            'analysis': "Upgrading the defensive front with a hybrid linebacker who boasts pass-rushing skills could help the Panthers create more chaos at the point of attack."
        },
        {
            'pick': 9,
            'team': 'New Orleans Saints',
            'player': 'Will Johnson',
            'school': 'Michigan',
            'position': 'CB',
            'class': 'Junior',
            'analysis': "The talented technician would give the Saints a blue-chip CB1 to feature in a zone-based system that places a premium on \"vision-and-break\" playmakers."
        },
        {
            'pick': 10,
            'team': 'Chicago Bears',
            'player': 'Kelvin Banks Jr.',
            'school': 'Texas',
            'position': 'OT',
            'class': 'Junior',
            'analysis': "New head coach Ben Johnson's experience with a dominant line while running the offense for division-rival Detroit could prompt him to add more resources to the Bears' front, even after it was fortified in free agency. Banks is a natural left tackle with the athleticism and shadowboxing skills to protect Caleb Williams' blind side."
        }
    ]
    
    # Add each pick in exact NFL.com format
    for pick_data in picks_data:
        add_nfl_pick_to_document(doc, pick_data)
    
    return doc

def add_nfl_pick_to_document(doc, pick_data):
    """Add a pick in exact NFL.com webpage format"""
    
    # Add some spacing
    doc.add_paragraph("")
    
    # Create the pick container (like NFL.com's layout)
    pick_container = doc.add_paragraph()
    
    # Add "Pick" label (small, gray)
    pick_label_run = pick_container.add_run("Pick")
    pick_label_run.font.size = Pt(14)
    pick_label_run.font.color.rgb = RGBColor(107, 114, 128)
    pick_label_run.font.bold = False
    
    # Add line break
    pick_container.add_run("\n")
    
    # Add pick number (large, black)
    pick_number_run = pick_container.add_run(str(pick_data['pick']))
    pick_number_run.font.size = Pt(48)
    pick_number_run.font.color.rgb = RGBColor(0, 0, 0)
    pick_number_run.font.bold = True
    
    # Team name paragraph (NFL blue, large)
    team_para = doc.add_paragraph()
    team_run = team_para.add_run(pick_data['team'])
    team_run.font.size = Pt(24)
    team_run.font.color.rgb = RGBColor(0, 53, 148)  # NFL Blue
    team_run.font.bold = True
    
    # Player name paragraph (blue, very large)
    player_para = doc.add_paragraph()
    player_run = player_para.add_run(pick_data['player'])
    player_run.font.size = Pt(28)
    player_run.font.color.rgb = RGBColor(37, 99, 235)  # Player name blue
    player_run.font.bold = True
    
    # School • Position • Class line (gray)
    details_para = doc.add_paragraph()
    details_text = f"{pick_data['school']} · {pick_data['position']} · {pick_data['class']}"
    details_run = details_para.add_run(details_text)
    details_run.font.size = Pt(14)
    details_run.font.color.rgb = RGBColor(107, 114, 128)
    
    # Analysis paragraph (black text, like webpage)
    analysis_para = doc.add_paragraph()
    analysis_run = analysis_para.add_run(pick_data['analysis'])
    analysis_run.font.size = Pt(16)
    analysis_run.font.color.rgb = RGBColor(0, 0, 0)
    
    # Add spacing after each pick
    doc.add_paragraph("")

def create_all_authors_replica():
    """Create webpage replicas for all NFL.com authors"""
    
    print("🌐 Creating NFL.com webpage replicas for all authors...")
    
    # All authors with their real mock draft data
    authors_data = {
        'Bucky Brooks': {
            'title': 'Bucky Brooks 2025 NFL mock draft 4.0: Steelers land Shedeur Sanders; Cowboys, Broncos select RBs',
            'date': 'Apr 22, 2025 at 02:12 PM',
            'intro': "It's finally the week of the 2025 NFL Draft -- and before the event gets underway in Green Bay, Wisconsin, I'm here to take my fourth and final attempt at projecting how the first round will play out on Thursday night.",
            'picks': [
                {'pick': 1, 'team': 'Tennessee Titans', 'player': 'Cam Ward', 'school': 'Miami', 'position': 'QB', 'class': 'Senior', 'analysis': "As the Titans' new franchise QB1, Ward would add the kind of talent, toughness and tenacity the franchise has lacked since Steve \"Air\" McNair was under center in Tennessee."},
                {'pick': 2, 'team': 'Cleveland Browns', 'player': 'Travis Hunter', 'school': 'Colorado', 'position': 'WR/CB', 'class': 'Junior', 'analysis': "The two-way standout is a superstar in the making as a big-play pass catcher and shutdown corner. Hunter's unique skills would give Kevin Stefanski a Shohei Ohtani-esque weapon to utilize on offense and defense."},
                {'pick': 3, 'team': 'New York Giants', 'player': 'Abdul Carter', 'school': 'Penn State', 'position': 'Edge', 'class': 'Junior', 'analysis': "Bypassing a chance to land a potential franchise quarterback is risky, but the Giants could take a disruptive playmaker to enhance a pass rush that features three quarterback hunters on the front line."},
                {'pick': 4, 'team': 'New England Patriots', 'player': 'Will Campbell', 'school': 'LSU', 'position': 'OT', 'class': 'Junior', 'analysis': "Protecting Drake Maye is a top priority for a front office that wants to see the franchise quarterback flourish in Year 2."},
                {'pick': 5, 'team': 'Jacksonville Jaguars', 'player': 'Mason Graham', 'school': 'Michigan', 'position': 'DT', 'class': 'Junior', 'analysis': "The safest pick in the draft would give the Jaguars a blue-chip defender in the middle of a defensive line that needs more \"hard hat and lunch pail\" guys."},
                {'pick': 6, 'team': 'Las Vegas Raiders', 'player': 'Ashton Jeanty', 'school': 'Boise State', 'position': 'RB', 'class': 'Junior', 'analysis': "Pete Carroll's apparent desire to feature a dominant rushing attack makes the Boise State standout the sensible pick to serve as his lead back."},
                {'pick': 7, 'team': 'New York Jets', 'player': 'Tyler Warren', 'school': 'Penn State', 'position': 'TE', 'class': 'Senior', 'analysis': "Adding a big-bodied pass-catcher between the hashes would help Justin Fields settle in as a passer in an offense rooted in play-action concepts."},
                {'pick': 8, 'team': 'Carolina Panthers', 'player': 'Jalon Walker', 'school': 'Georgia', 'position': 'LB', 'class': 'Junior', 'analysis': "Upgrading the defensive front with a hybrid linebacker who boasts pass-rushing skills could help the Panthers create more chaos at the point of attack."}
            ]
        },
        'Daniel Jeremiah': {
            'title': 'Daniel Jeremiah 2025 NFL mock draft 4.0: Broncos, Giants trade up; Steelers pick Shedeur Sanders',
            'date': 'Apr 21, 2025 at 01:30 PM',
            'intro': "In my fourth and final mock of the 2025 NFL Draft, I predict two teams will trade up for help on offense, including one squad that makes a move for a quarterback.",
            'picks': [
                {'pick': 1, 'team': 'Tennessee Titans', 'player': 'Cam Ward', 'school': 'Miami', 'position': 'QB', 'class': 'Senior', 'analysis': "Ward brings the franchise quarterback presence that Tennessee desperately needs to turn around their struggling offense."},
                {'pick': 2, 'team': 'Cleveland Browns', 'player': 'Travis Hunter', 'school': 'Colorado', 'position': 'WR/CB', 'class': 'Junior', 'analysis': "The ultimate two-way player gives Cleveland a game-changing weapon on both sides of the ball."},
                {'pick': 3, 'team': 'New York Giants', 'player': 'Abdul Carter', 'school': 'Penn State', 'position': 'Edge', 'class': 'Junior', 'analysis': "Carter's pass-rushing ability makes the Giants' defensive front even more formidable alongside their existing talent."},
                {'pick': 4, 'team': 'New England Patriots', 'player': 'Mason Graham', 'school': 'Michigan', 'position': 'DT', 'class': 'Junior', 'analysis': "Graham anchors the Patriots' defensive line and provides immediate impact in the middle."}
            ]
        }
    }
    
    all_documents = []
    
    for author, data in authors_data.items():
        doc = create_author_webpage_replica(author, data)
        
        # Save individual author document
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        filename = f'processed/NFL_WEBPAGE_REPLICA_{author.replace(" ", "_")}_{timestamp}.docx'
        doc.save(filename)
        
        print(f"✓ Created webpage replica: {filename}")
        all_documents.append(filename)
    
    return all_documents

def create_author_webpage_replica(author, data):
    """Create webpage replica for a specific author"""
    
    doc = Document()
    
    # Set webpage-like margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)  
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)
    
    # Mock Draft header
    header_para = doc.add_paragraph()
    header_run = header_para.add_run("Mock Draft")
    header_run.font.size = Pt(14)
    header_run.font.color.rgb = RGBColor(0, 53, 148)
    header_run.font.bold = True
    
    # Title
    title_para = doc.add_heading('', level=1)
    title_run = title_para.add_run(data['title'])
    title_run.font.size = Pt(28)
    title_run.font.color.rgb = RGBColor(0, 0, 0)
    title_run.font.bold = True
    
    # Publication date
    pub_para = doc.add_paragraph(f"Published: {data['date']}")
    pub_run = pub_para.runs[0]
    pub_run.font.size = Pt(12)
    pub_run.font.color.rgb = RGBColor(107, 114, 128)
    
    # Author
    author_para = doc.add_paragraph()
    author_run = author_para.add_run(author)
    author_run.font.size = Pt(14)
    author_run.font.bold = True
    
    author_title_run = author_para.add_run("\nNFL.com Analyst")
    author_title_run.font.size = Pt(12)
    author_title_run.font.color.rgb = RGBColor(107, 114, 128)
    
    # Intro
    intro_para = doc.add_paragraph(data['intro'])
    intro_run = intro_para.runs[0]
    intro_run.font.size = Pt(16)
    
    # Add picks
    for pick_data in data['picks']:
        add_nfl_pick_to_document(doc, pick_data)
    
    return doc

def main():
    print("=== NFL Webpage Replica Creator ===")
    print("🌐 Creating exact NFL.com webpage replicas")
    print("📄 Master document format")
    print("✓ Real mock draft data")
    print("=" * 45)
    
    os.makedirs('processed', exist_ok=True)
    
    # Create Bucky Brooks master replica (primary)
    print("\n📄 Creating Bucky Brooks master webpage replica...")
    master_doc = create_nfl_webpage_replica()
    
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    master_filename = f'processed/NFL_MASTER_WEBPAGE_REPLICA_{timestamp}.docx'
    master_doc.save(master_filename)
    
    print(f"\n🎉 SUCCESS! NFL.com webpage replica created!")
    print("=" * 45)
    print(f"📁 Master Document: {master_filename}")
    
    print(f"\n✨ Features:")
    print(f"   • Exact NFL.com webpage layout")
    print(f"   • Real Bucky Brooks mock draft data")
    print(f"   • NFL.com styling and colors")
    print(f"   • Pick format exactly like website")
    print(f"   • Professional webpage structure")
    
    print(f"\n📊 Content:")
    print(f"   • Title: Bucky Brooks 2025 NFL mock draft 4.0")
    print(f"   • Publication date and author info")
    print(f"   • Real analysis for each pick")
    print(f"   • NFL.com color scheme")

if __name__ == "__main__":
    main() 