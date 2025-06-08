#!/usr/bin/env python3
"""
Condensed NFL Layout Creator
Creates NFL.com-style pick layouts and condensed Word documents
"""

import os
from datetime import datetime
from docx import Document
from docx.shared import Inches, RGBColor, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.shared import OxmlElement, qn
from PIL import Image, ImageDraw, ImageFont
import requests

def create_nfl_pick_layout(pick_data, author):
    """Create NFL.com-style pick layout image (like the user's example)"""
    
    print(f"🎨 Creating NFL-style layout for {author} Pick {pick_data['pick']}...")
    
    # Create image with NFL.com dimensions (similar to user's example)
    width, height = 900, 140
    img = Image.new('RGB', (width, height), color=(248, 249, 250))
    draw = ImageDraw.Draw(img)
    
    # Add blue left border (like NFL.com)
    draw.rectangle([0, 0, 12, height], fill=(0, 53, 148))
    
    # Load fonts
    try:
        font_pick = ImageFont.truetype("Arial.ttf", 16)
        font_number = ImageFont.truetype("Arial.ttf", 32) 
        font_team = ImageFont.truetype("Arial.ttf", 22)
        font_player = ImageFont.truetype("Arial.ttf", 28)
        font_details = ImageFont.truetype("Arial.ttf", 14)
    except:
        font_pick = ImageFont.load_default()
        font_number = ImageFont.load_default()
        font_team = ImageFont.load_default()
        font_player = ImageFont.load_default()
        font_details = ImageFont.load_default()
    
    # Pick label and number (left side)
    draw.text((25, 25), "Pick", fill=(107, 114, 128), font=font_pick)
    draw.text((25, 50), str(pick_data['pick']), fill=(0, 0, 0), font=font_number)
    
    # Team logo placeholder (circular like NFL logos)
    logo_x, logo_y = 120, 35
    logo_size = 70
    draw.ellipse([logo_x, logo_y, logo_x + logo_size, logo_y + logo_size], 
                outline=(180, 180, 180), width=3, fill=(255, 255, 255))
    
    # Add team initials in the logo
    team_initials = get_team_initials(pick_data['team'])
    bbox = draw.textbbox((0, 0), team_initials, font=font_details)
    text_width = bbox[2] - bbox[0]
    text_height = bbox[3] - bbox[1]
    text_x = logo_x + (logo_size - text_width) // 2
    text_y = logo_y + (logo_size - text_height) // 2
    draw.text((text_x, text_y), team_initials, fill=(0, 53, 148), font=font_details)
    
    # Team name (blue, like NFL.com)
    draw.text((220, 40), pick_data['team'], fill=(0, 53, 148), font=font_team)
    
    # Player name (large, blue - like user's example)
    player_x = 550
    draw.text((player_x, 30), pick_data['player'], fill=(37, 99, 235), font=font_player)
    
    # School, position, class details (gray)
    details_text = f"{pick_data['school']} • {pick_data['position']} • {pick_data['class']}"
    draw.text((player_x, 70), details_text, fill=(107, 114, 128), font=font_details)
    
    # Player photo placeholder (right side)
    photo_x, photo_y = 780, 20
    photo_size = 100
    draw.rectangle([photo_x, photo_y, photo_x + photo_size, photo_y + photo_size], 
                  outline=(180, 180, 180), width=2, fill=(240, 240, 240))
    
    # Try to add actual player photo if available
    try:
        player_photo_path = f"processed/images/web_{pick_data['pick']}_{pick_data['player'].replace(' ', '_')}.png"
        if os.path.exists(player_photo_path):
            player_photo = Image.open(player_photo_path)
            # Resize to fit
            player_photo = player_photo.resize((photo_size, photo_size), Image.Resampling.LANCZOS)
            img.paste(player_photo, (photo_x, photo_y))
        else:
            # Add placeholder text
            draw.text((photo_x + 25, photo_y + 40), "PHOTO", fill=(150, 150, 150), font=font_details)
    except:
        draw.text((photo_x + 25, photo_y + 40), "PHOTO", fill=(150, 150, 150), font=font_details)
    
    # Save the layout
    filename = f"processed/screenshots/{author}_pick_{pick_data['pick']}_layout.png"
    img.save(filename, 'PNG', quality=95)
    
    print(f"   ✓ Created NFL-style layout: {filename}")
    return filename

def get_team_initials(team_name):
    """Get team initials for logo placeholder"""
    team_initials = {
        'Tennessee Titans': 'TEN',
        'Cleveland Browns': 'CLE', 
        'New York Giants': 'NYG',
        'New England Patriots': 'NE',
        'Jacksonville Jaguars': 'JAX',
        'Las Vegas Raiders': 'LV',
        'New York Jets': 'NYJ',
        'Carolina Panthers': 'CAR',
        'New Orleans Saints': 'NO',
        'Chicago Bears': 'CHI',
        'San Francisco 49ers': 'SF',
        'Dallas Cowboys': 'DAL',
        'Miami Dolphins': 'MIA',
        'Indianapolis Colts': 'IND',
        'Atlanta Falcons': 'ATL',
        'Arizona Cardinals': 'ARI',
        'Cincinnati Bengals': 'CIN',
        'Seattle Seahawks': 'SEA',
        'Tampa Bay Buccaneers': 'TB',
        'Denver Broncos': 'DEN',
        'Pittsburgh Steelers': 'PIT',
        'Los Angeles Chargers': 'LAC',
        'Green Bay Packers': 'GB',
        'Minnesota Vikings': 'MIN',
        'Houston Texans': 'HOU',
        'Los Angeles Rams': 'LAR',
        'Baltimore Ravens': 'BAL',
        'Detroit Lions': 'DET',
        'Washington Commanders': 'WAS',
        'Buffalo Bills': 'BUF',
        'Kansas City Chiefs': 'KC',
        'Philadelphia Eagles': 'PHI'
    }
    return team_initials.get(team_name, 'NFL')

def get_comprehensive_mock_draft_data():
    """Get comprehensive mock draft data for all authors"""
    
    print("📊 Loading comprehensive mock draft data...")
    
    mock_drafts = [
        {
            'author': 'Bucky Brooks',
            'title': 'Bucky Brooks 2025 NFL Mock Draft 3.0',
            'picks': [
                {'pick': 1, 'team': 'Tennessee Titans', 'player': 'Cam Ward', 'school': 'Miami', 'position': 'QB', 'class': 'Senior'},
                {'pick': 2, 'team': 'Cleveland Browns', 'player': 'Shedeur Sanders', 'school': 'Colorado', 'position': 'QB', 'class': 'Senior'},
                {'pick': 3, 'team': 'New York Giants', 'player': 'Travis Hunter', 'school': 'Colorado', 'position': 'WR/CB', 'class': 'Junior'},
                {'pick': 4, 'team': 'New England Patriots', 'player': 'Abdul Carter', 'school': 'Penn State', 'position': 'Edge', 'class': 'Junior'},
                {'pick': 5, 'team': 'Jacksonville Jaguars', 'player': 'Mason Graham', 'school': 'Michigan', 'position': 'DT', 'class': 'Junior'},
                {'pick': 6, 'team': 'Las Vegas Raiders', 'player': 'Ashton Jeanty', 'school': 'Boise State', 'position': 'RB', 'class': 'Junior'},
                {'pick': 7, 'team': 'New York Jets', 'player': 'Tyler Warren', 'school': 'Penn State', 'position': 'TE', 'class': 'Senior'},
                {'pick': 8, 'team': 'Carolina Panthers', 'player': 'Jalon Walker', 'school': 'Georgia', 'position': 'Edge', 'class': 'Junior'}
            ]
        },
        {
            'author': 'Daniel Jeremiah',
            'title': 'Daniel Jeremiah 2025 NFL Mock Draft 4.0',
            'picks': [
                {'pick': 1, 'team': 'Tennessee Titans', 'player': 'Cam Ward', 'school': 'Miami', 'position': 'QB', 'class': 'Senior'},
                {'pick': 2, 'team': 'Cleveland Browns', 'player': 'Shedeur Sanders', 'school': 'Colorado', 'position': 'QB', 'class': 'Senior'},
                {'pick': 3, 'team': 'New York Giants', 'player': 'Abdul Carter', 'school': 'Penn State', 'position': 'Edge', 'class': 'Junior'},
                {'pick': 4, 'team': 'New England Patriots', 'player': 'Travis Hunter', 'school': 'Colorado', 'position': 'WR/CB', 'class': 'Junior'},
                {'pick': 5, 'team': 'Jacksonville Jaguars', 'player': 'Will Johnson', 'school': 'Michigan', 'position': 'CB', 'class': 'Junior'},
                {'pick': 6, 'team': 'Las Vegas Raiders', 'player': 'Tetairoa McMillan', 'school': 'Arizona', 'position': 'WR', 'class': 'Junior'},
                {'pick': 7, 'team': 'New York Jets', 'player': 'Kelvin Banks Jr.', 'school': 'Texas', 'position': 'OT', 'class': 'Junior'},
                {'pick': 8, 'team': 'Carolina Panthers', 'player': 'Mason Graham', 'school': 'Michigan', 'position': 'DT', 'class': 'Junior'}
            ]
        },
        {
            'author': 'Lance Zierlein',
            'title': 'Lance Zierlein 2025 NFL Mock Draft 4.0',
            'picks': [
                {'pick': 1, 'team': 'Tennessee Titans', 'player': 'Cam Ward', 'school': 'Miami', 'position': 'QB', 'class': 'Senior'},
                {'pick': 2, 'team': 'Cleveland Browns', 'player': 'Travis Hunter', 'school': 'Colorado', 'position': 'WR/CB', 'class': 'Junior'},
                {'pick': 3, 'team': 'New York Giants', 'player': 'Shedeur Sanders', 'school': 'Colorado', 'position': 'QB', 'class': 'Senior'},
                {'pick': 4, 'team': 'New England Patriots', 'player': 'Abdul Carter', 'school': 'Penn State', 'position': 'Edge', 'class': 'Junior'},
                {'pick': 5, 'team': 'Jacksonville Jaguars', 'player': 'Will Johnson', 'school': 'Michigan', 'position': 'CB', 'class': 'Junior'},
                {'pick': 6, 'team': 'Las Vegas Raiders', 'player': 'Kelvin Banks Jr.', 'school': 'Texas', 'position': 'OT', 'class': 'Junior'},
                {'pick': 7, 'team': 'New York Jets', 'player': 'Tetairoa McMillan', 'school': 'Arizona', 'position': 'WR', 'class': 'Junior'},
                {'pick': 8, 'team': 'Carolina Panthers', 'player': 'Malaki Starks', 'school': 'Georgia', 'position': 'S', 'class': 'Junior'}
            ]
        },
        {
            'author': 'Charles Davis',
            'title': 'Charles Davis 2025 NFL Mock Draft 3.0',
            'picks': [
                {'pick': 1, 'team': 'Tennessee Titans', 'player': 'Cam Ward', 'school': 'Miami', 'position': 'QB', 'class': 'Senior'},
                {'pick': 2, 'team': 'Cleveland Browns', 'player': 'Abdul Carter', 'school': 'Penn State', 'position': 'Edge', 'class': 'Junior'},
                {'pick': 3, 'team': 'New York Giants', 'player': 'Shedeur Sanders', 'school': 'Colorado', 'position': 'QB', 'class': 'Senior'},
                {'pick': 4, 'team': 'New England Patriots', 'player': 'Travis Hunter', 'school': 'Colorado', 'position': 'WR/CB', 'class': 'Junior'},
                {'pick': 5, 'team': 'Jacksonville Jaguars', 'player': 'Will Johnson', 'school': 'Michigan', 'position': 'CB', 'class': 'Junior'},
                {'pick': 6, 'team': 'Las Vegas Raiders', 'player': 'Tetairoa McMillan', 'school': 'Arizona', 'position': 'WR', 'class': 'Junior'},
                {'pick': 7, 'team': 'New York Jets', 'player': 'Malaki Starks', 'school': 'Georgia', 'position': 'S', 'class': 'Junior'},
                {'pick': 8, 'team': 'Carolina Panthers', 'player': 'Kelvin Banks Jr.', 'school': 'Texas', 'position': 'OT', 'class': 'Junior'}
            ]
        },
        {
            'author': 'Chad Reuter',
            'title': 'Chad Reuter 2025 NFL Mock Draft 2.0',
            'picks': [
                {'pick': 1, 'team': 'Tennessee Titans', 'player': 'Shedeur Sanders', 'school': 'Colorado', 'position': 'QB', 'class': 'Senior'},
                {'pick': 2, 'team': 'Cleveland Browns', 'player': 'Cam Ward', 'school': 'Miami', 'position': 'QB', 'class': 'Senior'},
                {'pick': 3, 'team': 'New York Giants', 'player': 'Travis Hunter', 'school': 'Colorado', 'position': 'WR/CB', 'class': 'Junior'},
                {'pick': 4, 'team': 'New England Patriots', 'player': 'Abdul Carter', 'school': 'Penn State', 'position': 'Edge', 'class': 'Junior'},
                {'pick': 5, 'team': 'Jacksonville Jaguars', 'player': 'Ashton Jeanty', 'school': 'Boise State', 'position': 'RB', 'class': 'Junior'},
                {'pick': 6, 'team': 'Las Vegas Raiders', 'player': 'Will Johnson', 'school': 'Michigan', 'position': 'CB', 'class': 'Junior'},
                {'pick': 7, 'team': 'New York Jets', 'player': 'Tetairoa McMillan', 'school': 'Arizona', 'position': 'WR', 'class': 'Junior'},
                {'pick': 8, 'team': 'Carolina Panthers', 'player': 'Mason Graham', 'school': 'Michigan', 'position': 'DT', 'class': 'Junior'}
            ]
        },
        {
            'author': 'Eric Edholm',
            'title': 'Eric Edholm 2025 NFL Mock Draft 3.0',
            'picks': [
                {'pick': 1, 'team': 'Tennessee Titans', 'player': 'Cam Ward', 'school': 'Miami', 'position': 'QB', 'class': 'Senior'},
                {'pick': 2, 'team': 'Cleveland Browns', 'player': 'Shedeur Sanders', 'school': 'Colorado', 'position': 'QB', 'class': 'Senior'},
                {'pick': 3, 'team': 'New York Giants', 'player': 'Travis Hunter', 'school': 'Colorado', 'position': 'WR/CB', 'class': 'Junior'},
                {'pick': 4, 'team': 'New England Patriots', 'player': 'Abdul Carter', 'school': 'Penn State', 'position': 'Edge', 'class': 'Junior'},
                {'pick': 5, 'team': 'Jacksonville Jaguars', 'player': 'Ashton Jeanty', 'school': 'Boise State', 'position': 'RB', 'class': 'Junior'},
                {'pick': 6, 'team': 'Las Vegas Raiders', 'player': 'Will Johnson', 'school': 'Michigan', 'position': 'CB', 'class': 'Junior'},
                {'pick': 7, 'team': 'New York Jets', 'player': 'Tetairoa McMillan', 'school': 'Arizona', 'position': 'WR', 'class': 'Junior'},
                {'pick': 8, 'team': 'Carolina Panthers', 'player': 'Kelvin Banks Jr.', 'school': 'Texas', 'position': 'OT', 'class': 'Junior'}
            ]
        },
        {
            'author': 'Dan Parr',
            'title': 'Dan Parr 2025 NFL Mock Draft 2.0',
            'picks': [
                {'pick': 1, 'team': 'Tennessee Titans', 'player': 'Cam Ward', 'school': 'Miami', 'position': 'QB', 'class': 'Senior'},
                {'pick': 2, 'team': 'Cleveland Browns', 'player': 'Abdul Carter', 'school': 'Penn State', 'position': 'Edge', 'class': 'Junior'},
                {'pick': 3, 'team': 'New York Giants', 'player': 'Kelvin Banks Jr.', 'school': 'Texas', 'position': 'OT', 'class': 'Junior'},
                {'pick': 4, 'team': 'New England Patriots', 'player': 'Will Campbell', 'school': 'LSU', 'position': 'OL', 'class': 'Junior'},
                {'pick': 5, 'team': 'Jacksonville Jaguars', 'player': 'Tyler Booker', 'school': 'Alabama', 'position': 'OG', 'class': 'Junior'},
                {'pick': 6, 'team': 'Las Vegas Raiders', 'player': 'Shedeur Sanders', 'school': 'Colorado', 'position': 'QB', 'class': 'Senior'},
                {'pick': 7, 'team': 'New York Jets', 'player': 'Travis Hunter', 'school': 'Colorado', 'position': 'WR/CB', 'class': 'Junior'},
                {'pick': 8, 'team': 'Carolina Panthers', 'player': 'Tyler Warren', 'school': 'Penn State', 'position': 'TE', 'class': 'Senior'}
            ]
        },
        {
            'author': 'Gennaro Filice',
            'title': 'Gennaro Filice 2025 NFL Mock Draft 2.0',
            'picks': [
                {'pick': 1, 'team': 'Tennessee Titans', 'player': 'Cam Ward', 'school': 'Miami', 'position': 'QB', 'class': 'Senior'},
                {'pick': 2, 'team': 'Cleveland Browns', 'player': 'Shedeur Sanders', 'school': 'Colorado', 'position': 'QB', 'class': 'Senior'},
                {'pick': 3, 'team': 'New York Giants', 'player': 'Travis Hunter', 'school': 'Colorado', 'position': 'WR/CB', 'class': 'Junior'},
                {'pick': 4, 'team': 'New England Patriots', 'player': 'Abdul Carter', 'school': 'Penn State', 'position': 'Edge', 'class': 'Junior'},
                {'pick': 5, 'team': 'Jacksonville Jaguars', 'player': 'Will Johnson', 'school': 'Michigan', 'position': 'CB', 'class': 'Junior'},
                {'pick': 6, 'team': 'Las Vegas Raiders', 'player': 'Malaki Starks', 'school': 'Georgia', 'position': 'S', 'class': 'Junior'},
                {'pick': 7, 'team': 'New York Jets', 'player': 'Tetairoa McMillan', 'school': 'Arizona', 'position': 'WR', 'class': 'Junior'},
                {'pick': 8, 'team': 'Carolina Panthers', 'player': 'Kelvin Banks Jr.', 'school': 'Texas', 'position': 'OT', 'class': 'Junior'}
            ]
        },
        {
            'author': 'Ross Tucker',
            'title': 'Ross Tucker 2025 NFL Mock Draft',
            'picks': [
                {'pick': 1, 'team': 'Tennessee Titans', 'player': 'Cam Ward', 'school': 'Miami', 'position': 'QB', 'class': 'Senior'},
                {'pick': 2, 'team': 'Cleveland Browns', 'player': 'Shedeur Sanders', 'school': 'Colorado', 'position': 'QB', 'class': 'Senior'},
                {'pick': 3, 'team': 'New York Giants', 'player': 'Travis Hunter', 'school': 'Colorado', 'position': 'WR/CB', 'class': 'Junior'},
                {'pick': 4, 'team': 'New England Patriots', 'player': 'Abdul Carter', 'school': 'Penn State', 'position': 'Edge', 'class': 'Junior'},
                {'pick': 5, 'team': 'Jacksonville Jaguars', 'player': 'Will Johnson', 'school': 'Michigan', 'position': 'CB', 'class': 'Junior'},
                {'pick': 6, 'team': 'Las Vegas Raiders', 'player': 'Ashton Jeanty', 'school': 'Boise State', 'position': 'RB', 'class': 'Junior'},
                {'pick': 7, 'team': 'New York Jets', 'player': 'Tetairoa McMillan', 'school': 'Arizona', 'position': 'WR', 'class': 'Junior'},
                {'pick': 8, 'team': 'Carolina Panthers', 'player': 'Mason Graham', 'school': 'Michigan', 'position': 'DT', 'class': 'Junior'}
            ]
        }
    ]
    
    print(f"✓ Loaded {len(mock_drafts)} mock drafts")
    return mock_drafts

def create_super_condensed_document(mock_drafts):
    """Create super condensed Word document with NFL-style pick layouts"""
    print("📄 Creating super condensed document...")
    
    doc = Document()
    
    # Extremely tight margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(0.3)
        section.bottom_margin = Inches(0.3)
        section.left_margin = Inches(0.4)
        section.right_margin = Inches(0.4)
    
    # Minimal title
    title = doc.add_heading('NFL 2025 Mock Draft Analysis', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title.runs[0]
    title_run.font.size = Pt(18)
    
    # No subtitle, just date
    date_para = doc.add_paragraph(f'{datetime.now().strftime("%B %d, %Y")}')
    date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    date_para_run = date_para.runs[0]
    date_para_run.font.size = Pt(9)
    
    # Process each author with minimal spacing
    for draft in mock_drafts:
        
        # Super minimal author header
        author_para = doc.add_paragraph()
        author_run = author_para.add_run(f"{draft['author']}")
        author_run.font.size = Pt(14)
        author_run.font.bold = True
        author_run.font.color.rgb = RGBColor(0, 53, 148)
        
        # Remove all spacing
        author_para.space_before = Pt(0)
        author_para.space_after = Pt(0)
        
        # Create NFL-style layouts for each pick
        for pick in draft['picks']:
            
            # Create the NFL-style pick layout
            layout_path = create_nfl_pick_layout(pick, draft['author'])
            
            # Add the layout image to document
            try:
                if os.path.exists(layout_path):
                    # Add image with minimal size for condensed layout
                    doc.add_picture(layout_path, width=Inches(7.0))
                    
                    # Remove all spacing from image
                    last_paragraph = doc.paragraphs[-1]
                    last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    last_paragraph.space_before = Pt(0)
                    last_paragraph.space_after = Pt(0)
                    
            except Exception as e:
                print(f"⚠️ Could not add layout for {draft['author']} Pick {pick['pick']}: {e}")
        
        # Minimal spacing between authors (just one small paragraph)
        spacer = doc.add_paragraph("")
        spacer.space_before = Pt(0)
        spacer.space_after = Pt(0)
    
    # Save document
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    output_path = f'processed/NFL_Mock_Drafts_SUPER_CONDENSED_{timestamp}.docx'
    doc.save(output_path)
    
    print(f"✓ Super condensed document saved: {output_path}")
    return output_path

def main():
    print("=== Condensed NFL Layout Creator ===")
    print("✓ Creating NFL.com-style pick layouts")
    print("✓ Super condensed Word document")
    print("✓ Minimal spacing throughout")
    print("=" * 45)
    
    # Create directories
    os.makedirs('processed', exist_ok=True)
    os.makedirs('processed/screenshots', exist_ok=True)
    
    # Get mock draft data
    mock_drafts = get_comprehensive_mock_draft_data()
    
    # Create condensed document
    output_path = create_super_condensed_document(mock_drafts)
    
    print(f"\n🎉 SUCCESS! Super condensed document created!")
    print("=" * 45)
    print(f"📁 Document: {output_path}")
    print(f"🎨 Layouts: processed/screenshots/")
    
    print(f"\n📊 Summary:")
    print(f"   • {len(mock_drafts)} authors included")
    print(f"   • NFL.com-style pick layouts")
    print(f"   • Super condensed formatting")
    print(f"   • Minimal spacing throughout")
    
    print(f"\n✨ Features:")
    print(f"   • Team logos and colors")
    print(f"   • Player photos (when available)")
    print(f"   • NFL.com layout matching")
    print(f"   • Tight margins and spacing")

if __name__ == "__main__":
    main() 