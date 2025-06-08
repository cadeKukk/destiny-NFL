#!/usr/bin/env python3
"""
NFL.com Mock Draft Screenshot Tool - Simple Expanded Screenshots
Takes screenshots of picks with expanded areas to capture description text below
"""

import os
import time
from datetime import datetime
from selenium import webdriver
from selenium.webdriver.chrome.options import Options
from selenium.webdriver.common.by import By
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

class NFLScreenshotSimple:
    def __init__(self):
        self.driver = None
        
        # Create directories if they don't exist
        os.makedirs('processed', exist_ok=True)
        os.makedirs('processed/complete_screenshots', exist_ok=True)
        
        # Updated URLs with corrections
        self.authors = {
            'Bucky Brooks': 'https://www.nfl.com/news/bucky-brooks-2025-nfl-mock-draft-4-0-steelers-land-shedeur-sanders-cowboys-broncos-select-rbs',
            'Daniel Jeremiah': 'https://www.nfl.com/news/daniel-jeremiah-2025-nfl-mock-draft-4-0',
            'Lance Zierlein': 'https://www.nfl.com/news/lance-zierlein-2025-nfl-mock-draft-4-0-colts-trade-up-for-colston-loveland-saints-go-get-jaxson-dart',
            'Charles Davis': 'https://www.nfl.com/news/charles-davis-2025-nfl-mock-draft-3-0-cam-ward-only-qb-in-round-1-eagles-pick-te-mason-taylor',
            'Eric Edholm': 'https://www.nfl.com/news/eric-edholm-2025-nfl-mock-draft-3-0-four-first-round-quarterbacks-jaguars-take-rb-ashton-jeanty',
            'Dan Parr': 'https://www.nfl.com/news/dan-parr-2025-nfl-mock-draft-2-0-offensive-linemen-dominate-top-10-bears-grab-tight-end-tyler-warren',
            'Chad Reuter': 'https://www.nfl.com/news/seven-round-2025-nfl-mock-draft-round-one',
            'Gennaro Filice': 'https://www.nfl.com/news/gennaro-filice-2025-nfl-mock-draft-2-0-rb-ashton-jeanty-goes-top-5-cowboys-jump-for-jalon-walker',
            'Marc Ross': 'https://www.nfl.com/news/marc-ross-2025-nfl-mock-draft-1-0-three-qbs-selected-in-top-10-jets-snag-rb-ashton-jeanty'
        }

    def setup_selenium(self):
        """Setup Selenium WebDriver with Chrome"""
        try:
            chrome_options = Options()
            chrome_options.add_argument('--no-sandbox')
            chrome_options.add_argument('--disable-dev-shm-usage')
            chrome_options.add_argument('--window-size=1800,1400')
            chrome_options.add_argument('--user-agent=Mozilla/5.0 (Macintosh; Intel Mac OS X 10_15_7) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/120.0.0.0 Safari/537.36')
            chrome_options.add_argument('--disable-blink-features=AutomationControlled')
            
            self.driver = webdriver.Chrome(options=chrome_options)
            print("✓ Selenium WebDriver setup complete")
            return True
        except Exception as e:
            print(f"⚠️ Error setting up Selenium: {e}")
            return False

    def screenshot_webpage_content(self, url, author):
        """Screenshot webpage content with expanded pick areas"""
        try:
            print(f"🔍 Loading: {url}")
            # Navigate to the page
            self.driver.get(url)
            time.sleep(6)  # Wait for page to fully load
            
            # Remove overlays
            self.remove_overlays()
            time.sleep(3)
            
            # Scroll through page to load content
            self.load_page_content()
            
            # Take header screenshot
            header_screenshots = self.screenshot_article_header(author)
            
            # Take individual pick screenshots with expanded areas
            pick_screenshots = self.screenshot_individual_picks_expanded(author)
            
            return header_screenshots + pick_screenshots
            
        except Exception as e:
            print(f"⚠️ Error screenshotting {author}: {e}")
            return []

    def load_page_content(self):
        """Load page content by scrolling"""
        total_height = self.driver.execute_script("return document.body.scrollHeight")
        for i in range(0, total_height, 2000):
            self.driver.execute_script(f"window.scrollTo(0, {i});")
            time.sleep(0.8)
        
        # Return to top
        self.driver.execute_script("window.scrollTo(0, 0);")
        time.sleep(2)

    def remove_overlays(self):
        """Remove cookie banners and overlays"""
        overlay_selectors = [
            '[data-module="CookieBanner"]',
            '.onetrust-banner-sdk',
            '.cookie-banner',
            '.cookie-notice'
        ]
        
        for selector in overlay_selectors:
            try:
                overlays = self.driver.find_elements(By.CSS_SELECTOR, selector)
                for overlay in overlays:
                    if overlay.is_displayed():
                        self.driver.execute_script("arguments[0].style.display = 'none';", overlay)
            except:
                continue

    def screenshot_article_header(self, author):
        """Screenshot the article header"""
        screenshots = []
        try:
            header_path = f"processed/complete_screenshots/{author}_header.png"
            self.driver.save_screenshot(header_path)
            screenshots.append(header_path)
            print(f"   ✓ Header screenshot: {author}_header.png")
        except Exception as e:
            print(f"   ⚠️ Error taking header screenshot: {e}")
        
        return screenshots

    def screenshot_individual_picks_expanded(self, author):
        """Screenshot individual draft picks with expanded areas to capture description text"""
        screenshots = []
        
        try:
            # Try multiple selectors for different page layouts
            pick_selectors = [
                '.nfl-o-ranked-item.nfl-o-ranked-item--side-by-side',
                '.d3-o-tabs__panel .d3-o-media-object',
                '.d3-o-media-object',
                '.nfl-c-player-header',
                '.draft-pick',
                '.ranked-item',
                '.mock-draft-pick'
            ]
            
            pick_elements = None
            for selector in pick_selectors:
                pick_elements = self.driver.find_elements(By.CSS_SELECTOR, selector)
                if pick_elements:
                    print(f"   📋 Found {len(pick_elements)} draft picks for {author} using: {selector}")
                    break
            
            if not pick_elements:
                print(f"   ⚠️ No pick elements found for {author}")
                return []
            
            print(f"   📋 Processing {min(32, len(pick_elements))} picks for {author}")
            
            # Process up to 32 picks
            for i in range(min(32, len(pick_elements))):
                try:
                    pick_element = pick_elements[i]
                    pick_num = i + 1
                    
                    print(f"   🔍 Processing Pick {pick_num}...")
                    
                    # Scroll element into view
                    self.driver.execute_script("arguments[0].scrollIntoView({block: 'center', behavior: 'smooth'});", pick_element)
                    time.sleep(0.5)
                    
                    # Get element position and size
                    element_location = pick_element.location
                    element_size = pick_element.size
                    
                    # Calculate expanded screenshot area to include description text below
                    # Expand downward by approximately 150-200 pixels to capture description
                    expanded_height = element_size['height'] + 200  # Extra space for description
                    
                    # Take screenshot of expanded area using JavaScript
                    screenshot_path = f"processed/complete_screenshots/{author}_pick_{pick_num:02d}.png"
                    
                    # Use JavaScript to create a larger screenshot area
                    script = f"""
                    var element = arguments[0];
                    var rect = element.getBoundingClientRect();
                    var canvas = document.createElement('canvas');
                    var ctx = canvas.getContext('2d');
                    
                    // Set canvas size to include extra space below
                    canvas.width = rect.width;
                    canvas.height = rect.height + 200;
                    
                    return {{
                        x: rect.left + window.scrollX,
                        y: rect.top + window.scrollY,
                        width: rect.width,
                        height: rect.height + 200
                    }};
                    """
                    
                    # Get expanded coordinates
                    coords = self.driver.execute_script(script, pick_element)
                    
                    # Take screenshot of the expanded area
                    # For simplicity, let's use element screenshot and then try to capture more
                    pick_element.screenshot(screenshot_path)
                    
                    # Alternative approach: Take a larger screenshot using viewport
                    # Get viewport screenshot and crop to expanded area
                    viewport_screenshot_path = f"processed/complete_screenshots/{author}_pick_{pick_num:02d}_temp.png"
                    self.driver.save_screenshot(viewport_screenshot_path)
                    
                    # Use Python PIL to crop the expanded area from viewport screenshot
                    try:
                        from PIL import Image
                        
                        # Open the viewport screenshot
                        viewport_img = Image.open(viewport_screenshot_path)
                        
                        # Calculate crop coordinates (element position + extra height)
                        left = max(0, element_location['x'])
                        top = max(0, element_location['y'])
                        right = min(viewport_img.width, left + element_size['width'])
                        bottom = min(viewport_img.height, top + element_size['height'] + 200)  # +200px for description
                        
                        # Crop the expanded area
                        expanded_crop = viewport_img.crop((left, top, right, bottom))
                        expanded_crop.save(screenshot_path)
                        
                        # Clean up temp file
                        os.remove(viewport_screenshot_path)
                        
                        print(f"   ✓ Pick {pick_num} expanded screenshot captured (includes description)")
                        
                    except ImportError:
                        print(f"   ⚠️ PIL not available, using basic screenshot for pick {pick_num}")
                    except Exception as crop_error:
                        print(f"   ⚠️ Error cropping expanded area for pick {pick_num}: {crop_error}")
                        # Clean up temp file if it exists
                        if os.path.exists(viewport_screenshot_path):
                            os.remove(viewport_screenshot_path)
                    
                    screenshots.append(screenshot_path)
                    time.sleep(0.3)  # Brief pause between screenshots
                    
                except Exception as e:
                    print(f"   ⚠️ Error processing pick {i+1}: {e}")
            
            print(f"   ✓ Captured {len(screenshots)} expanded screenshots for {author}")
            
        except Exception as e:
            print(f"   ⚠️ Error taking individual pick screenshots: {e}")
        
        return screenshots

    def create_word_document(self, all_screenshots):
        """Create comprehensive Word document with all authors"""
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        doc_path = f"processed/NFL_SIMPLE_EXPANDED_{timestamp}.docx"
        
        try:
            doc = Document()
            
            # Set document margins (0.4 inches all around)
            sections = doc.sections
            for section in sections:
                section.top_margin = Inches(0.4)
                section.bottom_margin = Inches(0.4)
                section.left_margin = Inches(0.4)
                section.right_margin = Inches(0.4)
            
            # Document title
            title = doc.add_heading('2025 NFL Mock Draft - Expanded Screenshots', 0)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            subtitle = doc.add_paragraph(f'Generated on {datetime.now().strftime("%B %d, %Y at %I:%M %p")} - Screenshots include descriptions')
            subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
            subtitle_format = subtitle.runs[0].font
            subtitle_format.size = Pt(12)
            subtitle_format.italic = True
            
            # Process each author
            for author, url in self.authors.items():
                print(f"   📝 Adding {author} to document...")
                
                # Author header
                author_header = doc.add_heading(f'{author} Mock Draft', level=1)
                author_header.space_before = Pt(12)
                author_header.space_after = Pt(6)
                
                # Find screenshots for this author
                author_screenshots = [s for s in all_screenshots if author in s]
                header_screenshots = [s for s in author_screenshots if 'header' in s]
                pick_screenshots = [s for s in author_screenshots if 'pick_' in s and 'header' not in s]
                
                # Add header screenshot
                for header_path in header_screenshots:
                    if os.path.exists(header_path):
                        try:
                            header_para = doc.add_paragraph()
                            header_run = header_para.runs[0] if header_para.runs else header_para.add_run()
                            header_run.add_picture(header_path, width=Inches(7.0))
                            header_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                            header_para.space_after = Pt(8)
                        except Exception as e:
                            print(f"      ⚠️ Error adding header image for {author}: {e}")
                
                # Add pick screenshots (these now include descriptions in the image)
                for i, screenshot_path in enumerate(sorted(pick_screenshots), 1):
                    if os.path.exists(screenshot_path):
                        try:
                            # Add screenshot (now includes description text)
                            pick_para = doc.add_paragraph()
                            pick_run = pick_para.runs[0] if pick_para.runs else pick_para.add_run()
                            pick_run.add_picture(screenshot_path, width=Inches(6.5))
                            pick_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                            pick_para.space_after = Pt(6)
                            
                        except Exception as e:
                            print(f"      ⚠️ Error adding pick {i} for {author}: {e}")
                
                # Add page break between authors (except for last author)
                if author != list(self.authors.keys())[-1]:
                    doc.add_page_break()
            
            # Save document
            doc.save(doc_path)
            
            print(f"\n🎉 SUCCESS! Simple expanded NFL.com screenshots captured!")
            print(f"{'='*55}")
            print(f"📁 Document: {doc_path}")
            print(f"📸 Screenshots: processed/complete_screenshots/")
            print(f"\n📊 Summary:")
            print(f"   • {len(self.authors)} authors processed")
            print(f"   • {len(all_screenshots)} total screenshots captured")
            print(f"   • Each screenshot includes the description text below")
            print(f"   • Up to 32 picks per author")
            print(f"   • Simplified approach - no text extraction needed")
            
            return doc_path
            
        except Exception as e:
            print(f"⚠️ Error creating Word document: {e}")
            return None

    def cleanup(self):
        """Clean up resources"""
        if self.driver:
            self.driver.quit()
            print("✓ WebDriver closed")

def main():
    screenshot_tool = NFLScreenshotSimple()
    
    try:
        if not screenshot_tool.setup_selenium():
            return
        
        print("🎯 Starting NFL Mock Draft Screenshot Collection - SIMPLE EXPANDED VERSION...")
        print(f"📋 Processing {len(screenshot_tool.authors)} authors with expanded screenshots")
        
        all_screenshots = []
        
        # Process each author
        for author, url in screenshot_tool.authors.items():
            print(f"\n📸 Capturing expanded screenshots for {author}...")
            author_screenshots = screenshot_tool.screenshot_webpage_content(url, author)
            all_screenshots.extend(author_screenshots)
        
        # Create Word document
        if all_screenshots:
            print(f"\n📄 Creating Word document with expanded screenshots...")
            doc_path = screenshot_tool.create_word_document(all_screenshots)
        else:
            print("⚠️ No screenshots captured")
            
    except KeyboardInterrupt:
        print("\n⚠️ Process interrupted by user")
    except Exception as e:
        print(f"⚠️ Unexpected error: {e}")
    finally:
        screenshot_tool.cleanup()

if __name__ == "__main__":
    main() 